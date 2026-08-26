#!/usr/bin/env python3
"""Create final edited documents from transcripts and slides (pipeline step 7).

Inputs per video:
- transcripts in <playlist>/<ORIG>/ and <playlist>/EN/ (<stem>.txt/.srt),
  where ORIG is the original-language folder (e.g. RU); for English
  originals only EN/ is used;
- slide texts and document structure in <playlist>/SLIDES/<short-key>/
  slides.json (produced by text_from_slides.py); with --no-slides a video
  may instead have <playlist>/INFO/<stem>.json (saved by
  transcribe_videos.py in remote mode) - then the document title is built
  in the summary format (<channel> | <playlist> : <title>) and the
  sections come from the video timecodes (chapters). A chapter titled by
  the author heads its section verbatim; chapters titled by YouTube
  ("<Untitled Chapter 1>", or a lowercase line of the speech) give way to
  the heading the editor writes (chapters_are_youtube_s_own).

Outputs: edited documents in <ORIG>/OUTPUT/ and EN/OUTPUT/ (--orig-only
limits this to <ORIG>/OUTPUT/), three formats each: <stem>.md, <stem>.docx
and <stem>.pdf.

How a document is built:

- The section list (with scene time ranges) comes from slides.json: the
  info slide opens the Foreword, foreword slides open unnumbered sections,
  the agenda slide opens the Agenda section, agenda-matched slides open the
  numbered sections, closing slides are merged into one Conclusion section.
- The boundary between two sections is then moved off the middle of a
  sentence: a slide is switched by hand, mid-speech, so the moment it
  changes is not where the text can be cut. The nearest sentence end is
  taken instead, looked for further back than forward, since a speaker
  announces the next topic before reaching for the slide. When that
  announcement is a whole sentence - a stop, then "Этот слайд такой
  себе..." or "Тогда предлагаю перейти...", and only then the new slide -
  the boundary moves back once more, to the start of it (announced_start).
- The .srt transcripts are sliced by those time ranges, so every section
  gets its original-language and English text fragments, already split
  into paragraphs at the pauses of the recording (see silences.py: the
  .srt has no pauses of its own, whisper butts its segments together).
- Each section is edited by the OpenAI API using both fragments at once:
  recognition errors in the original are fixed against the English
  translation, the text is repunctuated and its grammar repaired (rules
  1b/1c: the recognizer knows two marks and loses endings), the paragraphs
  are kept (or joined and split by rule 1a) and the wording lightly
  polished; awkward English slide wording used for headings is fixed.
  What follows needs no model and is not asked of it: paragraphs over 200
  words are halved at a sentence end, the Russian text gets the letter «ё»
  where only «ё» can stand, and the phrases that are misspelled whatever
  the sentence («имею ввиду») are put right (see polish_section).
- Subsections: when the section text discusses the slide bullets one by
  one in recognizable fragments (e.g. "Benefits of Test Automation"), one
  subsection per bullet is created; when bullets are only mentioned in
  passing and the text does not decompose (e.g. "What Is Test
  Automation?"), the section stays flat. The model decides per section,
  but it is told where each bullet is named in the speech (bullet_anchors,
  matched offline) so that a subsection starts there and not two sentences
  later. A lone subsection on a multi-bullet slide is folded back into the
  intro; a split that covers most of the bullets but not all is asked for
  once more (see close_the_table_of_contents).
- Coverage check: every edited section is diffed (in the original
  language) against the transcript fragment it was made from. Runs of
  words the edit lost - typically a sentence at a section boundary - and
  fragments it emitted twice (in the intro and again in a subsection) make
  the section be edited once more, with those fragments quoted back to the
  model; what is still wrong after that is repaired fragment by fragment.
  A section missing a fifth of itself is a different failure - the model
  edited its first half and summarized the rest away, and asked for the
  whole tail it summarizes it again - so that one is edited once more in
  parts of about 1200 tokens (parts_for_a_losing_section); the parts come
  back flat, and are kept only if they really carry more of the
  transcript. What is still wrong after all that is reported in the log
  and recorded in the cache. Sections cached by an earlier run are re-checked
  offline and re-edited when they fail, so an already-generated document
  can be repaired without paying for the whole video again. A section
  edited under older rules (EDIT_RULES_VERSION) is re-edited as well: a
  document half in one style and half in another is worse than one paid
  for again. Paragraphs that are still over 200 words after the split
  (no sentence end to cut at), and questions written with a full stop, are
  reported after the run - both checks are free.

Formats:
- .md - a Table of Contents with full multi-level numbering (1., 2., 2.1.);
  body lines are at most 80 characters, no hyphenation, justified to the
  right edge by turning some single spaces into double spaces.
- .docx - styles follow the --doc template (default: built-in defaults
  modeled on RU/Docs_Edited/01_Introduction.V2.docx: heading 2 for the
  title, heading 3/4 for sections/subsections, "ds-markdown-paragraph"
  for verbatim slide text). The Table of Contents is a real Word TOC
  field: entries are internal hyperlinks with right-aligned page numbers
  (regardless of the ToC source - slides or timecodes); the file asks
  Word/WPS to update fields on open.
- .pdf - the .docx rendered to PDF via WPS Writer when installed (no
  activation nags, no clashes with a user's Word session), otherwise via
  Microsoft Word; fields (the ToC) are updated before the export. --pdf
  may name a different .docx style template for the PDF rendering.

Costs and safety:
- before editing, the script prints the estimated gpt-4o cost for the video
  and waits for a y/n confirmation (--yes skips the prompt);
- editing results are cached in SLIDES/<key>/edited_sections.json after
  every section, so a re-run after a failure (or a formatting-only change)
  does not pay the API again; delete that file to force fresh editing;
- the PDF conversion runs in a disposable child process: a native Word COM
  crash cannot kill the run, and a Word instance left behind by a crashed
  attempt is killed automatically (the user's own Word is never touched).

Examples:
  python src/docs/create_final_docs.py _Autotesting lectures
  python src/docs/create_final_docs.py _Autotesting lectures --next 3
  python src/docs/create_final_docs.py _Autotesting lectures --doc my_styles.docx
"""

from __future__ import annotations

import argparse
import difflib
import json
import re
import sys
import textwrap
import time
from dataclasses import dataclass, field, replace
from pathlib import Path

_SRC_DIR = Path(__file__).resolve().parents[1]
if str(_SRC_DIR) not in sys.path:
    sys.path.insert(0, str(_SRC_DIR))

from shared.api_key import read_api_key
from shared.glossary import GLOSSARY_FILENAME, find_glossary, load_terms
from shared.media_files import INFO_DIRNAME, list_videos
from shared.openai_chat import (
    MODEL,
    USD_PER_MTOKEN_COMPLETION,
    USD_PER_MTOKEN_PROMPT,
    chat_json,
)
from shared.project_paths import (
    WORKSPACE_ROOT,
    channels_dir,
    require_channel_ref,
)
from shared.settings_file import read_settings
from shared.silences import SilenceIndex, load_silences
from shared.transcripts import (
    LANGUAGE_NAMES,
    Segment,
    ends_sentence,
    group_paragraphs,
)
from slides.extract_slides import SLIDES_DIRNAME, short_slide_keys, slides_out_dir
from slides.text_from_slides import RESULT_FILENAME

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

OUTPUT_DIRNAME = "OUTPUT"
MD_WIDTH = 80
TOC_HEADING = "Table of Contents"
CONCLUSION_HEADINGS = {"EN": "Conclusion", "RU": "Заключение"}
# Heading of the single section when a video has neither slides nor timecodes.
TRANSCRIPT_HEADINGS = {"EN": "Transcript", "RU": "Транскрипт"}

# Editing results are cached next to slides.json, so a re-run after a
# failure (or after a formatting-only change) does not pay the API again.
EDIT_CACHE_FILENAME = "edited_sections.json"
# Bumped when the editing rules change enough that a section edited by the
# old ones would not match a freshly edited neighbour - a document half in
# one style and half in another is worse than paying for the whole of it
# again. The cost is shown and confirmed as usual before anything is sent.
#   2: paragraphs follow the pauses of the recording (see silences.py) and
#      rule 1a of the prompt.
#   3: punctuation and grammar spelled out as rules 1b and 1c.
#   4: rule 1a states the 150/200-word limits as limits, not as aims.
#   5: rule 4a - a subsection starts where its bullet is named.
#   6: rule 4a takes the anchors found for the bullets in the transcript.
#   7: a subsection heading is the name of its bullet, not the whole line;
#      a flat section whose every bullet was named is questioned once.
#   8: that heading is translated into the original language, not copied.
EDIT_RULES_VERSION = 8

# Rough tokenizer ratios for the cost estimate: ~4 chars per token for
# English, ~2.5 for other languages (Cyrillic packs fewer chars per token).
CHARS_PER_TOKEN_EN = 4.0
CHARS_PER_TOKEN_OTHER = 2.5
# Fixed per-section prompt overhead: system prompt + payload wrapper +
# slide title/bullets.
SECTION_OVERHEAD_TOKENS = 900

EDIT_SYSTEM_PROMPT = """\
You edit one section of an auto-transcribed lecture into its final form.
You get the section fragment of the original-language transcript, the
English machine translation of the same fragment, and the slide that opens
the section (title + bullet lines + scene timing).

Editing rules:
1. Original text: fix obvious speech recognition errors (the English
   translation often shows what was meant), especially garbled English
   product, company and technology names; punctuate it properly (see 1b)
   and make every sentence grammatical (see 1c); keep the paragraphs (see
   1a); lightly polish the spoken wording (drop filler words, false starts,
   broken repetitions) while fully preserving the content, details, examples
   and the speaker's tone. Never summarize, never drop content, never add
   content of your own.
1b. Punctuation. The recognizer typed the whole lecture with two marks, the
   full stop and the comma, and put them where it heard a break rather than
   where the sentence needs one. Repunctuate the text: the words stay as
   they are, only the marks change.
   - A dash where speech leaves out the verb or the copula, in parallel
     constructions and before a conclusion: «Хотите – посещайте, хотите –
     нет», «Автоматизация – это не панацея».
   - A colon before a list, or before the explanation of what was just
     said.
   - A question mark on a question - including the rhetorical ones the
     speaker asks and then answers herself: «Что это значит на практике?»,
     «Почему пестицидный парадокс?».
   - A comma, a colon or a dash in place of a full stop that cuts one
     sentence in two - between a cause and its effect, between a clause and
     the one it depends on.
   - A full stop where one sentence ends and the next begins without one.
   - Away with a comma that distorts the sense; in with the one that was
     dropped.
   A sentence that already reads well needs no new marks: do not reword a
   sentence to make room for one.
1c. Grammar. Endings are what the recognizer loses first, so the transcript
   is full of sentences nobody said: «тестировщиков обучают автоматизация»,
   «плюсы, которые дают нам автоматизацию», «какие-то самолётостроения»,
   «сталкивалось» for «сталкивалась». Every sentence must parse in its own
   language: restore the case, gender, number and verb form the sense calls
   for (the English translation shows which reading is meant), the endings
   of inflected loanwords included («adjustment-а», «сценариев»). This
   repairs the recognition, it does not rewrite the speaker.
   When the original language is Russian, write «ё» in every word spelled
   with it: «ещё», «её», «идёт», «даёт», «счёт», «чётко», «серьёзно», and
   «всё» when it is the pronoun; «все» and «чем» keep «е» only where they
   really are those words.
1a. Paragraphs. The blank lines of the given transcript are not arbitrary:
   the text is already split where the speaker stopped talking, so keep
   those breaks as they are. Depart from them only for a reason, and there
   are three:
   - the break falls in the middle of a thought (the sentence after it
     finishes the sentence before it, or answers a question just asked) -
     join the two paragraphs;
   - a paragraph runs past roughly 150 words or turns to a new thought
     halfway - split it, at the place the speaker moves on ("Тогда
     предлагаю перейти...", "Следующий пункт...", "Now let us look at...")
     or right after a list closes ("...и так далее.");
   - a paragraph is one short sentence that plainly belongs with its
     neighbour - join it. A single sentence standing alone is fine only
     when it announces what follows or sums up what came before.
   Aim for paragraphs of 50-120 words: several sentences on one thought,
   not a wall of text and not a stack of one-liners. Two limits are not
   aims but rules: never join two paragraphs into one longer than 150
   words, and never leave a paragraph of more than 200 words standing -
   split it at the best sentence end you can find, and if the recognizer
   left none, at the place your own punctuation ends a sentence.
1d. Order. The document follows the recording. The paragraphs come in the
   order the speaker said them, and a passage is never carried to where it
   would read better, however tempting - a reader watching the video has to
   find the two side by side. Joining and splitting paragraphs (see 1a) is
   the only rearranging there is.
2. English text: edit the same way; its punctuation usually needs little
   work, but polish awkward constructions - the speaker is not a native
   English speaker.
2a. The two transcripts are two recognitions of the same speech, so where
   they disagree the English one often carries what the original lost: a
   whole clause the original recognizer swallowed, or a self-correction
   the speaker made ("...of our country - of our planet"). Restore that
   content in the original language as well: neither version may end up
   saying less than the other, or contradicting it.
2b. Terminology and loanwords. "glossary" lists the names this course uses:
   spell them exactly as listed wherever they occur, in both languages, and
   use them to repair what the recognizer made of them ("PlevRite",
   "Pleuride", "плеврайт" -> "Playwright"; "OCD" -> "CI/CD"). A term that
   also stands on the slide is spelled as on the slide. The loanwords the
   speaker inflects in the original language are their voice, not errors:
   keep them, but write an established loanword the way that language
   already writes it («деплой», «фреймворк»), keep a Latin stem Latin and
   attach the ending after a hyphen («environment-ах», «benefit-ы»), and
   leave a phrase that keeps its English grammar in English in full
   ("course description"). Spell the same term the same way throughout.
3. Keep the two languages parallel: the same paragraph breaks, the same
   subsection boundaries, the same order.
4. Section layout - decide between exactly two layouts:
   (a) flat: no subsections, all text in the intro paragraphs;
   (b) fully decomposed: a short intro (possibly empty) followed by one
       subsection per slide bullet - every bullet, in bullet order -
       together covering all the remaining text. Digressions stay inside
       the        subsection of the bullet being discussed; the last subsection
       extends to the end of the section. Each subsection gets a concise
       English heading derived from its bullet plus a translation into the
       original language: the name of the bullet, two or three words - the
       part before the colon ("Test Isolation", not "Test Isolation: the
       result of one test shouldn't be depended on the another"). The
       original-language heading is that name translated («Изоляция
       тестов»), never a copy of the English one; only a name that is not
       translated on this course stays as it is (Playwright, Smoke).
   Choose (b) only when the bullets announce a list of parallel items
   (features, types, practices, tools) that the speaker then demonstrably
   covers one after another, so that the fragment of every bullet can be
   located confidently. Explanatory slides whose lines merely characterize
   the topic (a "Definition: ... Purpose: ... Example: ..." pattern and the
   like) usually do not structure the speech that follows - such sections
   stay flat even when the opening sentences touch the lines in order. If
   some bullets get no dedicated discussion, or the boundaries are unclear,
   or there are no bullets, choose (a). A partial split (subsections for
   only some bullets) is not allowed. When in doubt, choose (a).
   The intro and the subsections partition the text: every fragment goes
   to exactly one of them. Text moved into a subsection must not stay in
   the intro as well, and two subsections must not retell each other - the
   same sentences twice is an error, not a summary.
4a. Where a subsection starts. At the sentence in which the speaker first
   names its bullet - most often reading it out, in English or in her own
   words ("Test Isolation: the result of one test shouldn't depend on
   another", "И последний поинт здесь - less adaptability to changes") -
   and not a few sentences later, where the explanation gets going. The
   naming sentence opens the subsection; whatever came before it belongs
   to the bullet before, however much it sounds like an introduction.
   "bullet_anchors", when given, has those sentences already found in the
   transcript, mechanically, by the words of the bullet: start the
   subsection of a bullet at its "named_at" sentence. A null "named_at"
   means nothing was matched - not that the bullet was skipped; decide
   that one from the text. Bullets named one after another are the case
   layout (b) is for; a bullet named in the same breath as its neighbours
   ("Есть три раннера: xUnit, NUnit и MSTest") is not a subsection of its
   own. One subsection alone is never a decomposition: give every
   discussed bullet one, or keep the section flat.
   "bullets_without_subsection", when given, lists the bullets your
   previous answer left out while giving the others a subsection. Either
   give each of them its own subsection, moving the text that belongs to
   it out of the intro and the neighbours, or return the whole section
   flat. Do not invent a subsection for a bullet the speaker never
   discussed, and do not repeat text to fill one.
   "every_bullet_was_named", when given, says that each bullet was found
   named in the text, in slide order, and that your previous answer was
   nevertheless flat. This is layout (b): give every bullet its
   subsection, starting at its anchor. Stay flat only if the text really
   does not divide - if it does not, say so by returning it flat again.
5. Propose the final section heading in English: keep the given heading if
   it is fine, otherwise fix grammar or awkward wording; keep it short. Use
   normal title casing, not ALL CAPS.

Reply with JSON only:
{
  "heading": string,
  "intro_orig": [paragraphs in the original language],
  "intro_en": [paragraphs in English],
  "subsections": [
    {"heading_en": string, "heading_orig": string,
     "paragraphs_orig": [strings], "paragraphs_en": [strings]}
  ]
}
"subsections" may be []. When the original language is English, fill only
the *_en fields and leave "intro_orig"/"paragraphs_orig"/"heading_orig" as
empty lists / empty strings. When "transcript_english" is empty, do NOT
translate: edit only the original language and leave "intro_en" and every
"paragraphs_en" as empty lists.

When "section_part" is present ("i/N"), you are editing part i of N of one
long section split only for processing: edit the COMPLETE given fragment
from its first sentence to its last (it may start or end mid-topic - that
is expected; do not add openings or conclusions of your own), keep the
layout flat ("subsections": []) and remember: re-emit ALL the content,
paragraph by paragraph - never compress the fragment.

The section boundaries follow the slides, so the last sentences often
already open the topic of the next slide ("now let us look at the tools on
this slide..."): they belong to this section - edit them, never drop them.

"dropped_fragments" and "repeated_fragments" appear only in a retry: your
previous reply for this very section lost those transcript fragments and
emitted those other ones twice. Edit the section again from scratch,
keeping every dropped fragment in its place in the flow and every repeated
one exactly once (in the subsection it belongs to)."""


FRAGMENT_SYSTEM_PROMPT = """\
You restore one fragment that a previous edit of a lecture transcript
dropped.

You get the fragment as the raw transcript in the original language, the
English transcript of the whole section (the counterpart of the fragment is
somewhere inside it - find it yourself), and the edited paragraphs the
fragment goes between.

Edit the fragment exactly like the text around it: fix obvious speech
recognition errors (spelling the names of "glossary" exactly as listed),
punctuate it (the recognizer knows the full stop and the comma and puts
both where it hears a break: end the sentences where they end, and use the
dash, the colon and the question mark where the language calls for them),
repair the endings the recognizer mangled, lightly polish the spoken
wording - keeping every detail, example and the speaker's tone. Never
summarize, never drop anything, never add anything of your own, and do not
repeat the surrounding paragraphs: reply with this fragment alone, one
paragraph per language. When the original language is Russian, use the
letter «ё» wherever standard orthography calls for it.

Reply with JSON only:
{"paragraph_orig": string, "paragraph_en": string}
When the original language is English, fill "paragraph_en" only."""


DEDUP_SYSTEM_PROMPT = """\
You remove text that a previous edit of a lecture section placed twice.

You get the intro paragraphs of the section, the paragraphs of its
subsections and the fragments that both of them carry. The subsections keep
their copy; the intro has to lose it.

Return the intro paragraphs with the repeated text removed and everything
else kept word for word - do not rewrite, do not summarize, do not merge
what is left with anything, keep the two languages parallel. A paragraph
that loses all of its text disappears from the list.

Reply with JSON only:
{"intro_orig": [strings], "intro_en": [strings]}"""


ANNOTATION_CACHE_KEY = "__annotation__"

ANNOTATION_SYSTEM_PROMPT = """\
You write a concise annotation (abstract) of an edited video transcript.

Requirements:
- STRICT length: 200-250 words per language. Never fewer than 200 words.
  Count the words before replying; expand with concrete details from the
  transcript if the draft is short;
- cover the topic, the key points and the conclusions of the video;
- neutral informative tone, plain prose in 1-3 paragraphs, no bullet lists;
- write each annotation natively in its own language (translate the
  content, do not transliterate);
- spell the names of "glossary" exactly as listed there;
- for Russian text, use the letter «ё» wherever standard Russian
  orthography calls for it; do not replace «ё» with «e».

Reply with JSON only:
{"annotations": {"XX": "text", ...}}
with exactly the two-letter upper-case language codes requested in
"languages"."""


def annotation_languages(orig_code: str) -> list[str]:
    """Annotation languages by the rules:
    Russian original -> Russian annotation only; any other original ->
    English annotation plus a Russian translation."""
    return ["RU"] if orig_code == "RU" else ["EN", "RU"]


@dataclass
class Section:
    heading: str
    number: str | None  # docx numbering (from the agenda); None = unnumbered
    slide: dict | None
    start: float
    end: float = 0.0
    fixed_headings: dict[str, str] | None = None  # e.g. Conclusion per lang
    heading_en: str = ""
    intro: dict[str, list[str]] = field(default_factory=dict)
    subsections: list[dict] = field(default_factory=list)
    moved: bool = False  # a boundary of this section was refined this run
    anchors: list[str] = field(default_factory=list)  # see bullet_anchors


# --------------------------------------------------------------------------
# Inputs


def parse_srt(path: Path) -> list[Segment]:
    def to_seconds(stamp: str) -> float:
        hours, minutes, rest = stamp.strip().split(":")
        seconds, millis = rest.split(",")
        return (
            int(hours) * 3600 + int(minutes) * 60 + int(seconds)
            + int(millis) / 1000.0
        )

    segments: list[Segment] = []
    for block in re.split(r"\n\s*\n", path.read_text(encoding="utf-8-sig")):
        lines = [line for line in block.splitlines() if line.strip()]
        if len(lines) < 2 or "-->" not in lines[1]:
            continue
        start, end = lines[1].split("-->")
        segments.append(
            Segment(
                start=to_seconds(start),
                end=to_seconds(end),
                text=" ".join(line.strip() for line in lines[2:]),
            )
        )
    return segments


def section_text(
    segments: list[Segment],
    start: float,
    end: float,
    pauses: SilenceIndex | None = None,
) -> str:
    part = [seg for seg in segments if start <= seg.start < end]
    return "\n\n".join(group_paragraphs(part, pauses=pauses))


def detect_original_code(playlist_dir: Path) -> str:
    """Original-language folder code (RU, DE, ...); EN when only EN exists."""
    codes = sorted(
        entry.name
        for entry in playlist_dir.iterdir()
        if entry.is_dir()
        and len(entry.name) == 2
        and entry.name.isalpha()
        and entry.name.isupper()
        and entry.name != "EN"
        and any(entry.glob("*.srt"))
    )
    return codes[0] if codes else "EN"


# --------------------------------------------------------------------------
# Section list from slides.json


def build_sections(slides_data: dict, video_end: float) -> list[Section]:
    toc_by_slide = {
        entry.get("slide"): entry
        for entry in (slides_data.get("document") or {}).get("toc") or []
        if entry.get("slide")
    }
    sections: list[Section] = []
    for slide in slides_data.get("slides") or []:
        role = slide.get("role")
        scene = slide.get("scene") or {}
        start = float(scene.get("start_seconds") or 0.0)
        entry = toc_by_slide.get(slide.get("file")) or {}
        if role == "lecture_info":
            # The info slide only feeds document metadata; its lines are not
            # repeated inside the Foreword section.
            sections.append(Section("Foreword", None, None, start))
        elif role == "foreword":
            heading = entry.get("title") or slide.get("title") or "Foreword"
            sections.append(Section(heading, None, slide, start))
        elif role == "agenda":
            sections.append(Section("Agenda", None, slide, start))
        elif role == "section":
            heading = entry.get("title") or slide.get("title") or ""
            sections.append(Section(heading, entry.get("number"), slide, start))
        elif role == "closing" and (
            not sections or sections[-1].fixed_headings is None
        ):
            sections.append(
                Section(
                    "Conclusion",
                    None,
                    slide,
                    start,
                    fixed_headings=CONCLUSION_HEADINGS,
                )
            )
    if not sections:
        sections.append(Section("Transcript", None, None, 0.0))
    close_section_ranges(sections, video_end)
    return sections


def close_section_ranges(sections: list[Section], video_end: float) -> None:
    sections[0].start = 0.0
    for current, following in zip(sections, sections[1:]):
        current.end = following.start
    sections[-1].end = max(video_end, sections[-1].start + 1.0)


# --------------------------------------------------------------------------
# Section boundaries: from the slide change to the end of the sentence


# A slide is switched by hand, in the middle of speech, so the moment it
# changes is not where a section can be cut: on this course it lands inside a
# sentence at nearly every second boundary, and the half-sentence left hanging
# is what makes the editor of the next section retell the end of the previous
# one. The cut is moved to the nearest sentence end instead. The window is
# asymmetric because the habit is one-sided: the speaker announces the next
# topic first and reaches for the slide afterwards.
BOUNDARY_LOOKBACK_SECONDS = 20.0
BOUNDARY_LOOKAHEAD_SECONDS = 12.0
# ... so a candidate after the slide has to be that much closer to win.
BOUNDARY_FORWARD_PENALTY = 1.5
# Neither of the two neighbours may be shortened below this.
BOUNDARY_MIN_SECTION_SECONDS = 15.0
# The sentence before the slide change belongs to the section that follows
# when the speaker stopped for at least this long before saying it: she
# announces the next topic and only then reaches for the slide. A shorter
# gap is the rhythm of speech, not a new beginning.
ANNOUNCE_PAUSE_SECONDS = 1.0
# An announcement is short ("Тогда предлагаю перейти к практике."); a long
# sentence before the slide change is the previous topic still running.
ANNOUNCE_MAX_WORDS = 25
# Words too common on this course to say which slide is coming.
ANNOUNCE_STOPWORDS = {
    "test", "tests", "testing", "automation", "automated", "what", "which",
    "when", "your", "you", "our", "this", "that", "with", "from", "into",
    "will", "have", "they", "them", "their", "there", "than", "then", "them",
    "about", "also", "such", "some", "more", "most", "other", "using", "used",
}
SLIDE_MENTION_RE = re.compile(
    r"\b(слайд\w*|this slide|next slide|last slide)", re.IGNORECASE
)
# The speaker saying, in so many words, that she is moving on.
ANNOUNCE_TRANSITIONS = (
    "предлагаю перейти", "перейдём", "перейдем", "переходим", "перейти к",
    "давайте посмотрим", "следующий пункт", "следующая тема", "идём дальше",
    "поехали дальше", "двигаемся дальше", "поговорим о", "поговорим про",
    "и последнее", "теперь давайте", "let us look", "let's look",
    "let's move on", "next point", "moving on",
)
# A sentence end followed by the start of the next one. Only half of the .srt
# segments end at a sentence, so most of the usable cuts are inside the text
# of a segment; the capital letter keeps abbreviations ("и т.д. и т.п.") from
# looking like one.
SENTENCE_END_RE = re.compile(
    r"[.!?\u2026]+[\"\u00bb)\]]*\s+(?=[\"\u00ab(\[\u2013\u2014A-Z\u0410-\u042f\u0401])"
)


def char_time(segment: Segment, offset: int) -> float:
    """When the character at that offset of the segment text was spoken.

    Interpolated over the segment: the .srt has no timing inside it.
    """
    if not segment.text:
        return segment.start
    share = min(max(offset / len(segment.text), 0.0), 1.0)
    return segment.start + (segment.end - segment.start) * share


def sentence_spots(
    segments: list[Segment], low: float, high: float
) -> list[float]:
    """Times between `low` and `high` at which a sentence ends."""
    spots: list[float] = []
    for index, segment in enumerate(segments):
        if segment.end < low or segment.start > high:
            continue
        for match in SENTENCE_END_RE.finditer(segment.text):
            spot = char_time(segment, match.end())
            if low <= spot <= high:
                spots.append(spot)
        following = segments[index + 1] if index + 1 < len(segments) else None
        if following is not None and ends_sentence(segment.text):
            if low <= following.start <= high:
                spots.append(following.start)
    return spots


def split_segment_at(segments: list[Segment], moment: float) -> bool:
    """Split the segment `moment` falls into at the sentence end nearest to it.

    The section boundary is then a segment boundary again, and slicing the
    transcript by time keeps working as it did.
    """
    for index, segment in enumerate(segments):
        if not segment.start < moment < segment.end:
            continue
        offsets = [match.end() for match in SENTENCE_END_RE.finditer(segment.text)]
        if not offsets:
            return False
        offset = min(offsets, key=lambda o: abs(char_time(segment, o) - moment))
        head = segment.text[:offset].strip()
        tail = segment.text[offset:].strip()
        if not head or not tail:
            return False
        cut = char_time(segment, offset)
        segments[index] = replace(segment, end=cut, text=head)
        segments.insert(index + 1, replace(segment, start=cut, text=tail))
        return True
    return False


def sentences_with_times(
    segments: list[Segment]
) -> list[tuple[float, str]]:
    """(start, sentence) over a whole transcript.

    A sentence the recognizer split across two segments is glued back
    together and keeps the time of its first word.
    """
    pieces: list[tuple[float, str]] = []
    for segment in segments:
        text = segment.text.strip()
        if not text:
            continue
        start = 0
        for match in SENTENCE_END_RE.finditer(text):
            pieces.append(
                (char_time(segment, start), text[start:match.end()].strip())
            )
            start = match.end()
        tail = text[start:].strip()
        if tail:
            pieces.append((char_time(segment, start), tail))
    whole: list[tuple[float, str]] = []
    for start, text in pieces:
        if whole and not ends_sentence(whole[-1][1]):
            whole[-1] = (whole[-1][0], whole[-1][1] + " " + text)
        else:
            whole.append((start, text))
    return whole


def slide_words(section: Section) -> set[str]:
    """The words of the heading and the bullets of a section, worth hearing."""
    slide = section.slide or {}
    text = " ".join(
        [section.heading, str(slide.get("title") or "")]
        + [str(line) for line in slide.get("body") or []]
    )
    return {
        word for word in re.findall(r"[a-zA-Z]{4,}", text.lower())
        if word not in ANNOUNCE_STOPWORDS
    }


def distinctive_words(sections: list[Section]) -> list[set[str]]:
    """Per section, the words of its slide that the other slides do not use.

    "Automation" stands in half the titles of a lecture on test automation:
    hearing it says nothing about which slide is coming next.
    """
    per_section = [slide_words(section) for section in sections]
    counts: dict[str, int] = {}
    for words in per_section:
        for word in words:
            counts[word] = counts.get(word, 0) + 1
    limit = max(1, len(sections) // 3)
    return [
        {word for word in words if counts[word] <= limit}
        for words in per_section
    ]


def announces_section(text: str, english: str, wanted: set[str]) -> bool:
    """Whether this sentence is already about the section that follows."""
    if SLIDE_MENTION_RE.search(text):
        return True
    if any(phrase in text.lower() for phrase in ANNOUNCE_TRANSITIONS):
        return True
    return any(
        re.search(rf"\b{re.escape(word)}", english.lower()) for word in wanted
    )


def announced_start(
    segments: dict[str, list[Segment]],
    orig_code: str,
    boundary: float,
    low: float,
    wanted: set[str],
    pauses: SilenceIndex,
) -> float | None:
    """When the section really starts, if the speaker announced it early.

    The slide is switched by hand and the hand is late: the speaker stops,
    says what comes next ("Этот слайд такой себе...", "Тогда предлагаю
    перейти..."), and only then reaches for the slide. That last sentence
    opens the new section, not closes the old one - but only when all three
    are there: the stop, the announcement, and nothing else between it and
    the slide change.
    """
    source = segments.get(orig_code) or segments.get("EN") or []
    sentences = sentences_with_times(source)
    tail = [item for item in sentences if low <= item[0] < boundary - 0.05]
    if not tail:
        return None
    start, text = tail[-1]
    if len(text.split()) > ANNOUNCE_MAX_WORDS:
        return None
    if pauses.duration_at(start) < ANNOUNCE_PAUSE_SECONDS:
        return None
    english = " ".join(
        line for moment, line in sentences_with_times(segments.get("EN") or [])
        if start <= moment < boundary
    )
    return start if announces_section(text, english, wanted) else None


def refine_section_boundaries(
    sections: list[Section],
    segments: dict[str, list[Segment]],
    orig_code: str,
    pauses: SilenceIndex | None = None,
) -> list[tuple[str, float]]:
    """Move the boundaries off the middle of a sentence, and off the middle
    of an announcement (see announced_start).

    Returns (heading, shift in seconds) for every section that was moved.
    """
    source = segments.get(orig_code) or segments.get("EN") or []
    if len(sections) < 2 or not source:
        return []
    wanted = distinctive_words(sections)
    moved: list[tuple[str, float]] = []
    for index, (previous, section) in enumerate(
        zip(sections, sections[1:]), start=1
    ):
        slide_change = section.start
        low = max(
            previous.start + BOUNDARY_MIN_SECTION_SECONDS,
            slide_change - BOUNDARY_LOOKBACK_SECONDS,
        )
        high = min(
            section.end - BOUNDARY_MIN_SECTION_SECONDS,
            slide_change + BOUNDARY_LOOKAHEAD_SECONDS,
        )
        if low >= high:
            continue
        target: float | None = None
        before = [seg for seg in source if seg.start < slide_change]
        if before and not ends_sentence(before[-1].text):
            spots = sentence_spots(source, low, high)
            if spots:
                target = min(
                    spots,
                    key=lambda spot: (
                        slide_change - spot
                        if spot < slide_change
                        else (spot - slide_change) * BOUNDARY_FORWARD_PENALTY
                    ),
                )
        if pauses:
            announced = announced_start(
                segments,
                orig_code,
                slide_change if target is None else target,
                low,
                wanted[index],
                pauses,
            )
            if announced is not None:
                target = announced
        if target is None:
            continue
        for language_segments in segments.values():
            split_segment_at(language_segments, target)
        previous.end = target
        section.start = target
        # A cut that lands exactly on the slide change moves no text: the
        # sections are as they were, only the segment is now split there.
        if abs(target - slide_change) >= 0.05:
            previous.moved = section.moved = True
            moved.append((section.heading, target - slide_change))
    return moved


# --------------------------------------------------------------------------
# Slide bullets: where the speaker names each of them
#
# The reviewer's complaint about the subsections is not that they are wrong
# but that they are late: the speaker reads a bullet out ("Test Isolation:
# the result of one test shouldn't be depended on the another"), and the
# subsection starts two sentences further on, where the explanation gets
# going - after which the whole table of contents drifts by one item. Which
# sentence names which bullet is not a matter of judgement, so it is not
# asked of the model: it is matched here and handed over as a fact.
#
# The matching is done on the English transcript against the English slide,
# and the answer is quoted from the original: the speaker mixes languages
# ("И последний поинт здесь, что это less adaptability"), the slide does not.

# A bullet is recognized by the words no other bullet of the same slide
# uses - "test" and "automation" do not tell one bullet from another.
BULLET_MIN_HITS = 2
# One rare word is enough when it is long and specific ("parallelization").
BULLET_RARE_WORD_CHARS = 9
# How much of the original sentence to quote as the anchor.
ANCHOR_WORDS = 8
# Fewer bullets than this is not a list to walk through.
SPLIT_MIN_BULLETS = 3


def bullet_keywords(bullets: list[str]) -> list[set[str]]:
    """Per bullet, the words that set it apart from the other bullets."""
    words = [
        {
            word for word in re.findall(r"[a-zA-Z]{4,}", bullet.lower())
            if word not in ANNOUNCE_STOPWORDS
        }
        for bullet in bullets
    ]
    counts: dict[str, int] = {}
    for group in words:
        for word in group:
            counts[word] = counts.get(word, 0) + 1
    return [{word for word in group if counts[word] == 1} for group in words]


def bullet_anchors(
    bullets: list[str],
    orig_sentences: list[tuple[float, str]],
    en_sentences: list[tuple[float, str]],
) -> list[str]:
    """For each bullet, the sentence of the section that first names it.

    Empty string where nothing matched. The anchors are kept in bullet
    order: a bullet is looked for only after the one before it was found,
    because the speaker walks the slide from top to bottom.
    """
    if not bullets or not en_sentences:
        return []
    keywords = bullet_keywords(bullets)
    anchors: list[str] = []
    searched_from = 0
    for wanted in keywords:
        found = ""
        for index in range(searched_from, len(en_sentences)):
            moment, english = en_sentences[index]
            hits = {
                word for word in wanted
                if re.search(rf"\b{re.escape(word)}", english.lower())
            }
            rare = any(len(word) >= BULLET_RARE_WORD_CHARS for word in hits)
            if len(hits) < BULLET_MIN_HITS and not rare:
                continue
            spoken = min(
                orig_sentences or [(moment, english)],
                key=lambda item: abs(item[0] - moment),
            )[1]
            found = " ".join(spoken.split()[:ANCHOR_WORDS])
            searched_from = index + 1
            break
        anchors.append(found)
    return anchors


def sentences_of_section(
    segments: list[Segment], start: float, end: float
) -> list[tuple[float, str]]:
    return [
        item for item in sentences_with_times(segments)
        if start <= item[0] < end
    ]


def bullet_name_words(bullet: str) -> list[str]:
    """The words of the name a bullet goes by - what is before the colon.

    "End-to-End Testing: Simulating real user scenarios" is called
    "End-to-End", and a heading that says so is a heading for it, however
    little its wording has in common with the rest of the line.
    """
    if ":" not in bullet:
        return []
    name = bullet.split(":", 1)[0]
    return [
        word for word in re.findall(r"[a-zA-Z]{3,}", name.lower())
        if word not in ANNOUNCE_STOPWORDS and word not in ("the", "and")
    ]


def every_bullet_named(section: Section) -> bool:
    """Whether the speaker named every bullet of the slide, in slide order.

    A rare thing - three sections of the eighty-nine on this course - and
    the very case rule 4 calls a list of parallel items covered one after
    another. When it happens, a flat section is worth questioning.
    """
    bullets = (section.slide or {}).get("body") or []
    return (
        len(bullets) >= SPLIT_MIN_BULLETS
        and len(section.anchors) == len(bullets)
        and all(section.anchors)
    )


def bullets_without_subsection(section: Section) -> list[str]:
    """Bullets left without a subsection of their own.

    Rule 4 allows a section to stay flat, but not to be split half way:
    a bullet that gets no subsection is an item the table of contents
    loses, and everything after it shifts by one - what the reviewer saw
    as "оглавление поплыло".

    Only reported when most of the bullets did get one. A slide the
    speaker never walks through as a list (a comparison table of forty
    rows) is a different case entirely, and not one to complain about.
    """
    bullets = [str(line) for line in (section.slide or {}).get("body") or []]
    if not section.subsections or len(bullets) < 2:
        return []
    headings = [
        str(sub.get("heading_en") or "").lower()
        for sub in section.subsections
    ]
    missing: list[str] = []
    for bullet, wanted in zip(bullets, bullet_keywords(bullets)):
        named = bullet_name_words(bullet)
        if not wanted and not named:
            continue
        if any(
            any(re.search(rf"\b{re.escape(word)}", heading) for word in wanted)
            or (named and all(word in heading for word in named))
            for heading in headings
        ):
            continue
        missing.append(bullet)
    covered = len(bullets) - len(missing)
    return missing if covered * 2 > len(bullets) else []


# --------------------------------------------------------------------------
# Section list and document header without slides (INFO/<stem>.json written
# by transcribe_videos.py in remote mode: title, date, duration, timecodes)


def load_video_info(playlist_dir: Path, stem: str) -> dict:
    path = playlist_dir / INFO_DIRNAME / f"{stem}.json"
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return {}


def info_document_header(info: dict, stem: str) -> tuple[str, list[str]]:
    """Document title (summary format: <channel> | <playlist> : <title>)
    and meta lines from the saved video metadata."""
    title = str(info.get("title") or stem)
    channel = str(info.get("channel") or "")
    playlist = str(info.get("playlist") or "")
    if channel and playlist:
        doc_title = f"{channel} | {playlist} : {title}"
    elif channel:
        doc_title = f"{channel} : {title}"
    else:
        doc_title = title
    meta_lines = [
        f"{label}: {value}"
        for label, value in (
            ("Date", info.get("date")),
            ("Duration", f"[{info['duration_text']}]"
             if info.get("duration_text") else None),
            ("URL", info.get("url")),
        )
        if value
    ]
    return doc_title, meta_lines


UNTITLED_CHAPTER = re.compile(r"^<untitled chapter \d+>$", re.I)


def chapters_are_youtube_s_own(titles: list[str]) -> bool:
    """Whether the timecodes were titled by YouTube rather than the author.

    Automatic chapters arrive either as "<Untitled Chapter 1>" or as a
    lowercase phrase lifted out of the speech ("create a new dotnet core
    console"); an author writes a heading and capitalizes it. The set is
    judged as a whole, since YouTube titles either all of them or none.
    """
    named = [title.strip() for title in titles if title.strip()]
    if not named:
        return False
    return all(
        UNTITLED_CHAPTER.match(title) or title == title.lower()
        for title in named
    )


def build_sections_from_info(
    info: dict, video_end: float, orig_code: str
) -> list[Section]:
    """One numbered section per video chapter (timecode); a single
    'Transcript' section when the video has no timecodes.

    A chapter title written by the author is his own heading and is used
    verbatim in every language. One generated by YouTube is not a heading
    at all - it is a lowercase line of the speech, and "<Untitled Chapter
    1>" is not even that - so those give way to the heading the editor
    writes for the section anyway.
    """
    sections: list[Section] = []
    chapters = info.get("chapters") or []
    theirs = chapters_are_youtube_s_own(
        [str(chapter.get("title") or "") for chapter in chapters]
    )
    for index, chapter in enumerate(chapters, start=1):
        title = str(chapter.get("title") or f"Part {index}").strip()
        sections.append(
            Section(
                title,
                str(index),
                None,
                float(chapter.get("start") or 0.0),
                fixed_headings=(
                    None if theirs else {orig_code: title, "EN": title}
                ),
            )
        )
    if not sections:
        sections.append(
            Section(
                "Transcript", None, None, 0.0,
                fixed_headings=TRANSCRIPT_HEADINGS,
            )
        )
    close_section_ranges(sections, video_end)
    return sections


# --------------------------------------------------------------------------
# Editing via the OpenAI API


def request_edit(
    api_key: str,
    section: Section,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    orig_text: str,
    en_text: str,
    terms: list[str],
    usage: dict[str, int],
    part: tuple[int, int] | None,
    issues: Coverage,
    missing_bullets: list[str] | None = None,
    all_named: bool = False,
) -> dict:
    slide = section.slide or {}
    orig_name = LANGUAGE_NAMES.get(orig_code, orig_code)
    payload = {
        "course": course,
        "document_title": doc_title,
        "original_language": orig_name,
        "glossary": terms,
        "section_heading": section.heading,
        "slide_title": slide.get("title"),
        "slide_bullets": slide.get("body") or [],
        "heading_flags": slide.get("flags") or [],
        "section_minutes": round((section.end - section.start) / 60.0, 1),
        "transcript_original": orig_text if orig_code != "EN" else "",
        "transcript_english": en_text,
    }
    if part is not None:
        payload["section_part"] = f"{part[0]}/{part[1]}"
    if any(section.anchors):
        payload["bullet_anchors"] = [
            {"bullet": str(bullet), "named_at": anchor or None}
            for bullet, anchor in zip(
                slide.get("body") or [], section.anchors
            )
        ]
    if missing_bullets:
        payload["bullets_without_subsection"] = missing_bullets
    if all_named:
        payload["every_bullet_was_named"] = True
    if issues.dropped:
        payload["dropped_fragments"] = [gap.text for gap in issues.dropped]
    if issues.repeated:
        payload["repeated_fragments"] = issues.repeated
    return chat_json(
        api_key,
        [
            {"role": "system", "content": EDIT_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False, indent=2),
            },
        ],
        usage,
    )


def apply_edit_result(section: Section, result: dict, orig_code: str) -> None:
    section.heading_en = str(result.get("heading") or section.heading)
    section.intro = {
        orig_code: [str(p) for p in result.get("intro_orig") or []],
        "EN": [str(p) for p in result.get("intro_en") or []],
    }
    section.subsections = []
    for sub in result.get("subsections") or []:
        section.subsections.append(
            {
                "heading_en": str(sub.get("heading_en") or ""),
                "heading_orig": str(sub.get("heading_orig") or ""),
                orig_code: [str(p) for p in sub.get("paragraphs_orig") or []],
                "EN": [str(p) for p in sub.get("paragraphs_en") or []],
            }
        )
    slide = section.slide or {}
    if section.subsections and (
        slide.get("role") == "agenda"
        # Agenda items are the document's ToC: per-item subsections would
        # duplicate the main sections.
        or explanatory_bullets(slide.get("body") or [])
        # Explanatory slides (Definition:/Purpose:/Example: lines) do not
        # structure the speech that follows (documented limit case 1);
        # the model still splits them sometimes, so this is enforced here.
        or (
            len(section.subsections) == 1
            and len(slide.get("body") or []) > 1
        )
        # A single subsection on a slide of several bullets is not a
        # decomposition: it is the tail of the section given a heading of
        # its own, and it leaves the rest of the bullets nowhere.
    ):
        flatten_section(section)


def flatten_section(section: Section) -> None:
    """Fold the subsections back into the intro paragraphs, in order."""
    for sub in section.subsections:
        for code in section.intro:
            section.intro[code].extend(sub.get(code) or [])
    section.subsections = []


def edit_section(
    api_key: str,
    section: Section,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    orig_text: str,
    en_text: str,
    terms: list[str],
    usage: dict[str, int],
    part: tuple[int, int] | None = None,
) -> Coverage:
    """Edit the section in place and check the result against its transcript
    fragment. A section that lost or duplicated content is edited once more,
    with the offending fragments quoted back to the model; the better of the
    attempts is kept (a retry can come back worse than what it repairs), and
    whatever is still wrong in it is returned for the caller to report."""
    code = coverage_code(orig_code)
    source = orig_text if orig_code != "EN" else en_text
    best: dict | None = None
    issues = Coverage([], [])
    for attempt in range(1 + COVERAGE_RETRIES):
        result = request_edit(
            api_key,
            section,
            course=course,
            doc_title=doc_title,
            orig_code=orig_code,
            orig_text=orig_text,
            en_text=en_text,
            terms=terms,
            usage=usage,
            part=part,
            issues=issues,
        )
        apply_edit_result(section, result, orig_code)
        attempt_issues = section_coverage(section, source, code)
        if best is None or attempt_issues.weight() < issues.weight():
            best, issues = result, attempt_issues
        if issues.clean:
            break
        if attempt < COVERAGE_RETRIES:
            print(f"    {attempt_issues.summary()}; editing the section "
                  f"again...", flush=True)
    apply_edit_result(section, best, orig_code)
    if not issues.clean:
        # The whole-section re-edit did not help: fix the problems one by
        # one instead, and keep the outcome only if it is really better.
        print(f"    {issues.summary()}; repairing fragment by fragment...",
              flush=True)
        snapshot = section_snapshot(section)
        repair_section(
            api_key, section, issues, orig_code=orig_code, en_text=en_text,
            terms=terms, usage=usage,
        )
        repaired = section_coverage(section, source, code)
        if repaired.weight() < issues.weight():
            issues = repaired
        else:
            restore_snapshot(section, snapshot)
    if not issues.clean:
        print(f"    WARNING: {issues.summary()} left:", flush=True)
        for gap in issues.dropped:
            print(f"      lost: {shorten_fragment(gap.text)}", flush=True)
        for fragment in issues.repeated:
            print(f"      twice: {shorten_fragment(fragment)}", flush=True)
    issues = close_the_table_of_contents(
        api_key, section, issues,
        course=course, doc_title=doc_title, orig_code=orig_code,
        orig_text=orig_text, en_text=en_text, terms=terms, usage=usage,
        part=part, source=source, code=code,
    )
    return issues


def close_the_table_of_contents(
    api_key: str,
    section: Section,
    issues: Coverage,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    orig_text: str,
    en_text: str,
    terms: list[str],
    usage: dict[str, int],
    part: tuple[int, int] | None,
    source: str,
    code: str,
) -> Coverage:
    """Ask again when the subsections do not line up with the bullets.

    Two cases, both of them the reviewer's "оглавление поплыло": a split
    that covers most of the bullets but not all, and a section left flat
    although every one of its bullets was found named in the speech. The
    answer is kept only if it settles the case without costing coverage -
    a subsection invented for a bullet nobody discussed is worse than the
    missing line in the table of contents.
    """
    missing = bullets_without_subsection(section)
    named = every_bullet_named(section) and not section.subsections
    if not missing and not named:
        return issues
    if missing:
        print(f"    {len(missing)} bullet(s) without a subsection; asking "
              f"for the whole list...", flush=True)
    else:
        print("    every bullet was named in the speech, the section came "
              "back flat; asking again...", flush=True)
    snapshot = section_snapshot(section)
    result = request_edit(
        api_key, section,
        course=course, doc_title=doc_title, orig_code=orig_code,
        orig_text=orig_text, en_text=en_text, terms=terms, usage=usage,
        part=part, issues=Coverage([], []), missing_bullets=missing,
        all_named=named,
    )
    apply_edit_result(section, result, orig_code)
    again = section_coverage(section, source, code)
    settled = (
        not bullets_without_subsection(section)
        if missing
        else bool(section.subsections)
    )
    if settled and again.weight() <= issues.weight():
        return again
    restore_snapshot(section, snapshot)
    left = (
        ", ".join(shorten_fragment(bullet) for bullet in missing)
        if missing
        else "the section stays flat"
    )
    print(f"      kept as it was: {left}", flush=True)
    return issues


# --------------------------------------------------------------------------
# Coverage check: does the edited section still carry the whole fragment?
#
# The editor is asked to keep every detail, but a model sometimes silently
# drops a sentence (most often the transitional one at a section boundary)
# or emits the same fragment twice - once in the intro and again inside the
# subsection it belongs to. Both are caught by diffing the edited text
# against the transcript slice it was made from: only the original language
# is checked, because there the edit stays close to the source wording (the
# English side is a machine translation the editor rewrites much more
# freely, which would make the diff meaningless).

WORD_RE = re.compile(r"\w+", re.UNICODE)
# Shorter dropped runs are normal polishing (filler words, false starts).
COVERAGE_MIN_RUN_WORDS = 5
# Shorter repeated runs are common speech, not a duplicated fragment.
COVERAGE_MIN_REPEAT_WORDS = 12
COVERAGE_RETRIES = 1
FRAGMENT_PREVIEW_CHARS = 120
# A paragraph longer than this is reported after the run: rule 1a of the
# prompt asks for 50-120 words, and this much past it is no longer a
# borderline case but the wall of text the rule is there to prevent.
PARAGRAPH_REPORT_WORDS = 200


@dataclass
class Gap:
    """A fragment of the transcript missing from the edited text, with the
    word offset in that text where it should have been."""
    text: str
    offset: int


@dataclass
class Coverage:
    dropped: list[Gap]
    repeated: list[str]

    @property
    def clean(self) -> bool:
        return not self.dropped and not self.repeated

    def weight(self) -> int:
        """How bad the result is, in words: attempts are compared by this."""
        return sum(len(gap.text.split()) for gap in self.dropped) + sum(
            len(fragment.split()) for fragment in self.repeated
        )

    def summary(self) -> str:
        parts = []
        if self.dropped:
            parts.append(
                f"{len(self.dropped)} fragment(s) dropped, "
                f"{sum(len(gap.text.split()) for gap in self.dropped)} words"
            )
        if self.repeated:
            parts.append(f"{len(self.repeated)} fragment(s) duplicated")
        return " and ".join(parts) or "complete"

    def as_cache(self) -> dict[str, list[str]]:
        return {
            "dropped": [gap.text for gap in self.dropped],
            "repeated": self.repeated,
        }


def coverage_code(orig_code: str) -> str:
    """Language the coverage check runs on: the original one."""
    return orig_code if orig_code != "EN" else "EN"


def section_paragraphs(section: Section, code: str) -> list[str]:
    paragraphs = list(section.intro.get(code) or [])
    for sub in section.subsections:
        paragraphs.extend(sub.get(code) or [])
    return [text for text in paragraphs if text.strip()]


def long_paragraphs(
    sections: list[Section], code: str, limit: int = PARAGRAPH_REPORT_WORDS
) -> list[tuple[str, int]]:
    """(heading, words) of every paragraph still too long after polishing.

    Rule 1a asks for 50-120 words and polish_section halves anything past
    this limit, so what is reported here is the one case it cannot fix: a
    paragraph without a single sentence end in it.
    """
    found: list[tuple[str, int]] = []
    for section in sections:
        for paragraph in section_paragraphs(section, code):
            words = len(paragraph.split())
            if words > limit:
                found.append((section.heading, words))
    return found


SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+")
QUESTION_OPENER_RE = re.compile(
    r"^[«\"(\s-]*(почему|зачем|каки[ем]|кака[яю]|како[йе]|каким|сколько|"
    r"что такое|что это|в чём|в чем)\b(?!-(?:то|нибудь|либо))",
    re.IGNORECASE,
)
# Past this a sentence opening with a question word is usually a statement
# ("Какие скиллы должны быть, какие навыки, на каком уровне и так далее.").
QUESTION_MAX_WORDS = 12


def punctuation_density(sections: list[Section], code: str) -> str:
    """The marks per 1000 words of the edited text, as one line of log.

    Rule 1b is the one rule whose result swings from run to run - the same
    lecture came back with 14 dashes once and with 3 the next time - so the
    number is printed rather than left to a probe.
    """
    text = "\n".join(
        paragraph
        for section in sections
        for paragraph in section_paragraphs(section, code)
    )
    words = len(text.split())
    if not words:
        return ""
    marks = {
        "dashes": len(re.findall(r"[\u2013\u2014]", text)),
        "colons": text.count(":"),
        "question marks": text.count("?"),
    }
    return ", ".join(
        f"{count * 1000 / words:.1f} {name}" for name, count in marks.items()
    ) + " per 1000 words"


def questions_with_a_full_stop(
    sections: list[Section], code: str
) -> list[str]:
    """The sentences that ask something and end with a full stop anyway.

    Only the plainest cases: a short sentence opening with a question word.
    The lecturer asks and answers herself all the time, and rule 1b keeps
    losing those question marks; this counts what is left, and, like the
    long paragraphs above, it costs nothing and repairs nothing.
    """
    if code != "RU":
        return []
    found: list[str] = []
    for section in sections:
        for paragraph in section_paragraphs(section, code):
            for sentence in SENTENCE_SPLIT_RE.split(paragraph):
                sentence = sentence.strip()
                if (
                    sentence.endswith(".")
                    and len(sentence.split()) <= QUESTION_MAX_WORDS
                    and QUESTION_OPENER_RE.match(sentence)
                ):
                    found.append(sentence)
    return found


# --------------------------------------------------------------------------
# The letter «ё»
#
# Rule 1c asks for it and the editor still writes «еще» more often than
# «ещё», so the words that can only be spelled with «ё» are put right here,
# for free and for certain. The ambiguous pairs stay with the editor: «все»
# and «всё», «чем» and «чём» are different words, and only the sentence
# says which one was meant.

YO_WORDS = """
    ещё её неё моё твоё своё
    идёт идём идёшь придёт придём пойдёт пойдём найдёт найдём
    перейдём зайдём подойдёт произойдёт обойдётся обойдёмся
    даёт даём отдаёт создаёт создаём передаёт остаётся остаёмся
    берёт берём возьмёт возьмём начнёт начнём поймёт поймём
    живёт живём ведёт ведём приведёт приведём несёт несём растёт
    зовёт зовём назовём разберём разберёмся вернёмся займёмся
    счёт счёта счёте отчёт отчёта отчёты учёт учёта расчёт расчёта зачёт
    чёткий чёткая чёткое чёткие чётких чёткую чётко чётче
    серьёзно серьёзный серьёзная серьёзные серьёзных
    надёжно надёжный надёжная надёжные надёжность
    объём объёма объёме объёмы приём приёма приёме приёмы
    лёгкий лёгкая лёгкое лёгкие лёгких
    шёл прошёл пришёл нашёл ушёл пошёл зашёл вошёл подошёл перешёл
    произошёл обошёлся тяжёлый тяжёлая тяжёлые жёсткий жёсткая жёсткие
""".split()
YO_BY_PLAIN = {word.replace("ё", "е"): word for word in YO_WORDS}
YO_RE = re.compile(
    r"\b(" + "|".join(sorted(YO_BY_PLAIN, key=len, reverse=True)) + r")\b",
    re.IGNORECASE,
)


def spell_yo(text: str) -> str:
    """«еще» -> «ещё», and the like: only where «е» is never right."""
    def fix(match: re.Match) -> str:
        word = match.group(0)
        spelled = YO_BY_PLAIN[word.lower()]
        if word.isupper():
            return spelled.upper()
        if word[0].isupper():
            return spelled[0].upper() + spelled[1:]
        return spelled

    return YO_RE.sub(fix, text)


# --------------------------------------------------------------------------
# Together or apart
#
# «имею ввиду», «в последствии»: the reviewer marked these five times, and
# they need no reading of the sentence either. «Ввиду» after a form of
# «иметь» is always two words - the preposition «ввиду» means "because of"
# and cannot be had; «в последствии» is always one word, the case that would
# make two of them is «в последствиях». Everything else about them is the
# same story as «ё»: rule 1c asks, the editor writes what it heard.

PHRASE_FIXES = [
    (
        re.compile(
            r"\b(име(?:ю|ешь|ет|ем|ете|ют|л|ла|ли|лось|ется))\s+ввиду\b",
            re.IGNORECASE,
        ),
        r"\1 в виду",
    ),
    (re.compile(r"\bимейте\s+ввиду\b", re.IGNORECASE), "имейте в виду"),
    (re.compile(r"\bв\s+последствии\b", re.IGNORECASE), "впоследствии"),
]


def respell(text: str) -> str:
    """«имею ввиду» -> «имею в виду», «в последствии» -> «впоследствии»."""
    def fix(match: re.Match, replacement: str) -> str:
        fixed = match.expand(replacement)
        if match.group(0)[:1].isupper():
            fixed = fixed[:1].upper() + fixed[1:]
        return fixed

    for pattern, replacement in PHRASE_FIXES:
        text = pattern.sub(lambda m, r=replacement: fix(m, r), text)
    return text


def split_paragraph(text: str, limit: int) -> list[str]:
    """Halve an over-long paragraph at a sentence end, and again if needed."""
    words = len(text.split())
    if words <= limit:
        return [text]
    sentences = SENTENCE_SPLIT_RE.split(text)
    if len(sentences) < 2:
        return [text]
    cut = running = 0
    nearest = words
    for index, sentence in enumerate(sentences[:-1], start=1):
        running += len(sentence.split())
        if abs(running - words / 2) < nearest:
            nearest = abs(running - words / 2)
            cut = index
    return split_paragraph(" ".join(sentences[:cut]), limit) + split_paragraph(
        " ".join(sentences[cut:]), limit
    )


def transcript_anchor(source_words: list[str], paragraph: str) -> int:
    """Where in the transcript the paragraph's longest verbatim run starts."""
    words = coverage_words(paragraph)
    if not words:
        return -1
    match = difflib.SequenceMatcher(
        None, source_words, words, autojunk=False
    ).find_longest_match(0, len(source_words), 0, len(words))
    return match.a if match.size >= COVERAGE_MIN_RUN_WORDS else -1


def out_of_order(anchors: list[int]) -> list[int]:
    """Which of the anchors have to move for the rest to run forwards.

    What stays is the longest run of anchors that never goes backwards;
    everything outside it is what the editor carried out of place, and
    moving exactly that is the smallest repair there is.
    """
    if not anchors:
        return []
    longest = [1] * len(anchors)
    after = [-1] * len(anchors)
    for index, anchor in enumerate(anchors):
        for earlier in range(index):
            if anchors[earlier] <= anchor and longest[earlier] >= longest[index]:
                longest[index] = longest[earlier] + 1
                after[index] = earlier
    stays: set[int] = set()
    index = max(range(len(anchors)), key=lambda at: longest[at])
    while index >= 0:
        stays.add(index)
        index = after[index]
    return [index for index in range(len(anchors)) if index not in stays]


# Anchoring a sentence can be a word or two out, so a handful of words
# apparently out of place is noise rather than a passage the editor moved.
ORDER_MIN_MOVED_WORDS = 30
# Carrying a passage out of the paragraph it was hung on can leave a
# sentence or two behind. Rule 1a has them join their neighbour rather than
# stand alone, and here the neighbour is the text they follow in the speech.
ORDER_MIN_PARAGRAPH_WORDS = 30


def in_spoken_order(source: str, paragraphs: list[str]) -> tuple[list[str], int]:
    """The paragraphs with any moved passage carried back, and its size.

    The editor moves a passage without moving a whole paragraph: here it
    took the hundred words that open a transcript paragraph and hung them
    on the end of the paragraph after it. So the work is done a sentence at
    a time, and the sentences that stay put keep the paragraph they were
    given; only the moved run has to become a paragraph of its own, which
    is what it is in the transcript anyway.
    """
    source_words = coverage_words(source)
    if not source_words or len(paragraphs) < 2:
        return paragraphs, 0
    numbers: list[int] = []
    sentences: list[str] = []
    anchors: list[int] = []
    last = 0
    for number, paragraph in enumerate(paragraphs):
        for sentence in SENTENCE_SPLIT_RE.split(paragraph):
            if not sentence.strip():
                continue
            found = transcript_anchor(source_words, sentence)
            last = last if found < 0 else found
            numbers.append(number)
            sentences.append(sentence)
            anchors.append(last)
    displaced = out_of_order(anchors)
    moved_words = sum(len(sentences[index].split()) for index in displaced)
    if moved_words < ORDER_MIN_MOVED_WORDS:
        return paragraphs, 0
    order = sorted(range(len(sentences)), key=lambda at: (anchors[at], at))
    rebuilt: list[str] = []
    previous: int | None = None
    for index in order:
        short = (
            rebuilt
            and len(rebuilt[-1].split()) < ORDER_MIN_PARAGRAPH_WORDS
        )
        if numbers[index] == previous or short:
            rebuilt[-1] = f"{rebuilt[-1]} {sentences[index]}"
        else:
            rebuilt.append(sentences[index])
        previous = numbers[index]
    return rebuilt, moved_words


def paragraphs_in_spoken_order(section: Section, sources: dict[str, str]) -> int:
    """Undo the passages the editor moved; return how many words came back.

    Rule 1d asks for the order of the recording and is mostly obeyed, but a
    passage placed after the one that used to follow it costs twice over:
    the document no longer runs alongside the video, and the coverage diff -
    which can align in one order only - reports the passage missing from the
    place it left. Putting it back is arithmetic, so it is done here rather
    than paid for. Each language is straightened against its own transcript
    and each container on its own, so no subsection takes in the text of its
    neighbour.
    """
    moved = 0
    for container in [section.intro, *section.subsections]:
        for code, paragraphs in list(container.items()):
            source = sources.get(code) or ""
            if not source or not isinstance(paragraphs, list):
                continue
            rebuilt, words = in_spoken_order(source, paragraphs)
            if words:
                container[code] = rebuilt
                moved += words
    return moved


def polish_section(section: Section, code: str) -> None:
    """Put the finished section through what needs no model: paragraphs of
    a readable length, the letter «ё» where only «ё» can stand, and the few
    phrases that are misspelled whatever the sentence.

    Rule 1a is asked for, reported on - and still broken: the editor merges
    paragraphs into blocks of 450 words, and the coverage repair inserts a
    lost fragment of the transcript as one paragraph of whatever length it
    happens to be. Splitting them is arithmetic, not judgement, so it is
    done here rather than paid for.
    """
    for container in [section.intro, *section.subsections]:
        for key, value in list(container.items()):
            if not isinstance(value, list):
                continue
            texts: list[str] = []
            for paragraph in value:
                texts.extend(
                    split_paragraph(str(paragraph), PARAGRAPH_REPORT_WORDS)
                )
            container[key] = (
                [respell(spell_yo(text)) for text in texts]
                if key == "RU"
                else texts
            )
    if code == "RU":
        for sub in section.subsections:
            if sub.get("heading_orig"):
                sub["heading_orig"] = respell(spell_yo(sub["heading_orig"]))


def coverage_words(text: str) -> list[str]:
    """Comparable words: case and «ё» are edited on purpose, so both are
    normalized away."""
    return [word.lower().replace("ё", "е") for word in WORD_RE.findall(text)]


def shorten_fragment(fragment: str) -> str:
    if len(fragment) <= FRAGMENT_PREVIEW_CHARS:
        return fragment
    return fragment[:FRAGMENT_PREVIEW_CHARS].rstrip() + "..."


def section_language_text(section: Section, code: str) -> str:
    paragraphs = list(section.intro.get(code) or [])
    for sub in section.subsections:
        paragraphs.extend(sub.get(code) or [])
    return "\n\n".join(paragraphs)


def section_word_count(section: Section, code: str) -> int:
    return len(coverage_words(section_language_text(section, code)))


def cache_incomplete(entry: dict) -> bool:
    """Cached edit still recorded dropped or duplicated transcript text."""
    recorded = entry.get("coverage") or {}
    return bool(recorded.get("dropped") or recorded.get("repeated"))


# How much of a run has to be found elsewhere in the edited text before it
# counts as moved rather than lost. Below this the run really is missing and
# what was found are the few words any two sentences share.
COVERAGE_MOVED_SHARE = 0.8


def carried_elsewhere(run: list[str], edited: list[str]) -> bool:
    """The run is in the edited text, only not where the diff looked.

    A diff aligns in one order, so a passage the model moved - here a
    paragraph placed after the one that used to follow it - can match in
    only one of its two places and is reported missing from the other.
    Restoring it then puts a second copy into the document, which is how a
    reordered paragraph turns into a repeated one.
    """
    matcher = difflib.SequenceMatcher(None, run, edited, autojunk=False)
    carried = sum(
        block.size
        for block in matcher.get_matching_blocks()
        if block.size >= COVERAGE_MIN_RUN_WORDS
    )
    return carried >= COVERAGE_MOVED_SHARE * len(run)


def dropped_fragments(source: str, edited: str) -> list[Gap]:
    """Runs of source words that the edited text does not carry at all.

    The gap is cut out of the transcript with its punctuation: it is handed
    back to the model to be edited and put in place, and a fragment served
    as a bare list of words comes back as a bare list of words.
    """
    words = list(WORD_RE.finditer(source))

    def cut(start: int, stop: int) -> str:
        """The run of source words with the marks between and after them."""
        end = words[stop].start() if stop < len(words) else len(source)
        text = source[words[start].start():end]
        # The mark that opens the next sentence came along with the rest.
        return " ".join(text.split()).rstrip("«\"([-–— ")

    source_words = coverage_words(source)
    edited_words = coverage_words(edited)
    matcher = difflib.SequenceMatcher(
        None, source_words, edited_words, autojunk=False
    )
    return [
        Gap(cut(start, stop), offset)
        for tag, start, stop, offset, _ in matcher.get_opcodes()
        if tag in ("delete", "replace")
        and stop - start >= COVERAGE_MIN_RUN_WORDS
        and not carried_elsewhere(source_words[start:stop], edited_words)
    ]


def count_runs(words: list[str], run: list[str]) -> int:
    """How many times the word run occurs (without overlaps)."""
    total = index = 0
    while index + len(run) <= len(words):
        if words[index:index + len(run)] == run:
            total += 1
            index += len(run)
        else:
            index += 1
    return total


def repeated_fragments(section: Section, source: str, code: str) -> list[str]:
    """Long word runs the section carries twice: an intro paragraph and the
    subsection it was also put into, two subsections retelling each other,
    two copies of one paragraph. A run the speaker really said twice is not
    a duplicate, so a fragment counts only when the edited text has more
    copies of it than the transcript does."""
    texts = [paragraph for paragraph in section.intro.get(code) or []]
    texts.extend(
        "\n\n".join(sub.get(code) or []) for sub in section.subsections
    )
    candidates: list[list[str]] = []
    for first, text in enumerate(texts):
        words_first = coverage_words(text)
        for second in range(first + 1, len(texts)):
            matcher = difflib.SequenceMatcher(
                None, words_first, coverage_words(texts[second]),
                autojunk=False,
            )
            candidates.extend(
                words_first[block.a:block.a + block.size]
                for block in matcher.get_matching_blocks()
                if block.size >= COVERAGE_MIN_REPEAT_WORDS
            )
    edited = coverage_words("\n\n".join(texts))
    source_words = coverage_words(source)
    found: list[str] = []
    for run in candidates:
        fragment = " ".join(run)
        if fragment in found:
            continue
        if count_runs(edited, run) > count_runs(source_words, run):
            found.append(fragment)
    return found


def language_containers(section: Section, code: str) -> list[list[str]]:
    """Paragraph lists of one language, in document order."""
    containers = [section.intro.setdefault(code, [])]
    containers.extend(sub.setdefault(code, []) for sub in section.subsections)
    return containers


def paragraph_index_for_offset(
    section: Section, code: str, offset: int
) -> int:
    """Where a fragment starting at that word offset of the edited text
    belongs, counted in whole paragraphs of the section."""
    used = index = 0
    for container in language_containers(section, code):
        for paragraph in container:
            if used >= offset:
                return index
            used += len(WORD_RE.findall(paragraph))
            index += 1
    return index


def insert_paragraph(
    section: Section, code: str, index: int, paragraph: str
) -> None:
    if not paragraph.strip():
        return
    position = index
    containers = language_containers(section, code)
    for container in containers:
        if position <= len(container):
            container.insert(position, paragraph)
            return
        position -= len(container)
    containers[-1].append(paragraph)


def restore_fragment(
    api_key: str,
    section: Section,
    gap: Gap,
    *,
    orig_code: str,
    en_text: str,
    terms: list[str],
    usage: dict[str, int],
) -> None:
    """Edit one dropped fragment on its own and put it back in place.

    Re-editing the whole section rarely brings back a fragment the model
    decided to skip (typically the transitional sentences at the very end
    of a section, which already announce the next slide); asked for that
    fragment alone, it does the job.
    """
    code = coverage_code(orig_code)
    index = paragraph_index_for_offset(section, code, gap.offset)
    neighbours = [
        paragraph
        for container in language_containers(section, code)
        for paragraph in container
    ][max(0, index - 1):index + 1]
    result = chat_json(
        api_key,
        [
            {"role": "system", "content": FRAGMENT_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "original_language": LANGUAGE_NAMES.get(
                            orig_code, orig_code
                        ),
                        "glossary": terms,
                        "dropped_fragment": gap.text,
                        "section_english_transcript": en_text,
                        "edited_neighbours": neighbours,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            },
        ],
        usage,
    )
    insert_paragraph(
        section, code, index, str(result.get("paragraph_orig") or "")
        if orig_code != "EN" else str(result.get("paragraph_en") or "")
    )
    if orig_code != "EN" and "EN" in section.intro:
        insert_paragraph(
            section, "EN", index, str(result.get("paragraph_en") or "")
        )


def drop_repeats(
    api_key: str,
    section: Section,
    fragments: list[str],
    *,
    orig_code: str,
    usage: dict[str, int],
) -> None:
    """Remove from the intro the text its subsections carry as well."""
    result = chat_json(
        api_key,
        [
            {"role": "system", "content": DEDUP_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "original_language": LANGUAGE_NAMES.get(
                            orig_code, orig_code
                        ),
                        "repeated_fragments": fragments,
                        "intro_orig": section.intro.get(orig_code) or [],
                        "intro_en": section.intro.get("EN") or [],
                        "subsections": [
                            {
                                "heading": sub.get("heading_en"),
                                "paragraphs_orig": sub.get(orig_code) or [],
                            }
                            for sub in section.subsections
                        ],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            },
        ],
        usage,
    )
    section.intro[orig_code] = [
        str(p) for p in result.get("intro_orig") or []
    ]
    if "EN" in section.intro:
        section.intro["EN"] = [str(p) for p in result.get("intro_en") or []]


# The restore prompt asks for the gap back as one paragraph, so a gap the
# size of half a section is condensed exactly the way the edit that lost it
# condensed it. Above this many words the gap is restored piece by piece,
# each piece a paragraph of the size rule 1a asks for.
FRAGMENT_RESTORE_WORDS = 120


def split_gap(gap: Gap, limit: int = FRAGMENT_RESTORE_WORDS) -> list[Gap]:
    """The gap cut into pieces of about `limit` words at sentence ends."""
    words = gap.text.split()
    if len(words) <= limit:
        return [gap]
    pieces: list[str] = []
    current: list[str] = []
    for word in words:
        current.append(word)
        if len(current) >= limit and ends_sentence(word):
            pieces.append(" ".join(current))
            current = []
    if current:
        if pieces and len(current) < COVERAGE_MIN_RUN_WORDS:
            pieces[-1] = f"{pieces[-1]} {' '.join(current)}"
        else:
            pieces.append(" ".join(current))
    return [Gap(text, gap.offset) for text in pieces]


def repair_section(
    api_key: str,
    section: Section,
    issues: Coverage,
    *,
    orig_code: str,
    en_text: str,
    terms: list[str],
    usage: dict[str, int],
) -> None:
    """Fix what is still wrong after the re-edit, one problem per call."""
    code = coverage_code(orig_code)
    for gap in sorted(issues.dropped, key=lambda item: -item.offset):
        pieces = split_gap(gap)
        if len(pieces) > 1:
            print(
                f"      {len(gap.text.split())} words in one gap; restoring "
                f"them in {len(pieces)} pieces...",
                flush=True,
            )
        # Every piece is inserted at the same paragraph index, so they go in
        # back to front to come out in the order the speaker said them. A
        # piece the section turns out to carry after all is left alone: the
        # gap was cut out of a passage the model only moved.
        for piece in reversed(pieces):
            here = coverage_words(section_language_text(section, code))
            if carried_elsewhere(coverage_words(piece.text), here):
                continue
            restore_fragment(
                api_key, section, piece, orig_code=orig_code, en_text=en_text,
                terms=terms, usage=usage,
            )
    if issues.repeated:
        drop_repeats(
            api_key, section, issues.repeated, orig_code=orig_code,
            usage=usage,
        )


def section_snapshot(section: Section) -> dict:
    return json.loads(json.dumps(section_state(section)))


def section_state(section: Section) -> dict:
    return {
        "heading_en": section.heading_en,
        "intro": section.intro,
        "subsections": section.subsections,
    }


def restore_snapshot(section: Section, snapshot: dict) -> None:
    section.heading_en = snapshot["heading_en"]
    section.intro = snapshot["intro"]
    section.subsections = snapshot["subsections"]


def section_coverage(section: Section, source: str, code: str) -> Coverage:
    if not source.strip():
        return Coverage([], [])
    return Coverage(
        dropped_fragments(source, section_language_text(section, code)),
        repeated_fragments(section, source, code),
    )


# gpt-4o silently compresses long re-emissions: asked to edit a whole
# 40-minute transcript in one call it returns a couple of paragraphs. Only
# sections whose edited text would be longer than this are therefore split
# into parts and concatenated - in practice the single section of a
# timecode-less video. Slide sections stay whole even when they run 15
# minutes: splitting them costs the subsection structure (parts are always
# flat) and, as the coverage check showed, loses text at the seams.
EDIT_CHUNK_TOKENS = 7000
# A single flat section (no slides, no chapters) longer than this is edited
# in parts up front. Whole-section calls still compress the middle even when
# the text is far below EDIT_CHUNK_TOKENS (a 7-minute monologue here lost
# subtitles 83-130 that way).
PROACTIVE_FLAT_SPLIT_WORDS = 700


def chunk_count(orig_text: str, en_text: str, orig_code: str) -> int:
    """Into how many parts a section has to be split for editing."""
    tokens = estimate_tokens(orig_text, orig_code) + estimate_tokens(
        en_text, "EN"
    )
    return max(1, -(-tokens // EDIT_CHUNK_TOKENS))


def split_section_ranges(
    section: Section,
    segments: list[Segment],
    parts: int,
) -> list[tuple[float, float]]:
    """Split [start, end) into roughly equal windows, snapping each boundary
    to the nearest transcript segment boundary so that no sentence is cut in
    half (both languages are sliced by the same time boundaries, staying
    parallel)."""
    duration = section.end - section.start
    if parts <= 1:
        return [(section.start, section.end)]
    candidates = [
        seg.start
        for seg in segments
        if section.start < seg.start < section.end
    ]
    boundaries: list[float] = []
    for index in range(1, parts):
        target = section.start + duration * index / parts
        snapped = (
            min(candidates, key=lambda t: abs(t - target))
            if candidates
            else target
        )
        if boundaries and snapped <= boundaries[-1]:
            continue
        boundaries.append(snapped)
    edges = [section.start, *boundaries, section.end]
    return [
        (edges[i], edges[i + 1])
        for i in range(len(edges) - 1)
        if edges[i + 1] > edges[i]
    ]


def edit_section_in_chunks(
    api_key: str,
    section: Section,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    segments: dict[str, list[Segment]],
    parts: int,
    terms: list[str],
    usage: dict[str, int],
    pauses: SilenceIndex | None = None,
) -> Coverage:
    """Edit one long section chunk by chunk; the result is always flat."""
    ranges = split_section_ranges(section, segments.get(orig_code) or [], parts)
    intro: dict[str, list[str]] = {orig_code: [], "EN": []}
    issues = Coverage([], [])
    heading_en = ""
    for index, (start, end) in enumerate(ranges, start=1):
        print(
            f"    part {index}/{len(ranges)} "
            f"({(end - start) / 60.0:.1f} min)...",
            flush=True,
        )
        part = Section(
            heading=section.heading,
            number=section.number,
            slide=section.slide,
            start=start,
            end=end,
        )
        orig_text = (
            section_text(segments[orig_code], start, end, pauses)
            if orig_code != "EN"
            else ""
        )
        en_text = (
            section_text(segments["EN"], start, end, pauses)
            if "EN" in segments
            else ""
        )
        part_issues = edit_section(
            api_key,
            part,
            course=course,
            doc_title=doc_title,
            orig_code=orig_code,
            orig_text=orig_text,
            en_text=en_text,
            terms=terms,
            usage=usage,
            part=(index, len(ranges)),
        )
        issues.dropped.extend(part_issues.dropped)
        issues.repeated.extend(part_issues.repeated)
        flatten_section(part)
        heading_en = heading_en or part.heading_en
        for code, paragraphs in part.intro.items():
            intro.setdefault(code, []).extend(paragraphs)
    section.heading_en = heading_en or section.heading
    section.intro = intro
    section.subsections = []
    return issues


# What a hard section loses is not a fragment but its tail: the model
# edits the first half attentively and summarizes the rest away. Asked
# again for the whole gap - 878 words in one lecture here - it condenses
# it a second time, so above this share of the section the fragment
# repair is the wrong tool and the section is edited again in parts,
# each short enough to be carried to its end. Below it the losses are
# scattered phrases, which the repair does handle.
COVERAGE_SPLIT_SHARE = 0.2
# Parts of about this size: the failing sections are the ones where the
# model gave up around the middle.
COVERAGE_SPLIT_TOKENS = 1200


def parts_for_a_losing_section(
    issues: Coverage, orig_text: str, en_text: str, orig_code: str
) -> int:
    """Into how many parts to re-edit a section that came back short."""
    source = orig_text if orig_code != "EN" else en_text
    total = len(coverage_words(source))
    if not total or issues.weight() < COVERAGE_SPLIT_SHARE * total:
        return 1
    tokens = estimate_tokens(orig_text, orig_code) + estimate_tokens(
        en_text, "EN"
    )
    return max(2, -(-tokens // COVERAGE_SPLIT_TOKENS))


def edit_section_again_in_parts(
    api_key: str,
    section: Section,
    issues: Coverage,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    orig_text: str,
    en_text: str,
    segments: dict[str, list[Segment]],
    parts: int,
    terms: list[str],
    usage: dict[str, int],
    pauses: SilenceIndex | None,
) -> Coverage:
    """Re-edit a section that lost its tail, in parts; keep the better one.

    The parts come back flat, so the section can lose its subsections
    here - a fair price for the paragraphs it would otherwise not have
    at all, and only paid when the whole-section attempts have failed.
    """
    print(
        f"    {issues.summary()}; editing the section in "
        f"{parts} parts...",
        flush=True,
    )
    check_code = coverage_code(orig_code)
    snapshot = section_snapshot(section)
    before_words = section_word_count(section, check_code)
    again = edit_section_in_chunks(
        api_key,
        section,
        course=course,
        doc_title=doc_title,
        orig_code=orig_code,
        segments=segments,
        parts=parts,
        terms=terms,
        usage=usage,
        pauses=pauses,
    )
    after_words = section_word_count(section, check_code)
    if (
        again.weight() < issues.weight()
        or after_words > before_words
    ):
        return again
    print(
        "    the parts came back no better; keeping the whole-section edit.",
        flush=True,
    )
    restore_snapshot(section, snapshot)
    return issues


EXPLANATORY_LABELS = {
    "definition", "purpose", "example", "goal", "examples", "note",
}


def explanatory_bullets(bullets: list[str]) -> bool:
    labeled = [
        bullet.split(":", 1)[0].strip().lower()
        for bullet in bullets
        if ":" in bullet
    ]
    known = [label for label in labeled if label in EXPLANATORY_LABELS]
    return len(known) >= 2 and 2 * len(known) >= len(bullets)


def section_heading_for(section: Section, lang: str) -> str:
    if section.fixed_headings and lang in section.fixed_headings:
        return section.fixed_headings[lang]
    return section.heading_en or section.heading


def subsection_heading_for(sub: dict, lang: str) -> str:
    if lang != "EN" and sub.get("heading_orig"):
        return sub["heading_orig"]
    return sub.get("heading_en") or ""


# --------------------------------------------------------------------------
# Annotation (200-250 words) generation and output


# Keep the annotation prompt well under the org TPM limit (30k for gpt-4o).
ANNOTATION_MAX_PROMPT_TOKENS = 18000


def edited_full_text(sections: list[Section], code: str) -> str:
    """All edited paragraphs of the given language, in document order."""
    parts: list[str] = []
    for section in sections:
        parts.extend(section.intro.get(code) or [])
        for sub in section.subsections:
            parts.extend(sub.get(code) or [])
    return "\n\n".join(parts)


def annotation_source_text(sections: list[Section], code: str) -> str:
    """Edited text condensed for the annotation prompt: when the full text
    would blow the request token limit, keep the section headings and a
    proportional beginning of every section, so the whole video is still
    covered."""
    ratio = CHARS_PER_TOKEN_EN if code == "EN" else CHARS_PER_TOKEN_OTHER
    budget = int(ANNOTATION_MAX_PROMPT_TOKENS * ratio)

    blocks: list[tuple[str, list[str]]] = []
    for section in sections:
        paragraphs: list[str] = list(section.intro.get(code) or [])
        for sub in section.subsections:
            paragraphs.extend(sub.get(code) or [])
        blocks.append((section.heading_en or section.heading, paragraphs))

    total = sum(len(p) for _, paras in blocks for p in paras)
    if total <= budget:
        return "\n\n".join(
            "\n\n".join([f"## {heading}", *paras])
            for heading, paras in blocks
            if paras or heading
        )

    scale = budget / total
    parts: list[str] = []
    for heading, paras in blocks:
        parts.append(f"## {heading}")
        section_budget = int(sum(len(p) for p in paras) * scale)
        used = 0
        for paragraph in paras:
            if used and used + len(paragraph) > section_budget:
                break
            parts.append(paragraph)
            used += len(paragraph)
    return "\n\n".join(parts)


def generate_annotation(
    api_key: str,
    *,
    doc_title: str,
    orig_code: str,
    sections: list[Section],
    terms: list[str],
    usage: dict[str, int],
) -> dict[str, str]:
    languages = annotation_languages(orig_code)
    source_code = orig_code if orig_code != "EN" else "EN"
    text = annotation_source_text(sections, source_code)
    payload = {
        "document_title": doc_title,
        "original_language": LANGUAGE_NAMES.get(orig_code, orig_code),
        "languages": languages,
        "glossary": terms,
        "edited_transcript": text,
    }
    result = chat_json(
        api_key,
        [
            {"role": "system", "content": ANNOTATION_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False, indent=2),
            },
        ],
        usage,
    )
    annotations = {
        code: str((result.get("annotations") or {}).get(code) or "").strip()
        for code in languages
    }
    # The model tends to undershoot the word count; one targeted retry per
    # too-short annotation.
    for code in languages:
        current = annotations.get(code)
        if not current or len(current.split()) >= 200:
            continue
        retry = chat_json(
            api_key,
            [
                {
                    "role": "system",
                    "content": (
                        "Expand the given annotation to 200-250 words "
                        "(STRICT: never fewer than 200 words - count them), "
                        "in the same language, adding concrete details "
                        "from the transcript. Keep the neutral informative "
                        "tone and plain prose. Reply with JSON only: "
                        '{"annotation": "text"}'
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "annotation": current,
                            "edited_transcript": text,
                        },
                        ensure_ascii=False,
                    ),
                },
            ],
            usage,
        )
        expanded = str(retry.get("annotation") or "").strip()
        if len(expanded.split()) > len(current.split()):
            annotations[code] = expanded
    return {code: text_ for code, text_ in annotations.items() if text_}


def annotation_path(playlist_dir: Path, code: str, stem: str) -> Path:
    return playlist_dir / code / OUTPUT_DIRNAME / f"{stem}.txt"


def write_annotation(path: Path, doc_title: str, text: str) -> None:
    """Annotation .txt: the document title, then the text formatted like the
    .md body (justified 80-character lines, no hyphenation, blank lines
    between paragraphs)."""
    lines: list[str] = textwrap.wrap(
        doc_title, width=MD_WIDTH, break_long_words=False,
        break_on_hyphens=False,
    )
    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = " ".join(paragraph.split())
        if paragraph:
            lines += ["", justify_paragraph(paragraph)]
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


# --------------------------------------------------------------------------
# Markdown output


def justify_line(line: str, width: int) -> str:
    """Pad the line to the width by doubling some single spaces (never more
    than doubling); spaces to double are spread evenly along the line."""
    deficit = width - len(line)
    words = line.split(" ")
    gaps = len(words) - 1
    if deficit <= 0 or gaps == 0 or deficit > gaps:
        return line
    chosen: set[int] = set()
    for index in range(deficit):
        # Center the doubled spaces inside their stretch of the line, so the
        # very first gaps are not systematically the widened ones.
        gap = min(gaps - 1, int((index + 0.5) * gaps / deficit))
        while gap in chosen:  # deficit <= gaps, so a free gap always exists
            gap = (gap + 1) % gaps
        chosen.add(gap)
    pieces: list[str] = []
    for index, word in enumerate(words[:-1]):
        pieces.append(word)
        pieces.append("  " if index in chosen else " ")
    pieces.append(words[-1])
    return "".join(pieces)


def justify_paragraph(text: str, width: int = MD_WIDTH) -> str:
    lines = textwrap.wrap(
        " ".join(text.split()),
        width=width,
        break_long_words=False,
        break_on_hyphens=False,
    )
    return "\n".join(
        [justify_line(line, width) for line in lines[:-1]] + lines[-1:]
    ) if lines else ""


def wrap_blockquote(text: str, width: int = MD_WIDTH) -> list[str]:
    lines = textwrap.wrap(
        text, width=width - 2, break_long_words=False, break_on_hyphens=False
    )
    return [f"> {line}" for line in lines] or ["> "]


def build_markdown(
    doc_title: str,
    meta_lines: list[str],
    sections: list[Section],
    lang: str,
) -> str:
    out: list[str] = [f"# {doc_title}", ""]
    for line in meta_lines:
        out.append(line + "  ")
    out += ["", f"## {TOC_HEADING}", ""]
    for number, section in enumerate(sections, start=1):
        out.append(f"- {number}. {section_heading_for(section, lang)}")
        for sub_number, sub in enumerate(section.subsections, start=1):
            out.append(
                f"  - {number}.{sub_number}. "
                f"{subsection_heading_for(sub, lang)}"
            )
    out.append("")
    for number, section in enumerate(sections, start=1):
        out += [f"## {number}. {section_heading_for(section, lang)}", ""]
        slide = section.slide or {}
        for bullet in slide.get("body") or []:
            out += wrap_blockquote(bullet)
        if slide.get("body"):
            out.append("")
        for paragraph in section.intro.get(lang) or []:
            out += [justify_paragraph(paragraph), ""]
        for sub_number, sub in enumerate(section.subsections, start=1):
            out += [
                f"### {number}.{sub_number}. "
                f"{subsection_heading_for(sub, lang)}",
                "",
            ]
            for paragraph in sub.get(lang) or []:
                out += [justify_paragraph(paragraph), ""]
    while out and out[-1] == "":
        out.pop()
    return "\n".join(out) + "\n"


# --------------------------------------------------------------------------
# DOCX / PDF output


def make_docx(
    doc_title: str,
    meta_lines: list[str],
    sections: list[Section],
    lang: str,
    template: Path | None,
    out_path: Path,
) -> None:
    from docx import Document
    from docx.shared import Pt

    if template is not None:
        document = Document(str(template))
        body = document.element.body
        for child in list(body):
            if not child.tag.endswith("}sectPr"):
                body.remove(child)
    else:
        document = Document()

    style_names = {style.name for style in document.styles}

    def pick(*candidates: str) -> str | None:
        for name in candidates:
            if name in style_names:
                return name
        return None

    # Style mapping modeled on RU/Docs_Edited/01_Introduction.V2.docx:
    # document title = heading 2, sections = heading 3, subsections =
    # heading 4, verbatim slide text = ds-markdown-paragraph.
    title_style = pick("Heading 2", "Title", "Heading 1")
    section_style = pick("Heading 3", "Heading 1")
    subsection_style = pick("Heading 4", "Heading 2")
    slide_style = pick("ds-markdown-paragraph")

    if template is None:
        apply_default_formatting(document)

    def para(text: str, style: str | None = None, *, italic: bool = False):
        paragraph = document.add_paragraph()
        if style:
            paragraph.style = document.styles[style]
        run = paragraph.add_run(text)
        if italic:
            run.italic = True
        return paragraph

    para(doc_title, title_style)
    for line in meta_lines:
        para(line)
    # The ToC caption must not use a heading style, or the TOC field below
    # would list the caption itself as an entry.
    caption = para(TOC_HEADING)
    caption_run = caption.runs[0]
    caption_run.bold = True
    caption_size = (
        document.styles[section_style].font.size if section_style else None
    )
    caption_run.font.size = caption_size or Pt(13)
    add_toc_field(
        document,
        heading_level(section_style, 1),
        heading_level(subsection_style, 2),
    )
    request_field_update_on_open(document)
    for section in sections:
        number = f"{section.number}. " if section.number else ""
        para(f"{number}{section_heading_for(section, lang)}", section_style)
        slide = section.slide or {}
        for bullet in slide.get("body") or []:
            para(bullet, slide_style, italic=slide_style is None)
        for paragraph in section.intro.get(lang) or []:
            para(paragraph)
        for index, sub in enumerate(section.subsections, start=1):
            sub_number = (
                f"{section.number}.{index}. " if section.number else ""
            )
            para(
                f"{sub_number}{subsection_heading_for(sub, lang)}",
                subsection_style,
            )
            for paragraph in sub.get(lang) or []:
                para(paragraph)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def heading_level(style_name: str | None, default: int) -> int:
    """Outline level of a built-in "Heading N" style name."""
    match = re.search(r"(\d+)$", style_name or "")
    return int(match.group(1)) if match else default


def add_toc_field(document, top_level: int, bottom_level: int) -> None:
    """Insert a real Word TOC field covering the given heading levels.

    Word/WPS builds it into a table of contents with internal hyperlinks and
    right-aligned page numbers when the fields are updated - the .docx asks
    for that on open (see request_field_update_on_open), and the PDF
    conversion updates the fields before exporting.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' TOC \\o "{top_level}-{bottom_level}" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    for element in (begin, instruction, separate):
        field_run._r.append(element)
    placeholder = paragraph.add_run(
        "Update the fields (Ctrl+A, F9) to build the table of contents."
    )
    placeholder.italic = True
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def request_field_update_on_open(document) -> None:
    """Ask Word/WPS to recalculate fields (the ToC) when the file is opened."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def apply_default_formatting(document) -> None:
    """Default look of the generated .docx (used when --doc is not given).

    - the page number is shown centered in the page header;
    - non-heading paragraph styles get Calibri, 6 pt spacing above and
      below and justified alignment; heading styles keep their spacing.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        name = style.name.lower()
        if name.startswith("heading") or name == "title":
            continue
        style.font.name = "Calibri"
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header = document.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


RPC_E_CALL_REJECTED = -2147418111  # "Call was rejected by callee": Word busy


def com_retry(call, attempts: int = 8, on_reject=None):
    """Run a COM call, waiting out transient 'call rejected' rejections
    (Word keeps servicing its message loop for a while after startup);
    `on_reject` runs before each retry (e.g. to dismiss a nag dialog)."""
    for attempt in range(1, attempts + 1):
        try:
            return call()
        except Exception as exc:
            if (
                getattr(exc, "hresult", None) != RPC_E_CALL_REJECTED
                or attempt == attempts
            ):
                raise
            if on_reject is not None:
                try:
                    on_reject()
                except Exception:
                    pass
            time.sleep(attempt)


def dismiss_word_dialogs(pids: set[int]) -> None:
    """Close nag dialogs of the given (hidden) Word processes.

    Word occasionally pops modal prompts even in automation - e.g. "Word
    is not your default program for documents, choose file types?" - and
    then rejects every COM call until somebody answers. Press the safest
    button (No/Cancel) or close the dialog.
    """
    import win32con
    import win32gui
    import win32process

    BM_CLICK = 0x00F5
    PREFERRED = ("нет", "no", "отмена", "cancel", "закрыть", "close")

    def handle_dialog(hwnd) -> None:
        buttons: list[tuple[int, str]] = []

        def collect(child, _):
            if win32gui.GetClassName(child) == "Button":
                text = (
                    win32gui.GetWindowText(child)
                    .replace("&", "").strip().lower()
                )
                buttons.append((child, text))
            return True

        try:
            win32gui.EnumChildWindows(hwnd, collect, None)
        except Exception:
            pass
        for preferred in PREFERRED:
            for handle, text in buttons:
                if preferred in text:
                    win32gui.PostMessage(handle, BM_CLICK, 0, 0)
                    return
        win32gui.PostMessage(hwnd, win32con.WM_CLOSE, 0, 0)

    def visit(hwnd, _):
        if (
            win32gui.IsWindowVisible(hwnd)
            and win32gui.GetClassName(hwnd) == "#32770"
            and win32process.GetWindowThreadProcessId(hwnd)[1] in pids
        ):
            handle_dialog(hwnd)
        return True

    win32gui.EnumWindows(visit, None)


def wps_progid() -> str | None:
    """COM ProgID of WPS Writer when WPS Office is installed, else None."""
    import winreg

    for progid in ("Kwps.Application", "wps.Application"):
        try:
            winreg.CloseKey(
                winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, progid)
            )
            return progid
        except OSError:
            continue
    return None


def run_tasklist(args: list[str]) -> str:
    """tasklist output, decoded safely.

    tasklist writes in the console OEM code page; the default locale codec
    (e.g. cp1252) chokes on Cyrillic window titles or localized messages,
    the decode error kills subprocess' reader thread and stdout comes back
    as None - so decode as OEM with replacement.
    """
    import subprocess

    try:
        return subprocess.run(
            ["tasklist", *args],
            capture_output=True,
            text=True,
            encoding="oem",
            errors="replace",
            timeout=30,
        ).stdout or ""
    except Exception:
        return ""


def image_pids(image_name: str) -> set[int]:
    """PIDs of the running processes with the given executable name."""
    listing = run_tasklist(
        ["/FI", f"IMAGENAME eq {image_name}", "/FO", "CSV", "/NH"]
    )
    pids: set[int] = set()
    for line in listing.splitlines():
        cells = [cell.strip('"') for cell in line.strip().split('","')]
        if len(cells) >= 2 and cells[0].lower() == image_name.lower():
            try:
                pids.add(int(cells[1]))
            except ValueError:
                pass
    return pids


def update_com_fields(document) -> None:
    """Rebuild the tables of contents and other fields of an open document
    (Word or WPS COM), so the exported PDF gets ToC page numbers and
    hyperlinks. Best-effort: a document without fields is left as is."""
    try:
        tocs = document.TablesOfContents
        for index in range(1, int(tocs.Count) + 1):
            tocs(index).Update()
    except Exception:
        pass
    try:
        document.Fields.Update()
    except Exception:
        pass


def convert_docx_to_pdf_wps(docx_path: Path, pdf_path: Path,
                            progid: str) -> int:
    """One PDF conversion attempt via WPS Writer COM (Word-compatible API).

    WPS shows no activation nags and does not conflict with a Microsoft
    Word window the user is working in, so it is preferred over Word when
    installed. The application is Quit() only when this call actually
    started the wps.exe process - if a running WPS instance was reused,
    quitting would close the user's own documents.
    """
    WD_FORMAT_PDF = 17
    import pythoncom
    import win32com.client

    before = image_pids("wps.exe")
    pythoncom.CoInitialize()
    try:
        app = win32com.client.DispatchEx(progid)
        started = image_pids("wps.exe") - before
        try:
            app.Visible = False
            try:
                app.DisplayAlerts = 0
            except Exception:
                pass
            document = app.Documents.Open(
                str(docx_path), ReadOnly=False, AddToRecentFiles=False
            )
            try:
                update_com_fields(document)
                document.ExportAsFixedFormat(str(pdf_path), WD_FORMAT_PDF)
            finally:
                document.Close(False)
        finally:
            if started:
                try:
                    app.Quit()
                except Exception:
                    pass
    finally:
        pythoncom.CoUninitialize()
    return 0 if pdf_path.is_file() else 1


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> int:
    """One PDF conversion attempt; runs in a child process.

    WPS Writer is used when installed (no activation nags, no clashes with
    a user's Word session); otherwise - or if WPS fails - Microsoft Word
    COM is the fallback.
    """
    # COM servers resolve relative paths against their own cwd, not ours.
    docx_path, pdf_path = docx_path.resolve(), pdf_path.resolve()
    progid = wps_progid()
    if progid is not None:
        try:
            if convert_docx_to_pdf_wps(docx_path, pdf_path, progid) == 0:
                return 0
        except Exception as exc:
            print(f"WPS conversion failed ({exc}); trying Word...",
                  file=sys.stderr, flush=True)
    return convert_docx_to_pdf_word(docx_path, pdf_path)


def convert_docx_to_pdf_word(docx_path: Path, pdf_path: Path) -> int:
    """One PDF conversion attempt via Word COM; runs in a child process.

    A dedicated Word instance (DispatchEx) is created and always Quit() in
    the end, so no zombie WINWORD process stays behind holding file locks,
    and a Word instance the user works in is never used (docx2pdf reuses a
    shared instance and is prone to both problems).

    Early binding (a generated type-library wrapper) is required: dynamic
    dispatch proved unreliable for Word Document objects on this setup -
    GetIDsOfNames fails with "AttributeError: Open.SaveAs". A corrupted
    gen_py cache (e.g. half-written by a crashed run) is dropped so that
    the next attempt regenerates it.
    """
    WD_FORMAT_PDF = 17
    import pythoncom
    import win32com.client

    drop_word_resiliency_entries(WORKSPACE_ROOT)
    before = winword_pids()
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        mine = winword_pids() - before  # the instance just started
        unblock = lambda: dismiss_word_dialogs(mine)  # noqa: E731
        try:
            word.Visible = False
            word.DisplayAlerts = 0
            document = com_retry(
                lambda: word.Documents.Open(
                    str(docx_path), ReadOnly=False, AddToRecentFiles=False
                ),
                on_reject=unblock,
            )
            try:
                com_retry(
                    lambda: update_com_fields(document), on_reject=unblock
                )
                com_retry(
                    lambda: document.SaveAs(
                        str(pdf_path), FileFormat=WD_FORMAT_PDF
                    ),
                    on_reject=unblock,
                )
            except AttributeError:
                # Broken name resolution: drop the (likely corrupted)
                # gen_py cache; the retry in the parent starts a fresh
                # child that rebuilds it.
                clear_gen_py_cache()
                raise
            com_retry(lambda: document.Close(False), on_reject=unblock)
        finally:
            try:
                com_retry(lambda: word.Quit(), attempts=3, on_reject=unblock)
            except Exception:
                pass
            # A Word that refused to Quit would linger, hold file locks and
            # make every following conversion fail too ("call rejected"
            # avalanche) - kill this child's own instance if still alive.
            kill_winword(mine)
    finally:
        pythoncom.CoUninitialize()
    return 0 if pdf_path.is_file() else 1


def drop_word_resiliency_entries(scope: Path) -> None:
    """Remove Word's crash-bookkeeping registry entries for our documents.

    When a WINWORD process dies while a document is open, Word records the
    document under HKCU ...\\Word\\Resiliency (DisabledItems /
    DocumentRecovery). On the next automated start Word then waits forever
    on a hidden prompt ("open the problem document?" / recovery / safe
    mode) and every COM call fails with 'Call was rejected by callee'.
    Entries for any document inside `scope` (the workspace) are removed;
    the user's own documents elsewhere are not touched.
    """
    import winreg

    target = str(scope).lower()

    def mentions_target(data) -> bool:
        return isinstance(data, bytes) and target in data.decode(
            "utf-16-le", errors="ignore"
        ).lower()

    def matching_values(key) -> list[str]:
        names = []
        index = 0
        while True:
            try:
                name, data, _ = winreg.EnumValue(key, index)
            except OSError:
                return names
            if mentions_target(data):
                names.append(name)
            index += 1

    def data_mentions_target(key) -> bool:
        return bool(matching_values(key))

    def subkeys(key) -> list[str]:
        names = []
        index = 0
        while True:
            try:
                names.append(winreg.EnumKey(key, index))
            except OSError:
                return names
            index += 1

    try:
        with winreg.OpenKey(
            winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Office"
        ) as office:
            versions = subkeys(office)
    except OSError:
        return
    for version in versions:
        base = (
            rf"Software\Microsoft\Office\{version}\Word\Resiliency"
        )
        # DisabledItems: flat values, the file path inside the binary data.
        try:
            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                base + r"\DisabledItems",
                0,
                winreg.KEY_ALL_ACCESS,
            ) as key:
                for name in matching_values(key):
                    winreg.DeleteValue(key, name)
        except OSError:
            pass
        # DocumentRecovery: one subkey per remembered document.
        try:
            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                base + r"\DocumentRecovery",
                0,
                winreg.KEY_ALL_ACCESS,
            ) as key:
                for name in subkeys(key):
                    with winreg.OpenKey(key, name) as sub:
                        drop = data_mentions_target(sub)
                    if drop:
                        winreg.DeleteKey(key, name)
        except OSError:
            pass
        # StartupItems: "Word crashed while starting" markers; they make a
        # hidden Word ask "start in safe mode?" and reject all COM calls.
        try:
            winreg.DeleteKey(
                winreg.HKEY_CURRENT_USER, base + r"\StartupItems"
            )
        except OSError:
            pass


def clear_gen_py_cache() -> None:
    import shutil

    try:
        import win32com

        shutil.rmtree(win32com.__gen_path__, ignore_errors=True)
    except Exception:
        pass


def winword_processes() -> dict[int, str]:
    """PID -> window title for the running WINWORD processes ("N/A" for
    hidden automation instances without a window)."""
    listing = run_tasklist(
        ["/V", "/FI", "IMAGENAME eq WINWORD.EXE", "/FO", "CSV", "/NH"]
    )
    processes: dict[int, str] = {}
    for line in listing.splitlines():
        cells = [cell.strip('"') for cell in line.strip().split('","')]
        if len(cells) >= 2 and cells[0].upper() == "WINWORD.EXE":
            try:
                processes[int(cells[1])] = cells[-1]
            except ValueError:
                pass
    return processes


def winword_pids() -> set[int]:
    return set(winword_processes())


# Window titles of hidden automation Word instances (no real document
# window): tasklist reports "N/A" or Word's internal helper window.
HIDDEN_WORD_TITLES = ("", "N/A", "HardwareMonitorWindow")


def user_word_running() -> bool:
    """True when a Word with a visible document window (the user's) runs."""
    return any(
        title not in HIDDEN_WORD_TITLES
        for title in winword_processes().values()
    )


def kill_winword(pids: set[int]) -> None:
    """Force-kill the given WINWORD processes, sparing windowed (user) ones."""
    import subprocess

    for pid, title in winword_processes().items():
        if pid in pids and title in HIDDEN_WORD_TITLES:
            subprocess.run(
                ["taskkill", "/PID", str(pid), "/T", "/F"],
                capture_output=True,
            )


def make_pdf(docx_path: Path, pdf_path: Path, attempts: int = 3) -> bool:
    """Render the .docx to PDF (WPS Writer, or Word); False on failure.

    Office COM occasionally dies with a native access violation (0xC0000005)
    that would kill the whole script and lose the money already spent on
    editing, so the conversion runs in a disposable child process. A Word
    instance left behind by a crashed attempt is killed before the retry -
    but only Word processes that appeared during the attempt, so a Word the
    user is working in is never touched.
    """
    import subprocess

    # Hidden Word instances surviving from earlier crashes make new
    # conversions fail too ("call rejected" avalanche) - clean them first.
    # Irrelevant when WPS does the conversion, so skip in that case.
    if wps_progid() is None:
        kill_winword(winword_pids())

    last_note = ""
    for attempt in range(1, attempts + 1):
        before = winword_pids()
        try:
            child = subprocess.run(
                [
                    sys.executable,
                    str(Path(__file__).resolve()),
                    "--convert-pdf",
                    str(docx_path),
                    str(pdf_path),
                ],
                capture_output=True,
                text=True,
                # The child writes UTF-8 (see the reconfigure at the top);
                # without this the parent would decode it in the ANSI code
                # page and lose the error message to a UnicodeDecodeError.
                encoding="utf-8",
                errors="replace",
                timeout=300,
            )
            failed = child.returncode != 0
            last_note = (
                (child.stderr or child.stdout).strip().splitlines()[-1]
                if failed and (child.stderr or child.stdout).strip()
                else f"exit code {child.returncode}"
            )
        except subprocess.TimeoutExpired:
            failed, last_note = True, "conversion timed out"
        if not failed and pdf_path.is_file():
            return True
        # Clean up Word instances left by the failed attempt; a windowed
        # Word (the user's, even one opened during the attempt) is spared.
        kill_winword(winword_pids() - before)
        if attempt < attempts:
            time.sleep(3)
    print(f"  WARNING: PDF conversion failed: {last_note}", flush=True)
    if wps_progid() is None and user_word_running():
        print(
            "  NOTE: a Word window is open; working in Word can block the "
            "hidden conversion. Close Word and re-run the script - missing "
            "PDFs are then rebuilt from the .docx at no API cost.",
            flush=True,
        )
    return False


# --------------------------------------------------------------------------
# Selection and session


def output_files(lang_dir: Path, stem: str) -> dict[str, Path]:
    out = lang_dir / OUTPUT_DIRNAME
    return {
        "md": out / f"{stem}.md",
        "docx": out / f"{stem}.docx",
        "pdf": out / f"{stem}.pdf",
    }


def transcripts_ready(lang_dir: Path, stem: str) -> bool:
    return (lang_dir / f"{stem}.srt").is_file()


def is_processed(
    playlist_dir: Path,
    stem: str,
    lang_codes: list[str],
    *,
    annotate: bool = False,
    orig_code: str = "EN",
) -> bool:
    """md + docx in every language OUTPUT folder (pdf is best-effort);
    with --annotate also the annotation .txt files."""
    docs_done = all(
        output_files(playlist_dir / code, stem)[kind].is_file()
        for code in lang_codes
        for kind in ("md", "docx")
    )
    if not docs_done:
        return False
    if annotate:
        return all(
            annotation_path(playlist_dir, code, stem).is_file()
            for code in annotation_languages(orig_code)
        )
    return True


# --------------------------------------------------------------------------
# Editing cache and cost estimate


def section_cache_key(section: Section) -> str:
    return f"{section.heading} @ {section.start:.1f}"


def load_edit_cache(cache_path: Path, stem: str) -> dict[str, dict]:
    try:
        data = json.loads(cache_path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return {}
    return data.get("sections") or {} if data.get("video") == stem else {}


def section_to_cache(
    section: Section, issues: Coverage, *, split_tried: bool = False
) -> dict:
    entry = {
        "heading_en": section.heading_en,
        "intro": section.intro,
        "subsections": section.subsections,
        # The cache key carries the start of the section but not its end, and
        # the boundary refinement moves both; without the range a section that
        # merely lost its last sentences would keep an edit that still has
        # them - and so would the section they moved to.
        "range": section_range(section),
        "rules": EDIT_RULES_VERSION,
        # Recorded so that a section whose gaps survived the retry is not
        # re-edited (and re-paid for) on every later run.
        "coverage": issues.as_cache(),
    }
    if split_tried:
        # Everything that could be tried on this section was: it is not to
        # be paid for again, however much of it is missing.
        entry["split_tried"] = True
    return entry


def cached_lost_share(entry: dict, source: str) -> float:
    """How much of the transcript the cached edit is missing, 0..1."""
    recorded = entry.get("coverage") or {}
    lost = sum(
        len(fragment.split())
        for key in ("dropped", "repeated")
        for fragment in recorded.get(key) or []
    )
    total = len(coverage_words(source))
    return lost / total if total else 0.0


def section_range(section: Section) -> list[float]:
    return [round(section.start, 1), round(section.end, 1)]


def stale_range(section: Section, entry: dict) -> bool:
    """Whether the cached edit was made for a different stretch of the video.

    Entries written before the range was recorded are trusted unless this run
    moved a boundary of the section.
    """
    cached = entry.get("range")
    if cached is None:
        return section.moved
    return list(cached) != section_range(section)


def section_from_cache(section: Section, entry: dict) -> None:
    section.heading_en = entry.get("heading_en") or section.heading
    section.intro = entry.get("intro") or {}
    section.subsections = entry.get("subsections") or []


def cached_coverage(
    section: Section, entry: dict, source: str, code: str
) -> Coverage | None:
    """Coverage of a cached section, or None when it was already checked
    when it was written."""
    if "coverage" in entry:
        return None
    probe = Section(section.heading, section.number, section.slide,
                    section.start, section.end)
    section_from_cache(probe, entry)
    return section_coverage(probe, source, code)


def save_edit_cache(
    cache_path: Path, stem: str, cache: dict[str, dict]
) -> None:
    cache_path.write_text(
        json.dumps(
            {"video": stem, "model": MODEL, "sections": cache},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def estimate_tokens(text: str, code: str) -> int:
    ratio = CHARS_PER_TOKEN_EN if code == "EN" else CHARS_PER_TOKEN_OTHER
    return int(len(text) / ratio)


def estimate_cost(
    pending: list[tuple[Section, str, str]], orig_code: str, terms: list[str]
) -> tuple[int, int, float]:
    """(prompt tokens, completion tokens, USD) for the sections still to edit.

    The prompt carries both transcripts, the glossary and a fixed
    per-section overhead; the completion re-emits the edited text in both
    languages, so it is close to the size of the transcripts themselves.
    """
    prompt = completion = 0
    glossary_tokens = estimate_tokens(", ".join(terms), "EN")
    for _, orig_text, en_text in pending:
        text_tokens = (
            estimate_tokens(orig_text, orig_code)
            + estimate_tokens(en_text, "EN")
        )
        prompt += SECTION_OVERHEAD_TOKENS + glossary_tokens + text_tokens
        completion += text_tokens
    cost = (
        prompt * USD_PER_MTOKEN_PROMPT
        + completion * USD_PER_MTOKEN_COMPLETION
    ) / 1_000_000
    return prompt, completion, cost


def process_video(
    video: Path,
    *,
    playlist_dir: Path,
    slides_folder: Path | None,
    api_key: str,
    course: str,
    orig_code: str,
    lang_codes: list[str],
    terms: list[str],
    doc_template: Path | None,
    pdf_template: Path | None,
    auto_yes: bool = False,
    annotate: bool = False,
) -> bool:
    """Process one video; False when the user declined the cost estimate.

    With slides (slides_folder holds slides.json) the document structure
    comes from the slides; otherwise (--no-slides) the title is built in the
    summary format and the sections come from the video timecodes, both
    taken from INFO/<stem>.json.
    """
    slides_data = None
    if slides_folder is not None and (slides_folder / RESULT_FILENAME).is_file():
        slides_data = json.loads(
            (slides_folder / RESULT_FILENAME).read_text(encoding="utf-8")
        )

    segments = {
        code: parse_srt(playlist_dir / code / f"{video.stem}.srt")
        for code in lang_codes
    }
    video_end = max(
        (segs[-1].end for segs in segments.values() if segs), default=0.0
    )

    if slides_data is not None:
        document = slides_data.get("document") or {}
        doc_title = document.get("title") or video.stem
        meta = document.get("meta") or {}
        meta_lines = [
            f"{label}: {value}"
            for label, value in (
                ("Date", meta.get("date")),
                ("Recorded by", meta.get("recorded_by")),
                ("Language", meta.get("language")),
            )
            if value
        ]
        sections = build_sections(slides_data, video_end)
        cache_path = slides_folder / EDIT_CACHE_FILENAME
    else:
        info = load_video_info(playlist_dir, video.stem)
        doc_title, meta_lines = info_document_header(info, video.stem)
        sections = build_sections_from_info(info, video_end, orig_code)
        info_dir = playlist_dir / INFO_DIRNAME
        info_dir.mkdir(parents=True, exist_ok=True)
        cache_path = info_dir / f"{video.stem}.{EDIT_CACHE_FILENAME}"

    # The pauses of the recording (see silences.py) place both the section
    # boundaries and the paragraphs; without the media file both fall back
    # to what the .srt alone can tell.
    pauses = SilenceIndex(load_silences(video))
    if not pauses:
        print(
            "  Pauses: no silence data for this video - paragraphs are cut "
            "by length.",
            flush=True,
        )

    moved = refine_section_boundaries(sections, segments, orig_code, pauses)
    if moved:
        shifts = sorted(shift for _, shift in moved)
        print(
            f"  Boundaries: {len(moved)} of {len(sections) - 1} moved off a "
            f"sentence or an announcement "
            f"({shifts[0]:+.1f}..{shifts[-1]:+.1f} s).",
            flush=True,
        )

    texts: dict[str, tuple[str, str]] = {}
    for section in sections:
        orig_text = (
            section_text(
                segments[orig_code], section.start, section.end, pauses
            )
            if orig_code != "EN"
            else ""
        )
        en_text = (
            section_text(segments["EN"], section.start, section.end, pauses)
            if "EN" in segments
            else ""
        )
        texts[section_cache_key(section)] = (orig_text, en_text)
        section.anchors = bullet_anchors(
            [str(line) for line in (section.slide or {}).get("body") or []],
            sentences_of_section(
                segments[orig_code], section.start, section.end
            ),
            sentences_of_section(
                segments.get("EN") or segments[orig_code],
                section.start,
                section.end,
            ),
        )
    anchored = sum(1 for section in sections if any(section.anchors))
    if anchored:
        found = sum(sum(1 for a in s.anchors if a) for s in sections)
        print(
            f"  Bullets: {found} named in the speech, in {anchored} "
            f"section(s) - the subsections start there.",
            flush=True,
        )

    cache = load_edit_cache(cache_path, video.stem)
    # Sections edited before the coverage check existed (or by an older
    # model) are re-checked offline; the ones that lost or duplicated text
    # are dropped from the cache, so they are edited again below and show up
    # in the cost estimate.
    stale = 0
    check_code = coverage_code(orig_code)
    for section in sections:
        key = section_cache_key(section)
        entry = cache.get(key)
        if entry is None:
            continue
        if stale_range(section, entry):
            print(
                f"  Re-editing '{section.heading}': its boundaries moved.",
                flush=True,
            )
            del cache[key]
            stale += 1
            continue
        if entry.get("rules", 1) != EDIT_RULES_VERSION:
            print(
                f"  Re-editing '{section.heading}': the editing rules "
                "changed since it was written.",
                flush=True,
            )
            del cache[key]
            stale += 1
            continue
        if cache_incomplete(entry):
            print(
                f"  Re-editing '{section.heading}': the cached version still "
                "has dropped or duplicated transcript text.",
                flush=True,
            )
            del cache[key]
            stale += 1
            continue
        source = texts[key][0] if orig_code != "EN" else texts[key][1]
        lost = cached_lost_share(entry, source)
        if lost >= COVERAGE_SPLIT_SHARE and not entry.get("split_tried"):
            print(
                f"  Re-editing '{section.heading}': {lost:.0%} of it is "
                "missing from the cached version, and a section that loses "
                "its tail is now edited in parts.",
                flush=True,
            )
            del cache[key]
            stale += 1
            continue
        issues = cached_coverage(section, entry, source, check_code)
        if issues is None:
            continue
        if issues.clean:
            entry["coverage"] = issues.as_cache()
        else:
            print(
                f"  Re-editing '{section.heading}': the cached version has "
                f"{issues.summary()}.",
                flush=True,
            )
            del cache[key]
            stale += 1

    pending = [
        (section, *texts[section_cache_key(section)])
        for section in sections
        if section_cache_key(section) not in cache
    ]

    prompt_est, completion_est, cost_est = estimate_cost(
        pending, orig_code, terms
    )
    annotation_cached = bool(
        (cache.get(ANNOTATION_CACHE_KEY) or {}).get("texts")
    )
    if annotate and not annotation_cached:
        # One extra call: the whole edited original text in the prompt,
        # 200-250 words per requested language in the completion.
        source = orig_code if orig_code != "EN" else "EN"
        full_len = sum(len(t[0] if source != "EN" else t[1])
                       for t in texts.values())
        ann_prompt = min(
            int(
                full_len / (CHARS_PER_TOKEN_EN if source == "EN"
                            else CHARS_PER_TOKEN_OTHER)
            ),
            ANNOTATION_MAX_PROMPT_TOKENS,
        ) + 400
        ann_completion = 400 * len(annotation_languages(orig_code))
        prompt_est += ann_prompt
        completion_est += ann_completion
        cost_est += (
            ann_prompt * USD_PER_MTOKEN_PROMPT
            + ann_completion * USD_PER_MTOKEN_COMPLETION
        ) / 1_000_000
    print(
        f"  Sections: {len(sections)} total, "
        f"{len(sections) - len(pending)} already edited (cached), "
        f"{len(pending)} to edit via the API"
        + (f" ({stale} of them dropped from the cache)" if stale else "")
        + (
            f"; annotation ({'/'.join(annotation_languages(orig_code))}): "
            + ("cached" if annotation_cached else "to generate")
            if annotate
            else ""
        )
        + ".",
        flush=True,
    )
    print(
        f"  Estimated cost: ~{prompt_est} prompt + ~{completion_est} "
        f"completion tokens -> ~${cost_est:.2f} ({MODEL}).",
        flush=True,
    )
    if not auto_yes:
        print("  Create the final documents for this video? (y/n)",
              flush=True)
        try:
            answer = input().strip().lower()
        except EOFError:
            answer = ""
        if answer not in ("y", "yes"):
            return False

    usage: dict[str, int] = {}
    incomplete: list[str] = []

    def keep_the_spoken_order(section: Section, key: str) -> None:
        section_orig, section_en = texts[key]
        moved = paragraphs_in_spoken_order(
            section, {orig_code: section_orig or section_en, "EN": section_en}
        )
        if moved:
            print(
                f"    Order: {moved} word(s) the editor moved were put back "
                f"where they were said.",
                flush=True,
            )

    for index, section in enumerate(sections, start=1):
        key = section_cache_key(section)
        minutes = (section.end - section.start) / 60.0
        if key in cache:
            print(
                f"  [{index}/{len(sections)}] {section.heading} "
                f"({minutes:.1f} min): cached.",
                flush=True,
            )
            section_from_cache(section, cache[key])
            keep_the_spoken_order(section, key)
            polish_section(section, check_code)
            continue
        print(
            f"  [{index}/{len(sections)}] {section.heading} "
            f"({minutes:.1f} min)...",
            flush=True,
        )
        orig_text, en_text = texts[key]
        parts = chunk_count(orig_text, en_text, orig_code)
        source_text = en_text if orig_code == "EN" else orig_text
        if parts == 1 and section.slide is None:
            words = len(coverage_words(source_text))
            if words >= PROACTIVE_FLAT_SPLIT_WORDS:
                parts = max(2, words // PROACTIVE_FLAT_SPLIT_WORDS)
        split_tried = parts > 1
        if parts > 1:
            issues = edit_section_in_chunks(
                api_key,
                section,
                course=course,
                doc_title=doc_title,
                orig_code=orig_code,
                segments=segments,
                parts=parts,
                terms=terms,
                usage=usage,
                pauses=pauses,
            )
        else:
            issues = edit_section(
                api_key,
                section,
                course=course,
                doc_title=doc_title,
                orig_code=orig_code,
                orig_text=orig_text,
                en_text=en_text,
                terms=terms,
                usage=usage,
            )
            again = parts_for_a_losing_section(
                issues, orig_text, en_text, orig_code
            )
            if again > 1:
                split_tried = True
                issues = edit_section_again_in_parts(
                    api_key,
                    section,
                    issues,
                    course=course,
                    doc_title=doc_title,
                    orig_code=orig_code,
                    orig_text=orig_text,
                    en_text=en_text,
                    segments=segments,
                    parts=again,
                    terms=terms,
                    usage=usage,
                    pauses=pauses,
                )
        if not issues.clean:
            incomplete.append(f"{section.heading}: {issues.summary()}")
        keep_the_spoken_order(section, key)
        polish_section(section, check_code)
        # Persist after every section: a crash later in the run (e.g. during
        # the PDF stage) must not lose paid editing results.
        cache[key] = section_to_cache(
            section, issues, split_tried=split_tried
        )
        save_edit_cache(cache_path, video.stem, cache)

    annotations: dict[str, str] = {}
    if annotate:
        annotations = (cache.get(ANNOTATION_CACHE_KEY) or {}).get("texts") or {}
        needed = annotation_languages(orig_code)
        if not all(code in annotations for code in needed):
            print(
                f"  Annotation ({'/'.join(needed)}): generating...",
                flush=True,
            )
            annotations = generate_annotation(
                api_key,
                doc_title=doc_title,
                orig_code=orig_code,
                sections=sections,
                terms=terms,
                usage=usage,
            )
            cache[ANNOTATION_CACHE_KEY] = {"texts": annotations}
            save_edit_cache(cache_path, video.stem, cache)
        else:
            print(
                f"  Annotation ({'/'.join(needed)}): cached.", flush=True
            )
        # Write the annotation files before the docs stage: a failure there
        # (e.g. PDF conversion) must not lose the generated annotation.
        for code in annotation_languages(orig_code):
            text = annotations.get(code)
            if not text:
                print(
                    f"  WARNING: no {code} annotation was generated.",
                    flush=True,
                )
                continue
            path = annotation_path(playlist_dir, code, video.stem)
            write_annotation(path, doc_title, text)
            print(
                f"  Saved: {path.relative_to(playlist_dir)}", flush=True
            )

    for code in lang_codes:
        files = output_files(playlist_dir / code, video.stem)
        markdown = build_markdown(doc_title, meta_lines, sections, code)
        files["md"].parent.mkdir(parents=True, exist_ok=True)
        files["md"].write_text(markdown, encoding="utf-8")
        make_docx(
            doc_title, meta_lines, sections, code, doc_template, files["docx"]
        )
        if pdf_template is not None:
            tmp_docx = files["pdf"].with_suffix(".pdf.tmp.docx")
            make_docx(
                doc_title, meta_lines, sections, code, pdf_template, tmp_docx
            )
            if make_pdf(tmp_docx, files["pdf"]):
                tmp_docx.unlink(missing_ok=True)
        else:
            make_pdf(files["docx"], files["pdf"])
        for kind in ("md", "docx", "pdf"):
            if files[kind].is_file():
                print(
                    f"  Saved: {files[kind].relative_to(playlist_dir)}",
                    flush=True,
                )
    if incomplete:
        print(
            f"  WARNING: {len(incomplete)} section(s) did not pass the "
            f"coverage check even after a re-edit:",
            flush=True,
        )
        for line in incomplete:
            print(f"    {line}", flush=True)
    overlong = long_paragraphs(sections, coverage_code(orig_code))
    if overlong:
        worst = sorted(overlong, key=lambda item: -item[1])[:5]
        print(
            f"  NOTE: {len(overlong)} paragraph(s) over "
            f"{PARAGRAPH_REPORT_WORDS} words:",
            flush=True,
        )
        for heading, words in worst:
            print(f"    {words} words in '{heading}'", flush=True)
    density = punctuation_density(sections, coverage_code(orig_code))
    if density:
        print(f"  Punctuation: {density}.", flush=True)
    unmarked = questions_with_a_full_stop(sections, coverage_code(orig_code))
    if unmarked:
        print(
            f"  NOTE: {len(unmarked)} question(s) written with a full stop:",
            flush=True,
        )
        for sentence in unmarked[:5]:
            print(f"    {shorten_fragment(sentence)}", flush=True)
    actual_cost = (
        usage.get("prompt_tokens", 0) * USD_PER_MTOKEN_PROMPT
        + usage.get("completion_tokens", 0) * USD_PER_MTOKEN_COMPLETION
    ) / 1_000_000
    print(
        f"  Tokens used: {usage.get('prompt_tokens', 0)} prompt + "
        f"{usage.get('completion_tokens', 0)} completion "
        f"(~${actual_cost:.2f}).",
        flush=True,
    )
    return True


SETTINGS_KEYS = ("CHANNEL", "PLAYLIST", "NEXT", "VIDEO")


def apply_run_settings(args: argparse.Namespace) -> None:
    """Fill channel / playlist / session size from the bat's settings file."""
    if not args.settings:
        if not (args.channel_folder or "").strip():
            raise SystemExit(
                "channel_folder is required (pass it or --settings)."
            )
        if not (args.playlist_folder or "").strip():
            raise SystemExit(
                "playlist_folder is required (pass it or --settings)."
            )
        return
    settings = read_settings(args.settings, SETTINGS_KEYS)
    print(f"Settings: {args.settings}", flush=True)
    if not (args.channel_folder or "").strip():
        args.channel_folder = settings.get("CHANNEL", "")
    if not (args.playlist_folder or "").strip():
        args.playlist_folder = settings.get("PLAYLIST", "")
    if not args.video:
        video = settings.get("VIDEO", "").strip()
        if video:
            args.video = video
    if args.video is None:
        next_raw = settings.get("NEXT", "").strip()
        if next_raw:
            args.next_count = int(next_raw)
    if not (args.channel_folder or "").strip():
        raise SystemExit(f"CHANNEL is empty. Set it in {args.settings}.")
    if not (args.playlist_folder or "").strip():
        raise SystemExit(f"PLAYLIST is empty. Set it in {args.settings}.")
    print(f"Channel: {args.channel_folder}", flush=True)
    print(f"Playlist: {args.playlist_folder}", flush=True)
    if args.video:
        print(f"Video: {args.video}", flush=True)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create final edited documents from transcripts + slides."
    )
    parser.add_argument(
        "channel_folder",
        nargs="?",
        default=None,
        help="Channel ref under _channels/ (e.g. _Autotesting or "
        "AI_for_Game_Design\\_BuildingAeon); may come from --settings",
    )
    parser.add_argument(
        "playlist_folder",
        nargs="?",
        default=None,
        help="Playlist folder name under <channel>/_playlists (e.g. lectures); "
        "may come from --settings",
    )
    parser.add_argument(
        "--next",
        dest="next_count",
        type=int,
        default=1,
        metavar="N",
        help="How many videos to process this session (default: 1)",
    )
    parser.add_argument(
        "--doc",
        type=Path,
        default=None,
        metavar="DOCX",
        help=(
            "A .docx file whose styles are used for the generated .docx "
            "(default: built-in styles)"
        ),
    )
    parser.add_argument(
        "--pdf",
        type=Path,
        default=None,
        metavar="DOCX",
        help=(
            "A .docx style template used only for the PDF rendering "
            "(default: the PDF is rendered from the generated .docx)"
        ),
    )
    parser.add_argument(
        "--orig-only",
        action="store_true",
        help=(
            "Create the OUTPUT documents only for the original language "
            "(skip the English versions)"
        ),
    )
    parser.add_argument(
        "--no-slides",
        action="store_true",
        help=(
            "Also process videos without slides.json: the document title is "
            "built in the summary format and the ToC comes from the video "
            "timecodes (INFO/<stem>.json saved by transcribe_videos.py); "
            "videos that do have slides.json still use the slides"
        ),
    )
    parser.add_argument(
        "--annotate",
        action="store_true",
        help=(
            "Also create a 200-250 word annotation .txt per video: Russian "
            "original -> Russian annotation in RU/OUTPUT; any other "
            "original -> English annotation in EN/OUTPUT plus a Russian "
            "translation in RU/OUTPUT"
        ),
    )
    parser.add_argument(
        "--video",
        default=None,
        metavar="STEM",
        help=(
            "Process only the video with this file stem (exact name without "
            "extension); overrides --next"
        ),
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help=(
            "Re-create the documents of videos that already have OUTPUT "
            "files (the cached editing is reused, except for sections that "
            "fail the coverage check)"
        ),
    )
    parser.add_argument(
        "--terms",
        type=Path,
        default=None,
        metavar="FILE",
        help=(
            "Course glossary the editor spells terms by, one term per line "
            f"(default: {GLOSSARY_FILENAME} of the playlist folder, then of "
            "the channel folder)"
        ),
    )
    parser.add_argument(
        "--yes",
        action="store_true",
        help="Skip the per-video cost confirmation prompt",
    )
    parser.add_argument(
        "--workspace",
        type=Path,
        default=WORKSPACE_ROOT,
        help="Workspace root (default: parent of src/)",
    )
    parser.add_argument(
        "--settings",
        type=Path,
        default=None,
        metavar="FILE",
        help=(
            "Settings file next to a channel bat (CHANNEL, PLAYLIST; optional "
            "NEXT, VIDEO); see shared/settings_file.py"
        ),
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    apply_run_settings(args)
    if args.next_count < 1:
        raise SystemExit("--next must be a positive number")
    for flag, path in (("--doc", args.doc), ("--pdf", args.pdf)):
        if path is not None and not path.is_file():
            raise SystemExit(f"{flag} template not found: {path}")

    try:
        channel_dir = require_channel_ref(
            channels_dir(args.workspace), args.channel_folder
        )
    except FileNotFoundError as exc:
        raise SystemExit(str(exc)) from exc

    playlist_dir = channel_dir / "_playlists" / args.playlist_folder
    if not playlist_dir.is_dir():
        raise SystemExit(f"Playlist folder not found: {playlist_dir}")

    glossary_path = find_glossary(playlist_dir, channel_dir, args.terms)
    terms = load_terms(glossary_path)
    if glossary_path is not None:
        print(f"Glossary: {glossary_path} ({len(terms)} term(s)).", flush=True)

    slides_dir = playlist_dir / SLIDES_DIRNAME
    course = channel_dir.name.lstrip("_")
    orig_code = detect_original_code(playlist_dir)
    lang_codes = (
        [orig_code]
        if args.orig_only or orig_code == "EN"
        else [orig_code, "EN"]
    )

    # SLIDES/ folder keys are derived from the set of *local media* stems -
    # transcript-only videos (remote transcription) must not change them.
    media_videos = list_videos(playlist_dir)
    keys = short_slide_keys([video.stem for video in media_videos])
    videos = media_videos
    if args.no_slides:
        known = {video.stem for video in media_videos}
        extra = [
            playlist_dir / f"{srt.stem}.mp4"
            for srt in sorted((playlist_dir / orig_code).glob("*.srt"))
            if srt.stem not in known
        ]
        videos = sorted(media_videos + extra, key=lambda path: path.name.lower())
    if not videos:
        raise SystemExit(
            f"No video/audio files{' or transcripts' if args.no_slides else ''}"
            f" found in {playlist_dir}"
        )
    api_key = read_api_key(args.workspace)

    def slides_folder_for(video: Path) -> Path | None:
        if video.stem not in keys:
            return None
        return slides_out_dir(slides_dir, video, keys)

    def inputs_ready(video: Path) -> bool:
        if not all(
            transcripts_ready(playlist_dir / code, video.stem)
            for code in lang_codes
        ):
            return False
        folder = slides_folder_for(video)
        has_slides = folder is not None and (folder / RESULT_FILENAME).is_file()
        return has_slides or args.no_slides

    if args.video is not None:
        videos = [video for video in videos if video.stem == args.video]
        if not videos:
            raise SystemExit(
                f"--video: no video named {args.video!r} in {playlist_dir}"
            )

    eligible = [video for video in videos if inputs_ready(video)]
    pending = [
        video
        for video in eligible
        if args.force
        or not is_processed(
            playlist_dir,
            video.stem,
            lang_codes,
            annotate=args.annotate,
            orig_code=orig_code,
        )
    ]
    session = pending[: args.next_count] if args.video is None else pending

    print(f"Playlist folder: {playlist_dir}", flush=True)
    print(
        f"Original language: {orig_code}; outputs: "
        + ", ".join(f"{code}/{OUTPUT_DIRNAME}/" for code in lang_codes),
        flush=True,
    )
    ready_note = (
        "with transcripts" if args.no_slides else "with transcripts + slides.json"
    )
    print(
        f"Videos: {len(videos)} total, {len(eligible)} {ready_note}, "
        f"{len(eligible) - len(pending)} already processed, "
        f"{len(session)} in this session (--next {args.next_count}).",
        flush=True,
    )
    # Self-healing: the PDF stage is best-effort (Word may fail or be
    # blocked), so processed videos may lack PDFs. Rebuild them from the
    # existing .docx - this costs no API tokens.
    if args.pdf is None:
        for video in eligible:
            if video in session:
                continue
            for code in lang_codes:
                files = output_files(playlist_dir / code, video.stem)
                if files["docx"].is_file() and not files["pdf"].is_file():
                    print(
                        "PDF backfill: "
                        f"{files['pdf'].relative_to(playlist_dir)}...",
                        flush=True,
                    )
                    make_pdf(files["docx"], files["pdf"])

    if not session:
        print("Nothing to do: final documents exist for every ready video.",
              flush=True)
        return 0

    done = 0
    for index, video in enumerate(session, start=1):
        print(f"[{index}/{len(session)}] {video.name}", flush=True)
        if not process_video(
            video,
            playlist_dir=playlist_dir,
            slides_folder=slides_folder_for(video),
            api_key=api_key,
            course=course,
            orig_code=orig_code,
            lang_codes=lang_codes,
            terms=terms,
            doc_template=args.doc,
            pdf_template=args.pdf,
            auto_yes=args.yes,
            annotate=args.annotate,
        ):
            print("Session stopped: not confirmed.", flush=True)
            break
        done += 1

    print(f"Session done: {done} video(s).", flush=True)
    return 0


if __name__ == "__main__":
    if len(sys.argv) == 4 and sys.argv[1] == "--convert-pdf":
        # Internal entry point: one PDF conversion in a disposable child
        # process (see make_pdf).
        raise SystemExit(
            convert_docx_to_pdf(Path(sys.argv[2]), Path(sys.argv[3]))
        )
    raise SystemExit(main())
