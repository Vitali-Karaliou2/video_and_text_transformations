#!/usr/bin/env python3
"""Create final edited documents from transcripts and slides (pipeline step 7).

Inputs per video:
- transcripts in <playlist>/<ORIG>/ and <playlist>/EN/ (<stem>.txt/.srt),
  where ORIG is the original-language folder (e.g. RU); for English
  originals only EN/ is used;
- slide texts and document structure in <playlist>/SLIDES/<short-key>/
  slides.json (produced by text_from_slides.py); with --no-slides a video
  may instead have <playlist>/INFO/<stem>.json (saved by
  transcribe_videos.py in remote mode) - then the document title is built
  in the summary format (<channel> | <playlist> : <title>) and the
  sections come from the video timecodes (chapters).

Outputs: edited documents in <ORIG>/OUTPUT/ and EN/OUTPUT/ (--orig-only
limits this to <ORIG>/OUTPUT/), three formats each: <stem>.md, <stem>.docx
and <stem>.pdf.

How a document is built:

- The section list (with scene time ranges) comes from slides.json: the
  info slide opens the Foreword, foreword slides open unnumbered sections,
  the agenda slide opens the Agenda section, agenda-matched slides open the
  numbered sections, closing slides are merged into one Conclusion section.
- The .srt transcripts are sliced by those time ranges, so every section
  gets its original-language and English text fragments.
- Each section is edited by the OpenAI API using both fragments at once:
  recognition errors in the original are fixed against the English
  translation, punctuation is enriched (colons, dashes - not only commas),
  the text is split into paragraphs and lightly polished; awkward
  English slide wording used for headings is fixed.
- Subsections: when the section text discusses the slide bullets one by
  one in recognizable fragments (e.g. "Benefits of Test Automation"), one
  subsection per bullet is created; when bullets are only mentioned in
  passing and the text does not decompose (e.g. "What Is Test
  Automation?"), the section stays flat. The model decides per section.

Formats:
- .md - a Table of Contents with full multi-level numbering (1., 2., 2.1.);
  body lines are at most 80 characters, no hyphenation, justified to the
  right edge by turning some single spaces into double spaces.
- .docx - styles follow the --doc template (default: built-in defaults
  modeled on RU/Docs_Edited/01_Introduction.V2.docx: heading 2 for the
  title, heading 3/4 for sections/subsections, "ds-markdown-paragraph"
  for verbatim slide text). The Table of Contents is a real Word TOC
  field: entries are internal hyperlinks with right-aligned page numbers
  (regardless of the ToC source - slides or timecodes); the file asks
  Word/WPS to update fields on open.
- .pdf - the .docx rendered to PDF via WPS Writer when installed (no
  activation nags, no clashes with a user's Word session), otherwise via
  Microsoft Word; fields (the ToC) are updated before the export. --pdf
  may name a different .docx style template for the PDF rendering.

Costs and safety:
- before editing, the script prints the estimated gpt-4o cost for the video
  and waits for a y/n confirmation (--yes skips the prompt);
- editing results are cached in SLIDES/<key>/edited_sections.json after
  every section, so a re-run after a failure (or a formatting-only change)
  does not pay the API again; delete that file to force fresh editing;
- the PDF conversion runs in a disposable child process: a native Word COM
  crash cannot kill the run, and a Word instance left behind by a crashed
  attempt is killed automatically (the user's own Word is never touched).

Examples:
  python src/create_final_docs.py _Autotesting lectures
  python src/create_final_docs.py _Autotesting lectures --next 3
  python src/create_final_docs.py _Autotesting lectures --doc my_styles.docx
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import textwrap
import time
from dataclasses import dataclass, field
from pathlib import Path

_SRC_DIR = Path(__file__).resolve().parent
if str(_SRC_DIR) not in sys.path:
    sys.path.insert(0, str(_SRC_DIR))

from extract_slides import SLIDES_DIRNAME, short_slide_keys, slides_out_dir
from project_paths import WORKSPACE_ROOT, channels_dir, require_channel_ref
from text_from_slides import LANGUAGE_NAMES, MODEL, RESULT_FILENAME, chat_json
from transcribe_videos import (
    INFO_DIRNAME,
    Segment,
    group_paragraphs,
    list_videos,
    read_api_key,
)

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

OUTPUT_DIRNAME = "OUTPUT"
MD_WIDTH = 80
TOC_HEADING = "Table of Contents"
CONCLUSION_HEADINGS = {"EN": "Conclusion", "RU": "Заключение"}
# Heading of the single section when a video has neither slides nor timecodes.
TRANSCRIPT_HEADINGS = {"EN": "Transcript", "RU": "Транскрипт"}

# Editing results are cached next to slides.json, so a re-run after a
# failure (or after a formatting-only change) does not pay the API again.
EDIT_CACHE_FILENAME = "edited_sections.json"

# gpt-4o API prices as of 2026-07, USD per 1M tokens (see
# https://platform.openai.com/docs/pricing).
USD_PER_MTOKEN_PROMPT = 2.50
USD_PER_MTOKEN_COMPLETION = 10.00
# Rough tokenizer ratios for the cost estimate: ~4 chars per token for
# English, ~2.5 for other languages (Cyrillic packs fewer chars per token).
CHARS_PER_TOKEN_EN = 4.0
CHARS_PER_TOKEN_OTHER = 2.5
# Fixed per-section prompt overhead: system prompt + payload wrapper +
# slide title/bullets.
SECTION_OVERHEAD_TOKENS = 900

EDIT_SYSTEM_PROMPT = """\
You edit one section of an auto-transcribed lecture into its final form.
You get the section fragment of the original-language transcript, the
English machine translation of the same fragment, and the slide that opens
the section (title + bullet lines + scene timing).

Editing rules:
1. Original text: fix obvious speech recognition errors (the English
   translation often shows what was meant), especially garbled English
   product, company and technology names; enrich the punctuation according
   to the rules of that language (colons, dashes, semicolons - the raw
   transcript uses almost nothing but commas); split into paragraphs;
   lightly polish the spoken wording (drop filler words, false starts,
   broken repetitions) while fully preserving the content, details, examples
   and the speaker's tone. Never summarize, never drop content, never add
   content of your own. When the original language is Russian, use the
   letter «ё» wherever standard Russian orthography calls for it (e.g.
   «ещё», «всё», «идёт», «счёт», «отчёт»); do not replace «ё» with «е».
2. English text: edit the same way; its punctuation usually needs little
   work, but polish awkward constructions - the speaker is not a native
   English speaker.
3. Keep the two languages parallel: the same paragraph breaks, the same
   subsection boundaries, the same order.
4. Section layout - decide between exactly two layouts:
   (a) flat: no subsections, all text in the intro paragraphs;
   (b) fully decomposed: a short intro (possibly empty) followed by one
       subsection per slide bullet - every bullet, in bullet order -
       together covering all the remaining text. Digressions stay inside
       the subsection of the bullet being discussed; the last subsection
       extends to the end of the section. Each subsection gets a concise
       English heading derived from its bullet plus a translation into the
       original language.
   Choose (b) only when the bullets announce a list of parallel items
   (features, types, practices, tools) that the speaker then demonstrably
   covers one after another, so that the fragment of every bullet can be
   located confidently. Explanatory slides whose lines merely characterize
   the topic (a "Definition: ... Purpose: ... Example: ..." pattern and the
   like) usually do not structure the speech that follows - such sections
   stay flat even when the opening sentences touch the lines in order. If
   some bullets get no dedicated discussion, or the boundaries are unclear,
   or there are no bullets, choose (a). A partial split (subsections for
   only some bullets) is not allowed. When in doubt, choose (a).
5. Propose the final section heading in English: keep the given heading if
   it is fine, otherwise fix grammar or awkward wording; keep it short. Use
   normal title casing, not ALL CAPS.

Reply with JSON only:
{
  "heading": string,
  "intro_orig": [paragraphs in the original language],
  "intro_en": [paragraphs in English],
  "subsections": [
    {"heading_en": string, "heading_orig": string,
     "paragraphs_orig": [strings], "paragraphs_en": [strings]}
  ]
}
"subsections" may be []. When the original language is English, fill only
the *_en fields and leave "intro_orig"/"paragraphs_orig"/"heading_orig" as
empty lists / empty strings. When "transcript_english" is empty, do NOT
translate: edit only the original language and leave "intro_en" and every
"paragraphs_en" as empty lists.

When "section_part" is present ("i/N"), you are editing part i of N of one
long section split only for processing: edit the COMPLETE given fragment
from its first sentence to its last (it may start or end mid-topic - that
is expected; do not add openings or conclusions of your own), keep the
layout flat ("subsections": []) and remember: re-emit ALL the content,
paragraph by paragraph - never compress the fragment."""


ANNOTATION_CACHE_KEY = "__annotation__"

ANNOTATION_SYSTEM_PROMPT = """\
You write a concise annotation (abstract) of an edited video transcript.

Requirements:
- STRICT length: 200-250 words per language. Never fewer than 200 words.
  Count the words before replying; expand with concrete details from the
  transcript if the draft is short;
- cover the topic, the key points and the conclusions of the video;
- neutral informative tone, plain prose in 1-3 paragraphs, no bullet lists;
- write each annotation natively in its own language (translate the
  content, do not transliterate);
- for Russian text, use the letter «ё» wherever standard Russian
  orthography calls for it; do not replace «ё» with «e».

Reply with JSON only:
{"annotations": {"XX": "text", ...}}
with exactly the two-letter upper-case language codes requested in
"languages"."""


def annotation_languages(orig_code: str) -> list[str]:
    """Annotation languages by the rules:
    Russian original -> Russian annotation only; any other original ->
    English annotation plus a Russian translation."""
    return ["RU"] if orig_code == "RU" else ["EN", "RU"]


@dataclass
class Section:
    heading: str
    number: str | None  # docx numbering (from the agenda); None = unnumbered
    slide: dict | None
    start: float
    end: float = 0.0
    fixed_headings: dict[str, str] | None = None  # e.g. Conclusion per lang
    heading_en: str = ""
    intro: dict[str, list[str]] = field(default_factory=dict)
    subsections: list[dict] = field(default_factory=list)


# --------------------------------------------------------------------------
# Inputs


def parse_srt(path: Path) -> list[Segment]:
    def to_seconds(stamp: str) -> float:
        hours, minutes, rest = stamp.strip().split(":")
        seconds, millis = rest.split(",")
        return (
            int(hours) * 3600 + int(minutes) * 60 + int(seconds)
            + int(millis) / 1000.0
        )

    segments: list[Segment] = []
    for block in re.split(r"\n\s*\n", path.read_text(encoding="utf-8-sig")):
        lines = [line for line in block.splitlines() if line.strip()]
        if len(lines) < 2 or "-->" not in lines[1]:
            continue
        start, end = lines[1].split("-->")
        segments.append(
            Segment(
                start=to_seconds(start),
                end=to_seconds(end),
                text=" ".join(line.strip() for line in lines[2:]),
            )
        )
    return segments


def section_text(segments: list[Segment], start: float, end: float) -> str:
    part = [seg for seg in segments if start <= seg.start < end]
    return "\n\n".join(group_paragraphs(part))


def detect_original_code(playlist_dir: Path) -> str:
    """Original-language folder code (RU, DE, ...); EN when only EN exists."""
    codes = sorted(
        entry.name
        for entry in playlist_dir.iterdir()
        if entry.is_dir()
        and len(entry.name) == 2
        and entry.name.isalpha()
        and entry.name.isupper()
        and entry.name != "EN"
        and any(entry.glob("*.srt"))
    )
    return codes[0] if codes else "EN"


# --------------------------------------------------------------------------
# Section list from slides.json


def build_sections(slides_data: dict, video_end: float) -> list[Section]:
    toc_by_slide = {
        entry.get("slide"): entry
        for entry in (slides_data.get("document") or {}).get("toc") or []
        if entry.get("slide")
    }
    sections: list[Section] = []
    for slide in slides_data.get("slides") or []:
        role = slide.get("role")
        scene = slide.get("scene") or {}
        start = float(scene.get("start_seconds") or 0.0)
        entry = toc_by_slide.get(slide.get("file")) or {}
        if role == "lecture_info":
            # The info slide only feeds document metadata; its lines are not
            # repeated inside the Foreword section.
            sections.append(Section("Foreword", None, None, start))
        elif role == "foreword":
            heading = entry.get("title") or slide.get("title") or "Foreword"
            sections.append(Section(heading, None, slide, start))
        elif role == "agenda":
            sections.append(Section("Agenda", None, slide, start))
        elif role == "section":
            heading = entry.get("title") or slide.get("title") or ""
            sections.append(Section(heading, entry.get("number"), slide, start))
        elif role == "closing" and (
            not sections or sections[-1].fixed_headings is None
        ):
            sections.append(
                Section(
                    "Conclusion",
                    None,
                    slide,
                    start,
                    fixed_headings=CONCLUSION_HEADINGS,
                )
            )
    if not sections:
        sections.append(Section("Transcript", None, None, 0.0))
    close_section_ranges(sections, video_end)
    return sections


def close_section_ranges(sections: list[Section], video_end: float) -> None:
    sections[0].start = 0.0
    for current, following in zip(sections, sections[1:]):
        current.end = following.start
    sections[-1].end = max(video_end, sections[-1].start + 1.0)


# --------------------------------------------------------------------------
# Section list and document header without slides (INFO/<stem>.json written
# by transcribe_videos.py in remote mode: title, date, duration, timecodes)


def load_video_info(playlist_dir: Path, stem: str) -> dict:
    path = playlist_dir / INFO_DIRNAME / f"{stem}.json"
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return {}


def info_document_header(info: dict, stem: str) -> tuple[str, list[str]]:
    """Document title (summary format: <channel> | <playlist> : <title>)
    and meta lines from the saved video metadata."""
    title = str(info.get("title") or stem)
    channel = str(info.get("channel") or "")
    playlist = str(info.get("playlist") or "")
    if channel and playlist:
        doc_title = f"{channel} | {playlist} : {title}"
    elif channel:
        doc_title = f"{channel} : {title}"
    else:
        doc_title = title
    meta_lines = [
        f"{label}: {value}"
        for label, value in (
            ("Date", info.get("date")),
            ("Duration", f"[{info['duration_text']}]"
             if info.get("duration_text") else None),
            ("URL", info.get("url")),
        )
        if value
    ]
    return doc_title, meta_lines


def build_sections_from_info(
    info: dict, video_end: float, orig_code: str
) -> list[Section]:
    """One numbered section per video chapter (timecode); a single
    'Transcript' section when the video has no timecodes. Chapter titles are
    the author's own headings, so they are used verbatim in every language."""
    sections: list[Section] = []
    for index, chapter in enumerate(info.get("chapters") or [], start=1):
        title = str(chapter.get("title") or f"Part {index}").strip()
        sections.append(
            Section(
                title,
                str(index),
                None,
                float(chapter.get("start") or 0.0),
                fixed_headings={orig_code: title, "EN": title},
            )
        )
    if not sections:
        sections.append(
            Section(
                "Transcript", None, None, 0.0,
                fixed_headings=TRANSCRIPT_HEADINGS,
            )
        )
    close_section_ranges(sections, video_end)
    return sections


# --------------------------------------------------------------------------
# Editing via the OpenAI API


def edit_section(
    api_key: str,
    section: Section,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    orig_text: str,
    en_text: str,
    usage: dict[str, int],
    part: tuple[int, int] | None = None,
) -> None:
    slide = section.slide or {}
    orig_name = LANGUAGE_NAMES.get(orig_code, orig_code)
    payload = {
        "course": course,
        "document_title": doc_title,
        "original_language": orig_name,
        "section_heading": section.heading,
        "slide_title": slide.get("title"),
        "slide_bullets": slide.get("body") or [],
        "heading_flags": slide.get("flags") or [],
        "section_minutes": round((section.end - section.start) / 60.0, 1),
        "transcript_original": orig_text if orig_code != "EN" else "",
        "transcript_english": en_text,
    }
    if part is not None:
        payload["section_part"] = f"{part[0]}/{part[1]}"
    result = chat_json(
        api_key,
        [
            {"role": "system", "content": EDIT_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False, indent=2),
            },
        ],
        usage,
    )
    section.heading_en = str(result.get("heading") or section.heading)
    section.intro = {
        orig_code: [str(p) for p in result.get("intro_orig") or []],
        "EN": [str(p) for p in result.get("intro_en") or []],
    }
    section.subsections = []
    for sub in result.get("subsections") or []:
        section.subsections.append(
            {
                "heading_en": str(sub.get("heading_en") or ""),
                "heading_orig": str(sub.get("heading_orig") or ""),
                orig_code: [str(p) for p in sub.get("paragraphs_orig") or []],
                "EN": [str(p) for p in sub.get("paragraphs_en") or []],
            }
        )
    slide = section.slide or {}
    if section.subsections and (
        slide.get("role") == "agenda"
        # Agenda items are the document's ToC: per-item subsections would
        # duplicate the main sections.
        or explanatory_bullets(slide.get("body") or [])
        # Explanatory slides (Definition:/Purpose:/Example: lines) do not
        # structure the speech that follows (documented limit case 1);
        # the model still splits them sometimes, so this is enforced here.
    ):
        flatten_section(section)


def flatten_section(section: Section) -> None:
    """Fold the subsections back into the intro paragraphs, in order."""
    for sub in section.subsections:
        for code in section.intro:
            section.intro[code].extend(sub.get(code) or [])
    section.subsections = []


# gpt-4o silently compresses long re-emissions: asked to edit a whole
# 20-minute transcript in one call it returns a couple of paragraphs. Long
# sections (typical for timecode-less videos, where the whole video is one
# section) are therefore edited in ~4-minute chunks and concatenated.
EDIT_CHUNK_SECONDS = 240.0
EDIT_CHUNK_TRIGGER_SECONDS = 360.0


def split_section_ranges(
    section: Section,
    segments: list[Segment],
    chunk_seconds: float = EDIT_CHUNK_SECONDS,
) -> list[tuple[float, float]]:
    """Split [start, end) into roughly equal windows of ~chunk_seconds,
    snapping each boundary to the nearest transcript segment boundary so
    that no sentence is cut in half (both languages are sliced by the same
    time boundaries, staying parallel)."""
    duration = section.end - section.start
    parts = max(1, round(duration / chunk_seconds))
    if parts == 1:
        return [(section.start, section.end)]
    candidates = [
        seg.start
        for seg in segments
        if section.start < seg.start < section.end
    ]
    boundaries: list[float] = []
    for index in range(1, parts):
        target = section.start + duration * index / parts
        snapped = (
            min(candidates, key=lambda t: abs(t - target))
            if candidates
            else target
        )
        if boundaries and snapped <= boundaries[-1]:
            continue
        boundaries.append(snapped)
    edges = [section.start, *boundaries, section.end]
    return [
        (edges[i], edges[i + 1])
        for i in range(len(edges) - 1)
        if edges[i + 1] > edges[i]
    ]


def edit_section_in_chunks(
    api_key: str,
    section: Section,
    *,
    course: str,
    doc_title: str,
    orig_code: str,
    segments: dict[str, list[Segment]],
    usage: dict[str, int],
) -> None:
    """Edit one long section chunk by chunk; the result is always flat."""
    ranges = split_section_ranges(section, segments.get(orig_code) or [])
    intro: dict[str, list[str]] = {orig_code: [], "EN": []}
    heading_en = ""
    for index, (start, end) in enumerate(ranges, start=1):
        print(
            f"    part {index}/{len(ranges)} "
            f"({(end - start) / 60.0:.1f} min)...",
            flush=True,
        )
        part = Section(
            heading=section.heading,
            number=section.number,
            slide=section.slide,
            start=start,
            end=end,
        )
        orig_text = (
            section_text(segments[orig_code], start, end)
            if orig_code != "EN"
            else ""
        )
        en_text = (
            section_text(segments["EN"], start, end)
            if "EN" in segments
            else ""
        )
        edit_section(
            api_key,
            part,
            course=course,
            doc_title=doc_title,
            orig_code=orig_code,
            orig_text=orig_text,
            en_text=en_text,
            usage=usage,
            part=(index, len(ranges)),
        )
        flatten_section(part)
        heading_en = heading_en or part.heading_en
        for code, paragraphs in part.intro.items():
            intro.setdefault(code, []).extend(paragraphs)
    section.heading_en = heading_en or section.heading
    section.intro = intro
    section.subsections = []


EXPLANATORY_LABELS = {
    "definition", "purpose", "example", "goal", "examples", "note",
}


def explanatory_bullets(bullets: list[str]) -> bool:
    labeled = [
        bullet.split(":", 1)[0].strip().lower()
        for bullet in bullets
        if ":" in bullet
    ]
    known = [label for label in labeled if label in EXPLANATORY_LABELS]
    return len(known) >= 2 and 2 * len(known) >= len(bullets)


def section_heading_for(section: Section, lang: str) -> str:
    if section.fixed_headings and lang in section.fixed_headings:
        return section.fixed_headings[lang]
    return section.heading_en or section.heading


def subsection_heading_for(sub: dict, lang: str) -> str:
    if lang != "EN" and sub.get("heading_orig"):
        return sub["heading_orig"]
    return sub.get("heading_en") or ""


# --------------------------------------------------------------------------
# Annotation (200-250 words) generation and output


# Keep the annotation prompt well under the org TPM limit (30k for gpt-4o).
ANNOTATION_MAX_PROMPT_TOKENS = 18000


def edited_full_text(sections: list[Section], code: str) -> str:
    """All edited paragraphs of the given language, in document order."""
    parts: list[str] = []
    for section in sections:
        parts.extend(section.intro.get(code) or [])
        for sub in section.subsections:
            parts.extend(sub.get(code) or [])
    return "\n\n".join(parts)


def annotation_source_text(sections: list[Section], code: str) -> str:
    """Edited text condensed for the annotation prompt: when the full text
    would blow the request token limit, keep the section headings and a
    proportional beginning of every section, so the whole video is still
    covered."""
    ratio = CHARS_PER_TOKEN_EN if code == "EN" else CHARS_PER_TOKEN_OTHER
    budget = int(ANNOTATION_MAX_PROMPT_TOKENS * ratio)

    blocks: list[tuple[str, list[str]]] = []
    for section in sections:
        paragraphs: list[str] = list(section.intro.get(code) or [])
        for sub in section.subsections:
            paragraphs.extend(sub.get(code) or [])
        blocks.append((section.heading_en or section.heading, paragraphs))

    total = sum(len(p) for _, paras in blocks for p in paras)
    if total <= budget:
        return "\n\n".join(
            "\n\n".join([f"## {heading}", *paras])
            for heading, paras in blocks
            if paras or heading
        )

    scale = budget / total
    parts: list[str] = []
    for heading, paras in blocks:
        parts.append(f"## {heading}")
        section_budget = int(sum(len(p) for p in paras) * scale)
        used = 0
        for paragraph in paras:
            if used and used + len(paragraph) > section_budget:
                break
            parts.append(paragraph)
            used += len(paragraph)
    return "\n\n".join(parts)


def generate_annotation(
    api_key: str,
    *,
    doc_title: str,
    orig_code: str,
    sections: list[Section],
    usage: dict[str, int],
) -> dict[str, str]:
    languages = annotation_languages(orig_code)
    source_code = orig_code if orig_code != "EN" else "EN"
    text = annotation_source_text(sections, source_code)
    payload = {
        "document_title": doc_title,
        "original_language": LANGUAGE_NAMES.get(orig_code, orig_code),
        "languages": languages,
        "edited_transcript": text,
    }
    result = chat_json(
        api_key,
        [
            {"role": "system", "content": ANNOTATION_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False, indent=2),
            },
        ],
        usage,
    )
    annotations = {
        code: str((result.get("annotations") or {}).get(code) or "").strip()
        for code in languages
    }
    # The model tends to undershoot the word count; one targeted retry per
    # too-short annotation.
    for code in languages:
        current = annotations.get(code)
        if not current or len(current.split()) >= 200:
            continue
        retry = chat_json(
            api_key,
            [
                {
                    "role": "system",
                    "content": (
                        "Expand the given annotation to 200-250 words "
                        "(STRICT: never fewer than 200 words - count them), "
                        "in the same language, adding concrete details "
                        "from the transcript. Keep the neutral informative "
                        "tone and plain prose. Reply with JSON only: "
                        '{"annotation": "text"}'
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "annotation": current,
                            "edited_transcript": text,
                        },
                        ensure_ascii=False,
                    ),
                },
            ],
            usage,
        )
        expanded = str(retry.get("annotation") or "").strip()
        if len(expanded.split()) > len(current.split()):
            annotations[code] = expanded
    return {code: text_ for code, text_ in annotations.items() if text_}


def annotation_path(playlist_dir: Path, code: str, stem: str) -> Path:
    return playlist_dir / code / OUTPUT_DIRNAME / f"{stem}.txt"


def write_annotation(path: Path, doc_title: str, text: str) -> None:
    """Annotation .txt: the document title, then the text formatted like the
    .md body (justified 80-character lines, no hyphenation, blank lines
    between paragraphs)."""
    lines: list[str] = textwrap.wrap(
        doc_title, width=MD_WIDTH, break_long_words=False,
        break_on_hyphens=False,
    )
    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = " ".join(paragraph.split())
        if paragraph:
            lines += ["", justify_paragraph(paragraph)]
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


# --------------------------------------------------------------------------
# Markdown output


def justify_line(line: str, width: int) -> str:
    """Pad the line to the width by doubling some single spaces (never more
    than doubling); spaces to double are spread evenly along the line."""
    deficit = width - len(line)
    words = line.split(" ")
    gaps = len(words) - 1
    if deficit <= 0 or gaps == 0 or deficit > gaps:
        return line
    chosen: set[int] = set()
    for index in range(deficit):
        # Center the doubled spaces inside their stretch of the line, so the
        # very first gaps are not systematically the widened ones.
        gap = min(gaps - 1, int((index + 0.5) * gaps / deficit))
        while gap in chosen:  # deficit <= gaps, so a free gap always exists
            gap = (gap + 1) % gaps
        chosen.add(gap)
    pieces: list[str] = []
    for index, word in enumerate(words[:-1]):
        pieces.append(word)
        pieces.append("  " if index in chosen else " ")
    pieces.append(words[-1])
    return "".join(pieces)


def justify_paragraph(text: str, width: int = MD_WIDTH) -> str:
    lines = textwrap.wrap(
        " ".join(text.split()),
        width=width,
        break_long_words=False,
        break_on_hyphens=False,
    )
    return "\n".join(
        [justify_line(line, width) for line in lines[:-1]] + lines[-1:]
    ) if lines else ""


def wrap_blockquote(text: str, width: int = MD_WIDTH) -> list[str]:
    lines = textwrap.wrap(
        text, width=width - 2, break_long_words=False, break_on_hyphens=False
    )
    return [f"> {line}" for line in lines] or ["> "]


def build_markdown(
    doc_title: str,
    meta_lines: list[str],
    sections: list[Section],
    lang: str,
) -> str:
    out: list[str] = [f"# {doc_title}", ""]
    for line in meta_lines:
        out.append(line + "  ")
    out += ["", f"## {TOC_HEADING}", ""]
    for number, section in enumerate(sections, start=1):
        out.append(f"- {number}. {section_heading_for(section, lang)}")
        for sub_number, sub in enumerate(section.subsections, start=1):
            out.append(
                f"  - {number}.{sub_number}. "
                f"{subsection_heading_for(sub, lang)}"
            )
    out.append("")
    for number, section in enumerate(sections, start=1):
        out += [f"## {number}. {section_heading_for(section, lang)}", ""]
        slide = section.slide or {}
        for bullet in slide.get("body") or []:
            out += wrap_blockquote(bullet)
        if slide.get("body"):
            out.append("")
        for paragraph in section.intro.get(lang) or []:
            out += [justify_paragraph(paragraph), ""]
        for sub_number, sub in enumerate(section.subsections, start=1):
            out += [
                f"### {number}.{sub_number}. "
                f"{subsection_heading_for(sub, lang)}",
                "",
            ]
            for paragraph in sub.get(lang) or []:
                out += [justify_paragraph(paragraph), ""]
    while out and out[-1] == "":
        out.pop()
    return "\n".join(out) + "\n"


# --------------------------------------------------------------------------
# DOCX / PDF output


def make_docx(
    doc_title: str,
    meta_lines: list[str],
    sections: list[Section],
    lang: str,
    template: Path | None,
    out_path: Path,
) -> None:
    from docx import Document
    from docx.shared import Pt

    if template is not None:
        document = Document(str(template))
        body = document.element.body
        for child in list(body):
            if not child.tag.endswith("}sectPr"):
                body.remove(child)
    else:
        document = Document()

    style_names = {style.name for style in document.styles}

    def pick(*candidates: str) -> str | None:
        for name in candidates:
            if name in style_names:
                return name
        return None

    # Style mapping modeled on RU/Docs_Edited/01_Introduction.V2.docx:
    # document title = heading 2, sections = heading 3, subsections =
    # heading 4, verbatim slide text = ds-markdown-paragraph.
    title_style = pick("Heading 2", "Title", "Heading 1")
    section_style = pick("Heading 3", "Heading 1")
    subsection_style = pick("Heading 4", "Heading 2")
    slide_style = pick("ds-markdown-paragraph")

    if template is None:
        apply_default_formatting(document)

    def para(text: str, style: str | None = None, *, italic: bool = False):
        paragraph = document.add_paragraph()
        if style:
            paragraph.style = document.styles[style]
        run = paragraph.add_run(text)
        if italic:
            run.italic = True
        return paragraph

    para(doc_title, title_style)
    for line in meta_lines:
        para(line)
    # The ToC caption must not use a heading style, or the TOC field below
    # would list the caption itself as an entry.
    caption = para(TOC_HEADING)
    caption_run = caption.runs[0]
    caption_run.bold = True
    caption_size = (
        document.styles[section_style].font.size if section_style else None
    )
    caption_run.font.size = caption_size or Pt(13)
    add_toc_field(
        document,
        heading_level(section_style, 1),
        heading_level(subsection_style, 2),
    )
    request_field_update_on_open(document)
    for section in sections:
        number = f"{section.number}. " if section.number else ""
        para(f"{number}{section_heading_for(section, lang)}", section_style)
        slide = section.slide or {}
        for bullet in slide.get("body") or []:
            para(bullet, slide_style, italic=slide_style is None)
        for paragraph in section.intro.get(lang) or []:
            para(paragraph)
        for index, sub in enumerate(section.subsections, start=1):
            sub_number = (
                f"{section.number}.{index}. " if section.number else ""
            )
            para(
                f"{sub_number}{subsection_heading_for(sub, lang)}",
                subsection_style,
            )
            for paragraph in sub.get(lang) or []:
                para(paragraph)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def heading_level(style_name: str | None, default: int) -> int:
    """Outline level of a built-in "Heading N" style name."""
    match = re.search(r"(\d+)$", style_name or "")
    return int(match.group(1)) if match else default


def add_toc_field(document, top_level: int, bottom_level: int) -> None:
    """Insert a real Word TOC field covering the given heading levels.

    Word/WPS builds it into a table of contents with internal hyperlinks and
    right-aligned page numbers when the fields are updated - the .docx asks
    for that on open (see request_field_update_on_open), and the PDF
    conversion updates the fields before exporting.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' TOC \\o "{top_level}-{bottom_level}" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    for element in (begin, instruction, separate):
        field_run._r.append(element)
    placeholder = paragraph.add_run(
        "Update the fields (Ctrl+A, F9) to build the table of contents."
    )
    placeholder.italic = True
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def request_field_update_on_open(document) -> None:
    """Ask Word/WPS to recalculate fields (the ToC) when the file is opened."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def apply_default_formatting(document) -> None:
    """Default look of the generated .docx (used when --doc is not given).

    - the page number is shown centered in the page header;
    - non-heading paragraph styles get Calibri, 6 pt spacing above and
      below and justified alignment; heading styles keep their spacing.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        name = style.name.lower()
        if name.startswith("heading") or name == "title":
            continue
        style.font.name = "Calibri"
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header = document.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


RPC_E_CALL_REJECTED = -2147418111  # "Call was rejected by callee": Word busy


def com_retry(call, attempts: int = 8, on_reject=None):
    """Run a COM call, waiting out transient 'call rejected' rejections
    (Word keeps servicing its message loop for a while after startup);
    `on_reject` runs before each retry (e.g. to dismiss a nag dialog)."""
    for attempt in range(1, attempts + 1):
        try:
            return call()
        except Exception as exc:
            if (
                getattr(exc, "hresult", None) != RPC_E_CALL_REJECTED
                or attempt == attempts
            ):
                raise
            if on_reject is not None:
                try:
                    on_reject()
                except Exception:
                    pass
            time.sleep(attempt)


def dismiss_word_dialogs(pids: set[int]) -> None:
    """Close nag dialogs of the given (hidden) Word processes.

    Word occasionally pops modal prompts even in automation - e.g. "Word
    is not your default program for documents, choose file types?" - and
    then rejects every COM call until somebody answers. Press the safest
    button (No/Cancel) or close the dialog.
    """
    import win32con
    import win32gui
    import win32process

    BM_CLICK = 0x00F5
    PREFERRED = ("нет", "no", "отмена", "cancel", "закрыть", "close")

    def handle_dialog(hwnd) -> None:
        buttons: list[tuple[int, str]] = []

        def collect(child, _):
            if win32gui.GetClassName(child) == "Button":
                text = (
                    win32gui.GetWindowText(child)
                    .replace("&", "").strip().lower()
                )
                buttons.append((child, text))
            return True

        try:
            win32gui.EnumChildWindows(hwnd, collect, None)
        except Exception:
            pass
        for preferred in PREFERRED:
            for handle, text in buttons:
                if preferred in text:
                    win32gui.PostMessage(handle, BM_CLICK, 0, 0)
                    return
        win32gui.PostMessage(hwnd, win32con.WM_CLOSE, 0, 0)

    def visit(hwnd, _):
        if (
            win32gui.IsWindowVisible(hwnd)
            and win32gui.GetClassName(hwnd) == "#32770"
            and win32process.GetWindowThreadProcessId(hwnd)[1] in pids
        ):
            handle_dialog(hwnd)
        return True

    win32gui.EnumWindows(visit, None)


def wps_progid() -> str | None:
    """COM ProgID of WPS Writer when WPS Office is installed, else None."""
    import winreg

    for progid in ("Kwps.Application", "wps.Application"):
        try:
            winreg.CloseKey(
                winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, progid)
            )
            return progid
        except OSError:
            continue
    return None


def run_tasklist(args: list[str]) -> str:
    """tasklist output, decoded safely.

    tasklist writes in the console OEM code page; the default locale codec
    (e.g. cp1252) chokes on Cyrillic window titles or localized messages,
    the decode error kills subprocess' reader thread and stdout comes back
    as None - so decode as OEM with replacement.
    """
    import subprocess

    try:
        return subprocess.run(
            ["tasklist", *args],
            capture_output=True,
            text=True,
            encoding="oem",
            errors="replace",
            timeout=30,
        ).stdout or ""
    except Exception:
        return ""


def image_pids(image_name: str) -> set[int]:
    """PIDs of the running processes with the given executable name."""
    listing = run_tasklist(
        ["/FI", f"IMAGENAME eq {image_name}", "/FO", "CSV", "/NH"]
    )
    pids: set[int] = set()
    for line in listing.splitlines():
        cells = [cell.strip('"') for cell in line.strip().split('","')]
        if len(cells) >= 2 and cells[0].lower() == image_name.lower():
            try:
                pids.add(int(cells[1]))
            except ValueError:
                pass
    return pids


def update_com_fields(document) -> None:
    """Rebuild the tables of contents and other fields of an open document
    (Word or WPS COM), so the exported PDF gets ToC page numbers and
    hyperlinks. Best-effort: a document without fields is left as is."""
    try:
        tocs = document.TablesOfContents
        for index in range(1, int(tocs.Count) + 1):
            tocs(index).Update()
    except Exception:
        pass
    try:
        document.Fields.Update()
    except Exception:
        pass


def convert_docx_to_pdf_wps(docx_path: Path, pdf_path: Path,
                            progid: str) -> int:
    """One PDF conversion attempt via WPS Writer COM (Word-compatible API).

    WPS shows no activation nags and does not conflict with a Microsoft
    Word window the user is working in, so it is preferred over Word when
    installed. The application is Quit() only when this call actually
    started the wps.exe process - if a running WPS instance was reused,
    quitting would close the user's own documents.
    """
    WD_FORMAT_PDF = 17
    import pythoncom
    import win32com.client

    before = image_pids("wps.exe")
    pythoncom.CoInitialize()
    try:
        app = win32com.client.DispatchEx(progid)
        started = image_pids("wps.exe") - before
        try:
            app.Visible = False
            try:
                app.DisplayAlerts = 0
            except Exception:
                pass
            document = app.Documents.Open(
                str(docx_path), ReadOnly=False, AddToRecentFiles=False
            )
            try:
                update_com_fields(document)
                document.ExportAsFixedFormat(str(pdf_path), WD_FORMAT_PDF)
            finally:
                document.Close(False)
        finally:
            if started:
                try:
                    app.Quit()
                except Exception:
                    pass
    finally:
        pythoncom.CoUninitialize()
    return 0 if pdf_path.is_file() else 1


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> int:
    """One PDF conversion attempt; runs in a child process.

    WPS Writer is used when installed (no activation nags, no clashes with
    a user's Word session); otherwise - or if WPS fails - Microsoft Word
    COM is the fallback.
    """
    # COM servers resolve relative paths against their own cwd, not ours.
    docx_path, pdf_path = docx_path.resolve(), pdf_path.resolve()
    progid = wps_progid()
    if progid is not None:
        try:
            if convert_docx_to_pdf_wps(docx_path, pdf_path, progid) == 0:
                return 0
        except Exception as exc:
            print(f"WPS conversion failed ({exc}); trying Word...",
                  file=sys.stderr, flush=True)
    return convert_docx_to_pdf_word(docx_path, pdf_path)


def convert_docx_to_pdf_word(docx_path: Path, pdf_path: Path) -> int:
    """One PDF conversion attempt via Word COM; runs in a child process.

    A dedicated Word instance (DispatchEx) is created and always Quit() in
    the end, so no zombie WINWORD process stays behind holding file locks,
    and a Word instance the user works in is never used (docx2pdf reuses a
    shared instance and is prone to both problems).

    Early binding (a generated type-library wrapper) is required: dynamic
    dispatch proved unreliable for Word Document objects on this setup -
    GetIDsOfNames fails with "AttributeError: Open.SaveAs". A corrupted
    gen_py cache (e.g. half-written by a crashed run) is dropped so that
    the next attempt regenerates it.
    """
    WD_FORMAT_PDF = 17
    import pythoncom
    import win32com.client

    drop_word_resiliency_entries(WORKSPACE_ROOT)
    before = winword_pids()
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        mine = winword_pids() - before  # the instance just started
        unblock = lambda: dismiss_word_dialogs(mine)  # noqa: E731
        try:
            word.Visible = False
            word.DisplayAlerts = 0
            document = com_retry(
                lambda: word.Documents.Open(
                    str(docx_path), ReadOnly=False, AddToRecentFiles=False
                ),
                on_reject=unblock,
            )
            try:
                com_retry(
                    lambda: update_com_fields(document), on_reject=unblock
                )
                com_retry(
                    lambda: document.SaveAs(
                        str(pdf_path), FileFormat=WD_FORMAT_PDF
                    ),
                    on_reject=unblock,
                )
            except AttributeError:
                # Broken name resolution: drop the (likely corrupted)
                # gen_py cache; the retry in the parent starts a fresh
                # child that rebuilds it.
                clear_gen_py_cache()
                raise
            com_retry(lambda: document.Close(False), on_reject=unblock)
        finally:
            try:
                com_retry(lambda: word.Quit(), attempts=3, on_reject=unblock)
            except Exception:
                pass
            # A Word that refused to Quit would linger, hold file locks and
            # make every following conversion fail too ("call rejected"
            # avalanche) - kill this child's own instance if still alive.
            kill_winword(mine)
    finally:
        pythoncom.CoUninitialize()
    return 0 if pdf_path.is_file() else 1


def drop_word_resiliency_entries(scope: Path) -> None:
    """Remove Word's crash-bookkeeping registry entries for our documents.

    When a WINWORD process dies while a document is open, Word records the
    document under HKCU ...\\Word\\Resiliency (DisabledItems /
    DocumentRecovery). On the next automated start Word then waits forever
    on a hidden prompt ("open the problem document?" / recovery / safe
    mode) and every COM call fails with 'Call was rejected by callee'.
    Entries for any document inside `scope` (the workspace) are removed;
    the user's own documents elsewhere are not touched.
    """
    import winreg

    target = str(scope).lower()

    def mentions_target(data) -> bool:
        return isinstance(data, bytes) and target in data.decode(
            "utf-16-le", errors="ignore"
        ).lower()

    def matching_values(key) -> list[str]:
        names = []
        index = 0
        while True:
            try:
                name, data, _ = winreg.EnumValue(key, index)
            except OSError:
                return names
            if mentions_target(data):
                names.append(name)
            index += 1

    def data_mentions_target(key) -> bool:
        return bool(matching_values(key))

    def subkeys(key) -> list[str]:
        names = []
        index = 0
        while True:
            try:
                names.append(winreg.EnumKey(key, index))
            except OSError:
                return names
            index += 1

    try:
        with winreg.OpenKey(
            winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Office"
        ) as office:
            versions = subkeys(office)
    except OSError:
        return
    for version in versions:
        base = (
            rf"Software\Microsoft\Office\{version}\Word\Resiliency"
        )
        # DisabledItems: flat values, the file path inside the binary data.
        try:
            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                base + r"\DisabledItems",
                0,
                winreg.KEY_ALL_ACCESS,
            ) as key:
                for name in matching_values(key):
                    winreg.DeleteValue(key, name)
        except OSError:
            pass
        # DocumentRecovery: one subkey per remembered document.
        try:
            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                base + r"\DocumentRecovery",
                0,
                winreg.KEY_ALL_ACCESS,
            ) as key:
                for name in subkeys(key):
                    with winreg.OpenKey(key, name) as sub:
                        drop = data_mentions_target(sub)
                    if drop:
                        winreg.DeleteKey(key, name)
        except OSError:
            pass
        # StartupItems: "Word crashed while starting" markers; they make a
        # hidden Word ask "start in safe mode?" and reject all COM calls.
        try:
            winreg.DeleteKey(
                winreg.HKEY_CURRENT_USER, base + r"\StartupItems"
            )
        except OSError:
            pass


def clear_gen_py_cache() -> None:
    import shutil

    try:
        import win32com

        shutil.rmtree(win32com.__gen_path__, ignore_errors=True)
    except Exception:
        pass


def winword_processes() -> dict[int, str]:
    """PID -> window title for the running WINWORD processes ("N/A" for
    hidden automation instances without a window)."""
    listing = run_tasklist(
        ["/V", "/FI", "IMAGENAME eq WINWORD.EXE", "/FO", "CSV", "/NH"]
    )
    processes: dict[int, str] = {}
    for line in listing.splitlines():
        cells = [cell.strip('"') for cell in line.strip().split('","')]
        if len(cells) >= 2 and cells[0].upper() == "WINWORD.EXE":
            try:
                processes[int(cells[1])] = cells[-1]
            except ValueError:
                pass
    return processes


def winword_pids() -> set[int]:
    return set(winword_processes())


# Window titles of hidden automation Word instances (no real document
# window): tasklist reports "N/A" or Word's internal helper window.
HIDDEN_WORD_TITLES = ("", "N/A", "HardwareMonitorWindow")


def user_word_running() -> bool:
    """True when a Word with a visible document window (the user's) runs."""
    return any(
        title not in HIDDEN_WORD_TITLES
        for title in winword_processes().values()
    )


def kill_winword(pids: set[int]) -> None:
    """Force-kill the given WINWORD processes, sparing windowed (user) ones."""
    import subprocess

    for pid, title in winword_processes().items():
        if pid in pids and title in HIDDEN_WORD_TITLES:
            subprocess.run(
                ["taskkill", "/PID", str(pid), "/T", "/F"],
                capture_output=True,
            )


def make_pdf(docx_path: Path, pdf_path: Path, attempts: int = 3) -> bool:
    """Render the .docx to PDF (WPS Writer, or Word); False on failure.

    Office COM occasionally dies with a native access violation (0xC0000005)
    that would kill the whole script and lose the money already spent on
    editing, so the conversion runs in a disposable child process. A Word
    instance left behind by a crashed attempt is killed before the retry -
    but only Word processes that appeared during the attempt, so a Word the
    user is working in is never touched.
    """
    import subprocess

    # Hidden Word instances surviving from earlier crashes make new
    # conversions fail too ("call rejected" avalanche) - clean them first.
    # Irrelevant when WPS does the conversion, so skip in that case.
    if wps_progid() is None:
        kill_winword(winword_pids())

    last_note = ""
    for attempt in range(1, attempts + 1):
        before = winword_pids()
        try:
            child = subprocess.run(
                [
                    sys.executable,
                    str(Path(__file__).resolve()),
                    "--convert-pdf",
                    str(docx_path),
                    str(pdf_path),
                ],
                capture_output=True,
                text=True,
                timeout=300,
            )
            failed = child.returncode != 0
            last_note = (
                (child.stderr or child.stdout).strip().splitlines()[-1]
                if failed and (child.stderr or child.stdout).strip()
                else f"exit code {child.returncode}"
            )
        except subprocess.TimeoutExpired:
            failed, last_note = True, "conversion timed out"
        if not failed and pdf_path.is_file():
            return True
        # Clean up Word instances left by the failed attempt; a windowed
        # Word (the user's, even one opened during the attempt) is spared.
        kill_winword(winword_pids() - before)
        if attempt < attempts:
            time.sleep(3)
    print(f"  WARNING: PDF conversion failed: {last_note}", flush=True)
    if wps_progid() is None and user_word_running():
        print(
            "  NOTE: a Word window is open; working in Word can block the "
            "hidden conversion. Close Word and re-run the script - missing "
            "PDFs are then rebuilt from the .docx at no API cost.",
            flush=True,
        )
    return False


# --------------------------------------------------------------------------
# Selection and session


def output_files(lang_dir: Path, stem: str) -> dict[str, Path]:
    out = lang_dir / OUTPUT_DIRNAME
    return {
        "md": out / f"{stem}.md",
        "docx": out / f"{stem}.docx",
        "pdf": out / f"{stem}.pdf",
    }


def transcripts_ready(lang_dir: Path, stem: str) -> bool:
    return (lang_dir / f"{stem}.srt").is_file()


def is_processed(
    playlist_dir: Path,
    stem: str,
    lang_codes: list[str],
    *,
    annotate: bool = False,
    orig_code: str = "EN",
) -> bool:
    """md + docx in every language OUTPUT folder (pdf is best-effort);
    with --annotate also the annotation .txt files."""
    docs_done = all(
        output_files(playlist_dir / code, stem)[kind].is_file()
        for code in lang_codes
        for kind in ("md", "docx")
    )
    if not docs_done:
        return False
    if annotate:
        return all(
            annotation_path(playlist_dir, code, stem).is_file()
            for code in annotation_languages(orig_code)
        )
    return True


# --------------------------------------------------------------------------
# Editing cache and cost estimate


def section_cache_key(section: Section) -> str:
    return f"{section.heading} @ {section.start:.1f}"


def load_edit_cache(cache_path: Path, stem: str) -> dict[str, dict]:
    try:
        data = json.loads(cache_path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return {}
    return data.get("sections") or {} if data.get("video") == stem else {}


def section_to_cache(section: Section) -> dict:
    return {
        "heading_en": section.heading_en,
        "intro": section.intro,
        "subsections": section.subsections,
    }


def section_from_cache(section: Section, entry: dict) -> None:
    section.heading_en = entry.get("heading_en") or section.heading
    section.intro = entry.get("intro") or {}
    section.subsections = entry.get("subsections") or []


def save_edit_cache(
    cache_path: Path, stem: str, cache: dict[str, dict]
) -> None:
    cache_path.write_text(
        json.dumps(
            {"video": stem, "model": MODEL, "sections": cache},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def estimate_tokens(text: str, code: str) -> int:
    ratio = CHARS_PER_TOKEN_EN if code == "EN" else CHARS_PER_TOKEN_OTHER
    return int(len(text) / ratio)


def estimate_cost(
    pending: list[tuple[Section, str, str]], orig_code: str
) -> tuple[int, int, float]:
    """(prompt tokens, completion tokens, USD) for the sections still to edit.

    The prompt carries both transcripts plus a fixed per-section overhead;
    the completion re-emits the edited text in both languages, so it is
    close to the size of the transcripts themselves.
    """
    prompt = completion = 0
    for _, orig_text, en_text in pending:
        text_tokens = (
            estimate_tokens(orig_text, orig_code)
            + estimate_tokens(en_text, "EN")
        )
        prompt += SECTION_OVERHEAD_TOKENS + text_tokens
        completion += text_tokens
    cost = (
        prompt * USD_PER_MTOKEN_PROMPT
        + completion * USD_PER_MTOKEN_COMPLETION
    ) / 1_000_000
    return prompt, completion, cost


def process_video(
    video: Path,
    *,
    playlist_dir: Path,
    slides_folder: Path | None,
    api_key: str,
    course: str,
    orig_code: str,
    lang_codes: list[str],
    doc_template: Path | None,
    pdf_template: Path | None,
    auto_yes: bool = False,
    annotate: bool = False,
) -> bool:
    """Process one video; False when the user declined the cost estimate.

    With slides (slides_folder holds slides.json) the document structure
    comes from the slides; otherwise (--no-slides) the title is built in the
    summary format and the sections come from the video timecodes, both
    taken from INFO/<stem>.json.
    """
    slides_data = None
    if slides_folder is not None and (slides_folder / RESULT_FILENAME).is_file():
        slides_data = json.loads(
            (slides_folder / RESULT_FILENAME).read_text(encoding="utf-8")
        )

    segments = {
        code: parse_srt(playlist_dir / code / f"{video.stem}.srt")
        for code in lang_codes
    }
    video_end = max(
        (segs[-1].end for segs in segments.values() if segs), default=0.0
    )

    if slides_data is not None:
        document = slides_data.get("document") or {}
        doc_title = document.get("title") or video.stem
        meta = document.get("meta") or {}
        meta_lines = [
            f"{label}: {value}"
            for label, value in (
                ("Date", meta.get("date")),
                ("Recorded by", meta.get("recorded_by")),
                ("Language", meta.get("language")),
            )
            if value
        ]
        sections = build_sections(slides_data, video_end)
        cache_path = slides_folder / EDIT_CACHE_FILENAME
    else:
        info = load_video_info(playlist_dir, video.stem)
        doc_title, meta_lines = info_document_header(info, video.stem)
        sections = build_sections_from_info(info, video_end, orig_code)
        info_dir = playlist_dir / INFO_DIRNAME
        info_dir.mkdir(parents=True, exist_ok=True)
        cache_path = info_dir / f"{video.stem}.{EDIT_CACHE_FILENAME}"

    texts: dict[str, tuple[str, str]] = {}
    for section in sections:
        orig_text = (
            section_text(segments[orig_code], section.start, section.end)
            if orig_code != "EN"
            else ""
        )
        en_text = (
            section_text(segments["EN"], section.start, section.end)
            if "EN" in segments
            else ""
        )
        texts[section_cache_key(section)] = (orig_text, en_text)

    cache = load_edit_cache(cache_path, video.stem)
    pending = [
        (section, *texts[section_cache_key(section)])
        for section in sections
        if section_cache_key(section) not in cache
    ]

    prompt_est, completion_est, cost_est = estimate_cost(pending, orig_code)
    annotation_cached = bool(
        (cache.get(ANNOTATION_CACHE_KEY) or {}).get("texts")
    )
    if annotate and not annotation_cached:
        # One extra call: the whole edited original text in the prompt,
        # 200-250 words per requested language in the completion.
        source = orig_code if orig_code != "EN" else "EN"
        full_len = sum(len(t[0] if source != "EN" else t[1])
                       for t in texts.values())
        ann_prompt = min(
            int(
                full_len / (CHARS_PER_TOKEN_EN if source == "EN"
                            else CHARS_PER_TOKEN_OTHER)
            ),
            ANNOTATION_MAX_PROMPT_TOKENS,
        ) + 400
        ann_completion = 400 * len(annotation_languages(orig_code))
        prompt_est += ann_prompt
        completion_est += ann_completion
        cost_est += (
            ann_prompt * USD_PER_MTOKEN_PROMPT
            + ann_completion * USD_PER_MTOKEN_COMPLETION
        ) / 1_000_000
    print(
        f"  Sections: {len(sections)} total, "
        f"{len(sections) - len(pending)} already edited (cached), "
        f"{len(pending)} to edit via the API"
        + (
            f"; annotation ({'/'.join(annotation_languages(orig_code))}): "
            + ("cached" if annotation_cached else "to generate")
            if annotate
            else ""
        )
        + ".",
        flush=True,
    )
    print(
        f"  Estimated cost: ~{prompt_est} prompt + ~{completion_est} "
        f"completion tokens -> ~${cost_est:.2f} ({MODEL}).",
        flush=True,
    )
    if not auto_yes:
        print("  Create the final documents for this video? (y/n)",
              flush=True)
        try:
            answer = input().strip().lower()
        except EOFError:
            answer = ""
        if answer not in ("y", "yes"):
            return False

    usage: dict[str, int] = {}
    for index, section in enumerate(sections, start=1):
        key = section_cache_key(section)
        minutes = (section.end - section.start) / 60.0
        if key in cache:
            print(
                f"  [{index}/{len(sections)}] {section.heading} "
                f"({minutes:.1f} min): cached.",
                flush=True,
            )
            section_from_cache(section, cache[key])
            continue
        print(
            f"  [{index}/{len(sections)}] {section.heading} "
            f"({minutes:.1f} min)...",
            flush=True,
        )
        if section.end - section.start > EDIT_CHUNK_TRIGGER_SECONDS:
            edit_section_in_chunks(
                api_key,
                section,
                course=course,
                doc_title=doc_title,
                orig_code=orig_code,
                segments=segments,
                usage=usage,
            )
        else:
            orig_text, en_text = texts[key]
            edit_section(
                api_key,
                section,
                course=course,
                doc_title=doc_title,
                orig_code=orig_code,
                orig_text=orig_text,
                en_text=en_text,
                usage=usage,
            )
        # Persist after every section: a crash later in the run (e.g. during
        # the PDF stage) must not lose paid editing results.
        cache[key] = section_to_cache(section)
        save_edit_cache(cache_path, video.stem, cache)

    annotations: dict[str, str] = {}
    if annotate:
        annotations = (cache.get(ANNOTATION_CACHE_KEY) or {}).get("texts") or {}
        needed = annotation_languages(orig_code)
        if not all(code in annotations for code in needed):
            print(
                f"  Annotation ({'/'.join(needed)}): generating...",
                flush=True,
            )
            annotations = generate_annotation(
                api_key,
                doc_title=doc_title,
                orig_code=orig_code,
                sections=sections,
                usage=usage,
            )
            cache[ANNOTATION_CACHE_KEY] = {"texts": annotations}
            save_edit_cache(cache_path, video.stem, cache)
        else:
            print(
                f"  Annotation ({'/'.join(needed)}): cached.", flush=True
            )
        # Write the annotation files before the docs stage: a failure there
        # (e.g. PDF conversion) must not lose the generated annotation.
        for code in annotation_languages(orig_code):
            text = annotations.get(code)
            if not text:
                print(
                    f"  WARNING: no {code} annotation was generated.",
                    flush=True,
                )
                continue
            path = annotation_path(playlist_dir, code, video.stem)
            write_annotation(path, doc_title, text)
            print(
                f"  Saved: {path.relative_to(playlist_dir)}", flush=True
            )

    for code in lang_codes:
        files = output_files(playlist_dir / code, video.stem)
        markdown = build_markdown(doc_title, meta_lines, sections, code)
        files["md"].parent.mkdir(parents=True, exist_ok=True)
        files["md"].write_text(markdown, encoding="utf-8")
        make_docx(
            doc_title, meta_lines, sections, code, doc_template, files["docx"]
        )
        if pdf_template is not None:
            tmp_docx = files["pdf"].with_suffix(".pdf.tmp.docx")
            make_docx(
                doc_title, meta_lines, sections, code, pdf_template, tmp_docx
            )
            if make_pdf(tmp_docx, files["pdf"]):
                tmp_docx.unlink(missing_ok=True)
        else:
            make_pdf(files["docx"], files["pdf"])
        for kind in ("md", "docx", "pdf"):
            if files[kind].is_file():
                print(
                    f"  Saved: {files[kind].relative_to(playlist_dir)}",
                    flush=True,
                )
    actual_cost = (
        usage.get("prompt_tokens", 0) * USD_PER_MTOKEN_PROMPT
        + usage.get("completion_tokens", 0) * USD_PER_MTOKEN_COMPLETION
    ) / 1_000_000
    print(
        f"  Tokens used: {usage.get('prompt_tokens', 0)} prompt + "
        f"{usage.get('completion_tokens', 0)} completion "
        f"(~${actual_cost:.2f}).",
        flush=True,
    )
    return True


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create final edited documents from transcripts + slides."
    )
    parser.add_argument(
        "channel_folder",
        help="Channel ref under _channels/ (e.g. _Autotesting or "
        "AI_for_Game_Design\\_BuildingAeon)",
    )
    parser.add_argument(
        "playlist_folder",
        help="Playlist folder name under <channel>/_playlists (e.g. lectures)",
    )
    parser.add_argument(
        "--next",
        dest="next_count",
        type=int,
        default=1,
        metavar="N",
        help="How many videos to process this session (default: 1)",
    )
    parser.add_argument(
        "--doc",
        type=Path,
        default=None,
        metavar="DOCX",
        help=(
            "A .docx file whose styles are used for the generated .docx "
            "(default: built-in styles)"
        ),
    )
    parser.add_argument(
        "--pdf",
        type=Path,
        default=None,
        metavar="DOCX",
        help=(
            "A .docx style template used only for the PDF rendering "
            "(default: the PDF is rendered from the generated .docx)"
        ),
    )
    parser.add_argument(
        "--orig-only",
        action="store_true",
        help=(
            "Create the OUTPUT documents only for the original language "
            "(skip the English versions)"
        ),
    )
    parser.add_argument(
        "--no-slides",
        action="store_true",
        help=(
            "Also process videos without slides.json: the document title is "
            "built in the summary format and the ToC comes from the video "
            "timecodes (INFO/<stem>.json saved by transcribe_videos.py); "
            "videos that do have slides.json still use the slides"
        ),
    )
    parser.add_argument(
        "--annotate",
        action="store_true",
        help=(
            "Also create a 200-250 word annotation .txt per video: Russian "
            "original -> Russian annotation in RU/OUTPUT; any other "
            "original -> English annotation in EN/OUTPUT plus a Russian "
            "translation in RU/OUTPUT"
        ),
    )
    parser.add_argument(
        "--video",
        default=None,
        metavar="STEM",
        help=(
            "Process only the video with this file stem (exact name without "
            "extension); overrides --next"
        ),
    )
    parser.add_argument(
        "--yes",
        action="store_true",
        help="Skip the per-video cost confirmation prompt",
    )
    parser.add_argument(
        "--workspace",
        type=Path,
        default=WORKSPACE_ROOT,
        help="Workspace root (default: parent of src/)",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    if args.next_count < 1:
        raise SystemExit("--next must be a positive number")
    for flag, path in (("--doc", args.doc), ("--pdf", args.pdf)):
        if path is not None and not path.is_file():
            raise SystemExit(f"{flag} template not found: {path}")

    try:
        channel_dir = require_channel_ref(
            channels_dir(args.workspace), args.channel_folder
        )
    except FileNotFoundError as exc:
        raise SystemExit(str(exc)) from exc

    playlist_dir = channel_dir / "_playlists" / args.playlist_folder
    if not playlist_dir.is_dir():
        raise SystemExit(f"Playlist folder not found: {playlist_dir}")

    slides_dir = playlist_dir / SLIDES_DIRNAME
    course = channel_dir.name.lstrip("_")
    orig_code = detect_original_code(playlist_dir)
    lang_codes = (
        [orig_code]
        if args.orig_only or orig_code == "EN"
        else [orig_code, "EN"]
    )

    # SLIDES/ folder keys are derived from the set of *local media* stems -
    # transcript-only videos (remote transcription) must not change them.
    media_videos = list_videos(playlist_dir)
    keys = short_slide_keys([video.stem for video in media_videos])
    videos = media_videos
    if args.no_slides:
        known = {video.stem for video in media_videos}
        extra = [
            playlist_dir / f"{srt.stem}.mp4"
            for srt in sorted((playlist_dir / orig_code).glob("*.srt"))
            if srt.stem not in known
        ]
        videos = sorted(media_videos + extra, key=lambda path: path.name.lower())
    if not videos:
        raise SystemExit(
            f"No video/audio files{' or transcripts' if args.no_slides else ''}"
            f" found in {playlist_dir}"
        )
    api_key = read_api_key(args.workspace)

    def slides_folder_for(video: Path) -> Path | None:
        if video.stem not in keys:
            return None
        return slides_out_dir(slides_dir, video, keys)

    def inputs_ready(video: Path) -> bool:
        if not all(
            transcripts_ready(playlist_dir / code, video.stem)
            for code in lang_codes
        ):
            return False
        folder = slides_folder_for(video)
        has_slides = folder is not None and (folder / RESULT_FILENAME).is_file()
        return has_slides or args.no_slides

    if args.video is not None:
        videos = [video for video in videos if video.stem == args.video]
        if not videos:
            raise SystemExit(
                f"--video: no video named {args.video!r} in {playlist_dir}"
            )

    eligible = [video for video in videos if inputs_ready(video)]
    pending = [
        video
        for video in eligible
        if not is_processed(
            playlist_dir,
            video.stem,
            lang_codes,
            annotate=args.annotate,
            orig_code=orig_code,
        )
    ]
    session = pending[: args.next_count] if args.video is None else pending

    print(f"Playlist folder: {playlist_dir}", flush=True)
    print(
        f"Original language: {orig_code}; outputs: "
        + ", ".join(f"{code}/{OUTPUT_DIRNAME}/" for code in lang_codes),
        flush=True,
    )
    ready_note = (
        "with transcripts" if args.no_slides else "with transcripts + slides.json"
    )
    print(
        f"Videos: {len(videos)} total, {len(eligible)} {ready_note}, "
        f"{len(eligible) - len(pending)} already processed, "
        f"{len(session)} in this session (--next {args.next_count}).",
        flush=True,
    )
    # Self-healing: the PDF stage is best-effort (Word may fail or be
    # blocked), so processed videos may lack PDFs. Rebuild them from the
    # existing .docx - this costs no API tokens.
    if args.pdf is None:
        for video in eligible:
            if video in session:
                continue
            for code in lang_codes:
                files = output_files(playlist_dir / code, video.stem)
                if files["docx"].is_file() and not files["pdf"].is_file():
                    print(
                        "PDF backfill: "
                        f"{files['pdf'].relative_to(playlist_dir)}...",
                        flush=True,
                    )
                    make_pdf(files["docx"], files["pdf"])

    if not session:
        print("Nothing to do: final documents exist for every ready video.",
              flush=True)
        return 0

    done = 0
    for index, video in enumerate(session, start=1):
        print(f"[{index}/{len(session)}] {video.name}", flush=True)
        if not process_video(
            video,
            playlist_dir=playlist_dir,
            slides_folder=slides_folder_for(video),
            api_key=api_key,
            course=course,
            orig_code=orig_code,
            lang_codes=lang_codes,
            doc_template=args.doc,
            pdf_template=args.pdf,
            auto_yes=args.yes,
            annotate=args.annotate,
        ):
            print("Session stopped: not confirmed.", flush=True)
            break
        done += 1

    print(f"Session done: {done} video(s).", flush=True)
    return 0


if __name__ == "__main__":
    if len(sys.argv) == 4 and sys.argv[1] == "--convert-pdf":
        # Internal entry point: one PDF conversion in a disposable child
        # process (see make_pdf).
        raise SystemExit(
            convert_docx_to_pdf(Path(sys.argv[2]), Path(sys.argv[3]))
        )
    raise SystemExit(main())
