"""
Assemble page Markdown into OUTPUT/Book_*.md / .docx / .pdf.

Two modes:

1. Single-volume (default for short books / one-level TOC), e.g. Wojna
   Futbolowa: reads pages_text/ (via pages_expanded.txt when present), joins
   cross-page hyphenation, builds H2 sections, inserts a generated TOC, and
   writes OUTPUT/Book_<LANG>.*.

2. Multi-volume for large books with a two-level TOC (Part → Chapter): when
   _books/<book>/assemble.settings.txt is present, SPLIT_BY_PARTS is auto/yes
   and the book has more than PART_PAGE_THRESHOLD pages, one set of files is
   written per part: OUTPUT/Book_<N>_<LANG>.* (plus a final volume for
   appendices when BACK_START is set).

Usage (from workspace root):
  python src\\book_ocr\\assemble_book.py
  python src\\book_ocr\\assemble_book.py --book Wojna_Futbolowa --lang PL
  python src\\book_ocr\\assemble_book.py --book Artificial_Intelligence_2006
  python src\\book_ocr\\assemble_book.py --book Artificial_Intelligence_2006 --parts 1 9
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from book_paths import DEFAULT_BOOK, WORKSPACE_ROOT, book_dir, pages_file, text_dir

LINE_WIDTH_PAD = 3


def clean_paragraph(text: str) -> str:
    text = text.replace("\u00ad", "")
    text = re.sub(r"[ \t]+", " ", text).strip()
    text = re.sub(r"(?<=\w)-\s+(?=\w)", "", text)
    return text


def justify_line(words: list[str], width: int) -> str:
    if len(words) == 1:
        return words[0]
    base_len = sum(len(w) for w in words) + (len(words) - 1)
    need = width - base_len
    if need <= 0:
        return " ".join(words)
    gaps = len(words) - 1
    extras = [0] * gaps
    i = 0
    while need > 0:
        extras[i % gaps] += 1
        need -= 1
        i += 1
    parts: list[str] = []
    for idx, word in enumerate(words):
        parts.append(word)
        if idx < gaps:
            parts.append(" " * (1 + extras[idx]))
    return "".join(parts)


def wrap_and_justify_paragraph(text: str, width: int) -> str:
    words = text.split()
    if not words:
        return ""
    lines: list[str] = []
    cur: list[str] = []
    cur_len = 0
    for word in words:
        add = len(word) if not cur else len(word) + 1
        if cur and cur_len + add > width:
            lines.append(justify_line(cur, width))
            cur = [word]
            cur_len = len(word)
        else:
            cur.append(word)
            cur_len += add
    if cur:
        lines.append(" ".join(cur))
    return "\n".join(lines)

DEFAULT_LANG = "PL"
DEFAULT_WRAP = 60  # median-ish; +0 already includes pad convention of OCR files
FOREWORD_TITLE = "Przedmowa"
TOC_TITLE = "Spis treści"
AFTERWORD_TITLE = "Nota końcowa"

# Story starts: printed page number -> title (from the book's Spis treści,
# with OCR typos corrected from body text). Used only as section boundaries;
# the document Spis treści is generated from the resulting H2 headings.
STORY_STARTS: list[tuple[int, str]] = [
    (7, "Hotel Metropol"),
    (16, "Bezdomny z Haarlemu"),
    (
        30,
        "Plan książki, która mogłaby zacząć się w tym miejscu "
        "(czyli moje tarapaty nigdy nie spisane)",
    ),
    (34, "Lumumba"),
    (44, "Prezesi"),
    (47, "Ofensywa"),
    (51, "Ciąg dalszy planu książki, która mogłaby zacząć się (itd.)"),
    (66, "Ożenek i wolność"),
    (69, "Parlament Tanganiki w sprawie alimentów"),
    (75, "Będziemy pławić konie we krwi"),
    (92, "Algieria zakrywa twarz"),
    (123, "Spór o sędziego zakończony upadkiem rządu"),
    (130, "Płonące bariery"),
    (137, "Nigeria, lato 66:"),
    (152, "Cd. planu nigdy nie napisanej książki, która mogłaby (itd.)"),
    (162, "Czas najwyższy, abym zaczął pisać następną nigdy nie napisaną książkę"),
    (167, "Wojna futbolowa"),
    (193, "Dostaniesz dziewczynę"),
    (200, "Victoriano Gomez przed kamerami TV"),
    (
        203,
        "Ciąg dalszy czasu najwyższego, czyli planu drugiej nigdy nie "
        "napisanej książki, która (itd.)",
    ),
    (212, "Buty"),
    (217, "Nie będzie raju"),
    (228, "Ogaden, jesień 76:"),
    (
        238,
        "Cd. czasu najwyższego, czyli planu drugiej nigdy nie napisanej "
        "książki, która (itd.)",
    ),
]

FRONT_MATTER_SLUGS = {
    "page_000_cover",
    "page_n1",
    "page_n3",
    "page_n5",
    "page_n7",
}

TITLE_LINES = (
    "Ryszard Kapuściński",
    "Wojna futbolowa",
    "WOJNA FUTBOLOWA",
    "Czytelnik · Warszawa",
    "Czytelnik · Warszawa",
)

SKIP_PARAGRAPH = re.compile(
    r"(?is)^("
    r"blank(\s+page)?.*"
    r"|strona\s+pusta.*"
    r"|\(blank.*"
    r"|\(strona\s+pusta.*"
    r"|digitized\s+by.*"
    r"|boston\s+public\s+library.*"
    r"|withdrawn.*"
    r"|brak\s+(widocznego\s+)?tekstu.*"
    r"|brak\s+widocznego.*"
    r"|\[strona\s+zawiera.*"
    r"|\[środkowa\s+część.*"
    r"|library\s+stamp.*"
    r"|page\s+otherwise\s+blank.*"
    r"|no\s+readable\s+text.*"
    r"|no\s+visible\s+text.*"
    r"|https?://archive\.org/.*"
    r"|kahle/austin.*"
    r")$"
)

YEAR_ONLY = re.compile(r"^\*?(\d{4})\*?$")


@dataclass
class PageBlock:
    slug: str
    heading: str
    page_num: int | None
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class Section:
    title: str
    paragraphs: list[str] = field(default_factory=list)


def slug_from_url(url: str) -> str:
    m = re.search(r"/page/([^/]+)/", url)
    if m:
        leaf = m.group(1)
        if leaf.isdigit():
            return f"page_{int(leaf):03d}"
        return f"page_{leaf}"
    return "page_000_cover"


def load_slugs(path: Path) -> list[str]:
    slugs: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        slugs.append(slug_from_url(line))
    return slugs


def parse_page_num(heading: str) -> int | None:
    m = re.search(r"\bPage\s+(\d+)\b", heading, re.I)
    return int(m.group(1)) if m else None


def is_spis_heading(heading: str) -> bool:
    return "spis" in heading.lower() and "treś" in heading.lower()


def is_junk_paragraph(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    if SKIP_PARAGRAPH.match(t):
        return True
    # Drop OCR notes that wrap the digitize banner into a long paragraph.
    if "Digitized by the Internet Archive" in t:
        return True
    if "archive.org/details/" in t and "Kapuściński" not in t:
        return True
    return False


def is_title_duplicate(text: str) -> bool:
    norm = re.sub(r"\s+", " ", text).strip().casefold()
    for line in TITLE_LINES:
        if norm == line.casefold():
            return True
    # Full title block jammed into one paragraph
    if "kapuściński" in norm and "wojna futbolowa" in norm:
        return True
    return False


def normalize_spaces(text: str) -> str:
    return clean_paragraph(text.replace("\n", " "))


def parse_md_file(path: Path, slug: str) -> list[PageBlock]:
    raw = path.read_text(encoding="utf-8")
    body = re.sub(r"(?m)^<!--.*?-->\s*", "", raw)
    parts = re.split(r"(?m)^(## .+)$", body)
    blocks: list[PageBlock] = []
    i = 0
    while i < len(parts):
        part = parts[i]
        if re.match(r"^## ", part or ""):
            heading = part.strip()
            content = parts[i + 1] if i + 1 < len(parts) else ""
            paragraphs: list[str] = []
            for par in re.split(r"\n\s*\n", content.strip("\n")):
                lines = [ln.strip() for ln in par.splitlines() if ln.strip()]
                lines = [ln for ln in lines if not re.fullmatch(r"\d{1,4}", ln)]
                if not lines:
                    continue
                cleaned = normalize_spaces(" ".join(lines))
                if not cleaned or is_junk_paragraph(cleaned):
                    continue
                if is_title_duplicate(cleaned):
                    continue
                paragraphs.append(cleaned)
            blocks.append(
                PageBlock(
                    slug=slug,
                    heading=heading,
                    page_num=parse_page_num(heading),
                    paragraphs=paragraphs,
                )
            )
            i += 2
            continue
        i += 1
    return blocks


def ends_with_hyphenated_break(text: str) -> bool:
    return bool(re.search(r"[A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż]-$", text.rstrip()))


def looks_like_sentence_end(text: str) -> bool:
    t = text.rstrip()
    if not t:
        return True
    if ends_with_hyphenated_break(t):
        return False
    return bool(re.search(r'[.!?…»"”\']$', t))


def join_across_break(left: str, right: str) -> str:
    left = left.rstrip()
    right = right.lstrip()
    if ends_with_hyphenated_break(left):
        # Ameryka- + nie → Amerykanie
        return left[:-1] + right
    return left + " " + right


def titles_match(a: str, b: str) -> bool:
    def norm(s: str) -> str:
        s = s.casefold()
        s = s.replace("ł", "l").replace("ó", "o")
        s = re.sub(r"[^a-z0-9]+", "", s)
        return s

    na, nb = norm(a), norm(b)
    if not na or not nb:
        return False
    return na == nb or na.startswith(nb[:20]) or nb.startswith(na[:20])


def strip_leading_title(paragraphs: list[str], title: str) -> list[str]:
    if not paragraphs:
        return paragraphs
    first = paragraphs[0]
    if titles_match(first, title) or first.casefold() == title.casefold():
        return paragraphs[1:]
    # Title may be first line only of a short paragraph
    if len(first) < 80 and titles_match(first, title):
        return paragraphs[1:]
    return paragraphs


def build_sections(slugs: list[str], text_root: Path) -> tuple[str, list[Section]]:
    """Return (book_title_h1_text, sections)."""
    all_blocks: list[PageBlock] = []
    for slug in slugs:
        path = text_root / f"{slug}.md"
        if not path.exists():
            print(f"WARN: missing {path.name}", flush=True)
            continue
        all_blocks.extend(parse_md_file(path, slug))

    h1 = "Ryszard Kapuściński\nWojna futbolowa\nCzytelnik · Warszawa"

    front_paras: list[str] = []
    main_blocks: list[PageBlock] = []
    after_paras: list[str] = []
    phase = "front"  # front | main | after

    for block in all_blocks:
        if phase == "front":
            if block.slug in FRONT_MATTER_SLUGS:
                front_paras.extend(block.paragraphs)
                continue
            phase = "main"

        if phase == "main":
            if is_spis_heading(block.heading):
                phase = "after"
                continue
            # Also stop if body itself is only "Spis treści"
            if block.paragraphs and block.paragraphs[0].casefold().startswith("spis treści"):
                phase = "after"
                continue
            main_blocks.append(block)
            continue

        # after
        after_paras.extend(block.paragraphs)

    sections: list[Section] = []
    if front_paras:
        sections.append(Section(FOREWORD_TITLE, front_paras))

    # Placeholder for Spis treści — filled after all H2 titles are known.
    toc_index = len(sections)
    sections.append(Section(TOC_TITLE, []))

    story_iter = iter(STORY_STARTS)
    next_story = next(story_iter, None)
    current: Section | None = None
    pending_join = False

    def open_story(title: str) -> None:
        nonlocal current, pending_join
        current = Section(title, [])
        sections.append(current)
        pending_join = False

    def ensure_stories_due(page_num: int | None) -> None:
        nonlocal next_story
        while next_story is not None and page_num is not None and page_num >= next_story[0]:
            open_story(next_story[1])
            next_story = next(story_iter, None)

    for block in main_blocks:
        if block.page_num is not None:
            ensure_stories_due(block.page_num)
        elif current is None and next_story is not None:
            # Unnumbered leaf before first story: open first story when main starts.
            open_story(next_story[1])
            next_story = next(story_iter, None)

        if current is None:
            continue

        paras = list(block.paragraphs)
        if paras:
            paras = strip_leading_title(paras, current.title)

        for par in paras:
            if pending_join and current.paragraphs:
                current.paragraphs[-1] = join_across_break(current.paragraphs[-1], par)
                pending_join = False
            else:
                current.paragraphs.append(par)

        if paras:
            last = current.paragraphs[-1]
            pending_join = not looks_like_sentence_end(last)

    if after_paras:
        sections.append(Section(AFTERWORD_TITLE, after_paras))

    # Build TOC entries from H2 titles except the TOC section itself.
    toc_lines = [
        s.title
        for i, s in enumerate(sections)
        if i != toc_index and s.title not in {TOC_TITLE}
    ]
    sections[toc_index].paragraphs = toc_lines
    return h1, sections


def md_anchor(title: str) -> str:
    a = title.casefold()
    a = (
        a.replace("ą", "a")
        .replace("ć", "c")
        .replace("ę", "e")
        .replace("ł", "l")
        .replace("ń", "n")
        .replace("ó", "o")
        .replace("ś", "s")
        .replace("ź", "z")
        .replace("ż", "z")
    )
    # Keep letters of any script (Cyrillic chapter titles etc.).
    a = re.sub(r"[^\w]+", "-", a, flags=re.UNICODE).strip("-")
    return a or "section"


def format_body_paragraph(text: str, width: int) -> str:
    if YEAR_ONLY.match(text.strip()):
        return text.strip().strip("*")
    if len(text) < width * 0.6:
        return text
    return wrap_and_justify_paragraph(text, width)


def write_markdown(
    path: Path,
    h1: str,
    sections: list[Section],
    width: int,
    *,
    toc_title: str | None = None,
) -> None:
    toc_name = toc_title or TOC_TITLE
    lines: list[str] = [
        f"<!-- wrap_width: {width}; assembled from pages_text -->",
        "",
        f"# {h1.replace(chr(10), ' / ')}",
        "",
    ]
    for sec in sections:
        lines.append(f"## {sec.title}")
        lines.append("")
        if sec.title == toc_name:
            for title in sec.paragraphs:
                lines.append(f"- [{title}](#{md_anchor(title)})")
            lines.append("")
            continue
        for par in sec.paragraphs:
            lines.append(format_body_paragraph(par, width))
            lines.append("")
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _add_field_runs(paragraph, instr: str, placeholder: str = "1") -> None:
    """Insert a Word field (PAGE, PAGEREF, …) into an existing paragraph."""
    r_begin = paragraph.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fc_begin)

    r_instr = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r_instr._r.append(instr_el)

    r_sep = paragraph.add_run()
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fc_sep)

    paragraph.add_run(placeholder)

    r_end = paragraph.add_run()
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fc_end)


def add_page_numbers(doc: Document) -> None:
    """Centered PAGE field in the footer of every section."""
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p.text = ""
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_field_runs(p, "PAGE", "1")


def bookmark_name(index: int) -> str:
    return f"sec_{index:03d}"


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str) -> None:
    """Hyperlink to a bookmark inside the same document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run.append(t_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        p.add_run(part.strip())


def add_h2_with_bookmark(
    doc: Document, text: str, name: str | None, bookmark_id: int | None
):
    p = doc.add_paragraph(text, style="Heading 2")
    if name is not None and bookmark_id is not None:
        add_bookmark(p, name, bookmark_id)
    return p


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_toc_entry(doc: Document, title: str, anchor: str) -> None:
    """One Spis treści line: hyperlink + dot leaders + PAGEREF page number."""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(2)
    # Right-aligned tab with dotted leader near the right margin.
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    add_internal_hyperlink(p, anchor, title)
    p.add_run("\t")
    _add_field_runs(p, f"PAGEREF {anchor} \\h", "?")


def write_docx(
    path: Path,
    template: Path,
    h1: str,
    sections: list[Section],
    *,
    toc_title: str | None = None,
) -> Path:
    toc_name = toc_title or TOC_TITLE
    doc = Document(str(template)) if template.exists() else Document()
    clear_document_body(doc)
    add_page_numbers(doc)
    add_h1(doc, h1)

    # Pre-assign bookmark names for every section except the TOC itself.
    bookmarks: dict[str, str] = {}
    bookmark_id = 1
    for sec in sections:
        if sec.title == toc_name:
            continue
        name = bookmark_name(bookmark_id)
        bookmarks[sec.title] = name
        bookmark_id += 1

    next_id = 1
    for sec in sections:
        if sec.title == toc_name:
            add_h2_with_bookmark(doc, sec.title, None, None)
            for title in sec.paragraphs:
                anchor = bookmarks.get(title)
                if not anchor:
                    continue
                add_toc_entry(doc, title, anchor)
            continue

        name = bookmarks[sec.title]
        add_h2_with_bookmark(doc, sec.title, name, next_id)
        next_id += 1
        for par in sec.paragraphs:
            add_body(doc, par)

    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    doc.save(str(tmp))
    try:
        tmp.replace(path)
        return path
    except PermissionError:
        alt = path.with_name(path.stem + "_new" + path.suffix)
        tmp.replace(alt)
        print(
            f"WARN: cannot overwrite {path.name} (file is open); wrote {alt.name}",
            flush=True,
        )
        return alt


def write_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Export PDF via Word, updating PAGE/PAGEREF fields first.

    Returns the path actually written (may be *_new.pdf if the target is locked).
    """
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    docx_abs = str(docx_path.resolve())
    pdf_abs = str(pdf_path.resolve())
    alt_pdf = pdf_path.with_name(pdf_path.stem + "_new" + pdf_path.suffix)
    doc = word.Documents.Open(docx_abs)
    written = pdf_path
    try:
        doc.Fields.Update()
        try:
            doc.Repaginate()
        except Exception:
            pass
        doc.Fields.Update()
        # Save resolved fields back into the .docx (may fail if read-only).
        try:
            doc.Save()
        except Exception:
            pass
        try:
            doc.ExportAsFixedFormat(
                OutputFileName=pdf_abs,
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,
                CreateBookmarks=1,
            )
        except Exception:
            doc.ExportAsFixedFormat(
                OutputFileName=str(alt_pdf.resolve()),
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,
                CreateBookmarks=1,
            )
            written = alt_pdf
    finally:
        doc.Close(False)
        word.Quit()
    return written


def write_volume(
    out_dir: Path,
    stem: str,
    h1: str,
    sections: list[Section],
    wrap_width: int,
    template: Path,
    *,
    no_pdf: bool,
    toc_title: str | None = None,
) -> int:
    """Write md/docx/(pdf) for one volume. Returns 0 on success."""
    md_path = out_dir / f"{stem}.md"
    docx_path = out_dir / f"{stem}.docx"
    pdf_path = out_dir / f"{stem}.pdf"

    write_markdown(md_path, h1, sections, wrap_width, toc_title=toc_title)
    print(f"Wrote {md_path}", flush=True)

    docx_written = write_docx(
        docx_path, template, h1, sections, toc_title=toc_title
    )
    print(f"Wrote {docx_written}", flush=True)

    if not no_pdf:
        try:
            written_pdf = write_pdf(docx_written, pdf_path)
            print(f"Wrote {written_pdf}", flush=True)
        except Exception as exc:
            print(f"PDF export failed for {stem}: {exc}", file=sys.stderr)
            return 1
    return 0


def resolve_slugs(book: str, text_root: Path) -> list[str]:
    """pages_expanded.txt when present; otherwise every page_*.md in order."""
    pf = pages_file(book)
    if pf.is_file():
        return load_slugs(pf)
    pages = []
    for path in text_root.glob("page_*.md"):
        m = re.fullmatch(r"page_(\d+)", path.stem)
        if m:
            pages.append((int(m.group(1)), path.stem))
    pages.sort(key=lambda x: x[0])
    return [stem for _, stem in pages]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Assemble Book_XX.md/.docx/.pdf from OCR pages")
    p.add_argument("--book", default=DEFAULT_BOOK)
    p.add_argument("--lang", default=None, help="ISO 639-1 language code for filenames")
    p.add_argument("--wrap-width", type=int, default=None)
    p.add_argument(
        "--template",
        type=Path,
        default=WORKSPACE_ROOT / "_books" / "sample_for_book_in_one_file.docx",
    )
    p.add_argument("--no-pdf", action="store_true")
    p.add_argument(
        "--parts",
        nargs="+",
        type=int,
        default=None,
        help="In multi-volume mode, assemble only these part numbers "
             "(e.g. --parts 1 9). Default: all parts.",
    )
    p.add_argument(
        "--force-single",
        action="store_true",
        help="Ignore assemble.settings.txt split and build one Book_LANG.*",
    )
    return p.parse_args()


def main() -> int:
    args = parse_args()
    book_root = book_dir(args.book)
    text_root = text_dir(args.book)
    out_dir = book_root / "OUTPUT"
    out_dir.mkdir(parents=True, exist_ok=True)

    settings_path = book_root / "assemble.settings.txt"
    if settings_path.is_file() and not args.force_single:
        from assemble_structure import (  # local import avoids cycle at load
            discover_page_files,
            iter_volumes,
            load_structure,
        )

        structure = load_structure(settings_path)
        page_count = len(discover_page_files(text_root))
        if structure.should_split(page_count):
            lang = (args.lang or structure.lang).upper()
            want = set(args.parts) if args.parts else None
            rc = 0
            built = 0
            for num, h1, sections, width in iter_volumes(structure, text_root):
                if want is not None and num not in want:
                    continue
                wrap = args.wrap_width if args.wrap_width is not None else width
                stem = f"Book_{num}_{lang}"
                print(f"=== volume {stem}: {h1.splitlines()[-1]} ===", flush=True)
                step = write_volume(
                    out_dir,
                    stem,
                    h1,
                    sections,
                    wrap,
                    args.template,
                    no_pdf=args.no_pdf,
                    toc_title=structure.toc_title,
                )
                if step != 0:
                    rc = step
                built += 1
            if built == 0:
                print("Nothing to assemble: --parts matched no volumes.", flush=True)
                return 2
            print(
                f"Done. volumes={built} pages_text={page_count} out={out_dir}",
                flush=True,
            )
            print("API cost: $0 (local assembly).", flush=True)
            return rc

    # Single-volume path (Wojna Futbolowa and other one-level books).
    lang = (args.lang or DEFAULT_LANG).upper()
    wrap = (
        args.wrap_width
        if args.wrap_width is not None
        else DEFAULT_WRAP + LINE_WIDTH_PAD
    )
    slugs = resolve_slugs(args.book, text_root)
    if not slugs:
        print(f"ERROR: no pages to assemble in {text_root}", file=sys.stderr)
        return 2
    h1, sections = build_sections(slugs, text_root)
    stem = f"Book_{lang}"
    rc = write_volume(
        out_dir, stem, h1, sections, wrap, args.template, no_pdf=args.no_pdf
    )
    story_count = sum(
        1
        for s in sections
        if s.title not in {FOREWORD_TITLE, TOC_TITLE, AFTERWORD_TITLE}
    )
    print(
        f"Done. sections={len(sections)} stories={story_count} "
        f"wrap_width={wrap} out={out_dir}",
        flush=True,
    )
    print("API cost: $0 (local assembly).", flush=True)
    return rc


if __name__ == "__main__":
    raise SystemExit(main())
